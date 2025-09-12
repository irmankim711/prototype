#!/usr/bin/env python3
"""
Test DOCX Report Generation System
Tests the complete workflow from file upload to report generation
"""

import os
import sys
import requests
import json
from pathlib import Path

# Configuration
BASE_URL = "http://localhost:5000"
API_BASE = f"{BASE_URL}/api"

def test_file_upload():
    """Test file upload endpoint"""
    print("🔧 Testing File Upload...")
    
    # Create a simple test DOCX file (or use existing one)
    test_file_path = "test_document.docx"
    
    if not os.path.exists(test_file_path):
        print(f"⚠️ Test file not found: {test_file_path}")
        print("   Please create a test DOCX file or update the path")
        return None
    
    try:
        with open(test_file_path, 'rb') as f:
            files = {'file': f}
            data = {'title': 'Test Document'}
            
            response = requests.post(f"{API_BASE}/reports/upload/file", files=files, data=data)
            
            if response.status_code == 201:
                result = response.json()
                print(f"✅ File uploaded successfully")
                print(f"   File path: {result.get('file_path')}")
                print(f"   File size: {result.get('file_size')} bytes")
                return result
            else:
                print(f"❌ File upload failed: {response.status_code}")
                print(f"   Response: {response.text}")
                return None
                
    except Exception as e:
        print(f"❌ File upload error: {str(e)}")
        return None

def test_report_generation(upload_result):
    """Test report generation with uploaded file"""
    print("\n🔧 Testing Report Generation...")
    
    if not upload_result:
        print("❌ No upload result to test with")
        return None
    
    try:
        report_data = {
            "title": "Test DOCX Report",
            "description": "Generated from uploaded DOCX file",
            "latex_file_path": upload_result['file_path'],
            "config": {
                "output_formats": ["docx", "pdf", "excel"],
                "template": "default"
            }
        }
        
        response = requests.post(f"{API_BASE}/reports/generate/latex", json=report_data)
        
        if response.status_code == 200:
            result = response.json()
            print(f"✅ Report generated successfully")
            print(f"   Report ID: {result.get('id')}")
            print(f"   Status: {result.get('status')}")
            print(f"   Type: {result.get('report_type')}")
            return result
        else:
            print(f"❌ Report generation failed: {response.status_code}")
            print(f"   Response: {response.text}")
            return None
            
    except Exception as e:
        print(f"❌ Report generation error: {str(e)}")
        return None

def test_report_status(report_id):
    """Test getting report status"""
    print(f"\n🔧 Testing Report Status for ID: {report_id}...")
    
    try:
        response = requests.get(f"{API_BASE}/reports/{report_id}/status")
        
        if response.status_code == 200:
            result = response.json()
            print(f"✅ Report status retrieved")
            print(f"   Status: {result.get('status')}")
            print(f"   Created: {result.get('created_at')}")
            return result
        else:
            print(f"❌ Status retrieval failed: {response.status_code}")
            print(f"   Response: {response.text}")
            return None
            
    except Exception as e:
        print(f"❌ Status retrieval error: {str(e)}")
        return None

def test_report_preview(report_id):
    """Test report preview endpoint"""
    print(f"\n🔧 Testing Report Preview for ID: {report_id}...")
    
    try:
        response = requests.get(f"{API_BASE}/reports/{report_id}/preview")
        
        if response.status_code == 200:
            result = response.json()
            print(f"✅ Report preview retrieved")
            print(f"   Preview data keys: {list(result.keys())}")
            return result
        else:
            print(f"❌ Preview retrieval failed: {response.status_code}")
            print(f"   Response: {response.text}")
            return None
            
    except Exception as e:
        print(f"❌ Preview retrieval error: {str(e)}")
        return None

def test_storage_usage():
    """Test storage usage endpoint"""
    print("\n🔧 Testing Storage Usage...")
    
    try:
        response = requests.get(f"{API_BASE}/reports/lifecycle/storage")
        
        if response.status_code == 200:
            result = response.json()
            print(f"✅ Storage usage retrieved")
            print(f"   Total reports: {result.get('total_reports', 0)}")
            print(f"   Total storage: {result.get('total_storage_mb', 0)} MB")
            return result
        else:
            print(f"❌ Storage usage failed: {response.status_code}")
            print(f"   Response: {response.text}")
            return None
            
    except Exception as e:
        print(f"❌ Storage usage error: {str(e)}")
        return None

def create_test_docx():
    """Create a simple test DOCX file if none exists"""
    print("\n🔧 Creating Test DOCX File...")
    
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document for report generation.')
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('This is the first section with some content.')
        doc.add_heading('Section 2', level=1)
        doc.add_paragraph('This is the second section with more content.')
        
        test_file_path = "test_document.docx"
        doc.save(test_file_path)
        
        print(f"✅ Test DOCX file created: {test_file_path}")
        return test_file_path
        
    except ImportError:
        print("⚠️ python-docx not available, skipping test file creation")
        return None
    except Exception as e:
        print(f"❌ Error creating test file: {str(e)}")
        return None

def main():
    """Run all tests"""
    print("🚀 Testing DOCX Report Generation System")
    print("=" * 60)
    
    # Check if test file exists, create if needed
    test_file = "test_document.docx"
    if not os.path.exists(test_file):
        create_test_docx()
    
    # Test file upload
    upload_result = test_file_upload()
    
    if upload_result:
        # Test report generation
        report_result = test_report_generation(upload_result)
        
        if report_result:
            report_id = report_result.get('id')
            
            # Test report status
            test_report_status(report_id)
            
            # Test report preview
            test_report_preview(report_id)
    
    # Test storage usage
    test_storage_usage()
    
    print("\n" + "=" * 60)
    print("📊 Test Summary")
    print("=" * 60)
    
    if upload_result and report_result:
        print("✅ DOCX Report Generation System is working!")
        print("   - File upload: SUCCESS")
        print("   - Report generation: SUCCESS")
        print("   - You can now use the frontend to generate reports")
    else:
        print("❌ Some tests failed. Check the errors above.")
        print("   - Make sure your backend is running")
        print("   - Check that all required services are started")
        print("   - Verify database connections")

if __name__ == "__main__":
    main()
