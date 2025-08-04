from celery import shared_task
from .services.report_service import report_service
from .services.ai_service import ai_service
from .services.form_automation import form_automation_service
from .models import Report, Form, FormSubmission, db
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime, timedelta
import json
import os
from docx import Document
from docx.shared import Inches
import io
import base64

@shared_task
def generate_report_task(user_id, data):
    try:
        # Get AI suggestions for the report
        suggestions = ai_service.generate_report_suggestions(data)
        
        # Merge suggestions with user data
        enriched_data = {**data, 'ai_suggestions': suggestions}
        
        # Generate the report
        output_path = report_service.generate_report(
            template_id=data.get('template_id'),
            data=enriched_data
        )
        
        # Update report status in database
        report = Report.query.get(data.get('report_id'))
        if report:
            report.status = 'completed'
            report.output_url = output_path
            db.session.commit()
        
        return {
            'status': 'success',
            'output_path': output_path,
            'suggestions': suggestions
        }
        
    except Exception as e:
        # Update report status to failed
        report = Report.query.get(data.get('report_id'))
        if report:
            report.status = 'failed'
            db.session.commit()
            
        return {
            'status': 'error',
            'error': str(e)
        }

@shared_task
def generate_automated_report_task(form_id, report_type='summary', date_range='last_30_days', 
                                 trigger_type='manual', user_id=None, submission_id=None):
    """
    Generate automated reports from form submissions using AI analysis and charts
    """
    try:
        # Get form data
        form = Form.query.get(form_id)
        if not form:
            raise Exception(f"Form {form_id} not found")
        
        # Calculate date range
        end_date = datetime.utcnow()
        if date_range == 'last_7_days':
            start_date = end_date - timedelta(days=7)
        elif date_range == 'last_30_days':
            start_date = end_date - timedelta(days=30)
        elif date_range == 'last_90_days':
            start_date = end_date - timedelta(days=90)
        else:
            start_date = end_date - timedelta(days=30)
        
        # Get form submissions
        submissions_query = FormSubmission.query.filter_by(form_id=form_id)\
            .filter(FormSubmission.submitted_at >= start_date)\
            .filter(FormSubmission.submitted_at <= end_date)
        
        if submission_id:
            # If triggered by specific submission, include it
            submissions_query = submissions_query.filter(
                FormSubmission.id >= submission_id
            )
        
        submissions = submissions_query.all()
        
        if not submissions:
            return {
                'status': 'warning',
                'message': 'No submissions found for the specified date range'
            }
        
        # Convert to pandas DataFrame for analysis
        df = pd.DataFrame([
            {
                'id': sub.id,
                'submitted_at': sub.submitted_at,
                'submitter_email': sub.submitter_email,
                'source': sub.submission_source,
                **sub.data  # Flatten the form data
            }
            for sub in submissions
        ])
        
        # Generate analysis and charts
        analysis_result = analyze_form_data(df, form, report_type)
        
        # Generate Word document
        doc_path = generate_word_report(form, analysis_result, submissions)
        
        # Create report record
        report = Report(
            title=f"Automated Report - {form.title}",
            description=f"Auto-generated {report_type} report for {form.title}",
            user_id=user_id or form.creator_id,
            status='completed',
            data=analysis_result,
            output_url=doc_path,
            template_id='automated_report'
        )
        
        db.session.add(report)
        db.session.commit()
        
        return {
            'status': 'success',
            'report_id': report.id,
            'output_path': doc_path,
            'analysis': analysis_result,
            'submissions_count': len(submissions),
            'date_range': f"{start_date.date()} to {end_date.date()}"
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e)
        }

def analyze_form_data(df, form, report_type):
    """Analyze form data using pandas and AI"""
    analysis = {
        'summary': {},
        'trends': {},
        'charts': {},
        'insights': []
    }
    
    # Basic statistics
    analysis['summary'] = {
        'total_submissions': len(df),
        'date_range': f"{df['submitted_at'].min().date()} to {df['submitted_at'].max().date()}",
        'unique_submitters': df['submitter_email'].nunique(),
        'submission_sources': df['source'].value_counts().to_dict()
    }
    
    # Time-based trends
    df['date'] = pd.to_datetime(df['submitted_at']).dt.date
    daily_submissions = df.groupby('date').size()
    
    analysis['trends'] = {
        'daily_submissions': daily_submissions.to_dict(),
        'submission_rate': len(df) / max(1, (df['submitted_at'].max() - df['submitted_at'].min()).days)
    }
    
    # Generate charts
    charts = generate_charts(df, form)
    analysis['charts'] = charts
    
    # AI-powered insights
    insights = generate_ai_insights(df, form, report_type)
    analysis['insights'] = insights
    
    return analysis

def generate_charts(df, form):
    """Generate matplotlib charts for the report"""
    charts = {}
    
    # Set style
    plt.style.use('seaborn-v0_8')
    
    # 1. Daily submissions trend
    fig, ax = plt.subplots(figsize=(10, 6))
    daily_counts = df.groupby(df['submitted_at'].dt.date).size()
    ax.plot(daily_counts.index, daily_counts.values, marker='o', linewidth=2, markersize=6)
    ax.set_title('Daily Submissions Trend', fontsize=14, fontweight='bold')
    ax.set_xlabel('Date')
    ax.set_ylabel('Number of Submissions')
    ax.grid(True, alpha=0.3)
    plt.xticks(rotation=45)
    plt.tight_layout()
    
    # Save chart
    chart_path = f"static/charts/daily_trend_{form.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
    os.makedirs(os.path.dirname(chart_path), exist_ok=True)
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    charts['daily_trend'] = chart_path
    
    # 2. Submission source distribution
    if 'source' in df.columns and df['source'].nunique() > 1:
        fig, ax = plt.subplots(figsize=(8, 6))
        source_counts = df['source'].value_counts()
        ax.pie(source_counts.values, labels=source_counts.index, autopct='%1.1f%%', startangle=90)
        ax.set_title('Submissions by Source', fontsize=14, fontweight='bold')
        plt.tight_layout()
        
        chart_path = f"static/charts/source_distribution_{form.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        charts['source_distribution'] = chart_path
    
    # 3. Hourly submission pattern
    df['hour'] = pd.to_datetime(df['submitted_at']).dt.hour
    hourly_counts = df['hour'].value_counts().sort_index()
    
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.bar(hourly_counts.index, hourly_counts.values, alpha=0.7, color='skyblue')
    ax.set_title('Submissions by Hour of Day', fontsize=14, fontweight='bold')
    ax.set_xlabel('Hour of Day')
    ax.set_ylabel('Number of Submissions')
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    
    chart_path = f"static/charts/hourly_pattern_{form.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    charts['hourly_pattern'] = chart_path
    
    return charts

def generate_ai_insights(df, form, report_type):
    """Generate AI-powered insights from form data"""
    insights = []
    
    # Prepare data for AI analysis
    data_summary = {
        'total_submissions': len(df),
        'date_range': f"{df['submitted_at'].min().date()} to {df['submitted_at'].max().date()}",
        'submission_rate': len(df) / max(1, (df['submitted_at'].max() - df['submitted_at'].min()).days),
        'unique_submitters': df['submitter_email'].nunique(),
        'form_title': form.title,
        'form_description': form.description
    }
    
    # Add form-specific data analysis
    if form.schema and 'fields' in form.schema:
        for field in form.schema['fields']:
            field_name = field.get('label', field.get('id', ''))
            if field_name in df.columns:
                field_data = df[field_name].dropna()
                if len(field_data) > 0:
                    if field_data.dtype == 'object':
                        # Categorical data
                        value_counts = field_data.value_counts()
                        data_summary[f'{field_name}_top_values'] = value_counts.head(5).to_dict()
                    else:
                        # Numerical data
                        data_summary[f'{field_name}_stats'] = {
                            'mean': float(field_data.mean()),
                            'median': float(field_data.median()),
                            'min': float(field_data.min()),
                            'max': float(field_data.max())
                        }
    
    # Generate AI insights
    try:
        ai_response = ai_service.analyze_form_data(data_summary, report_type)
        insights = ai_response.get('insights', [])
    except Exception as e:
        # Fallback insights if AI service fails
        insights = [
            f"Generated {len(df)} submissions over {data_summary['date_range']}",
            f"Average submission rate: {data_summary['submission_rate']:.1f} per day",
            f"Unique submitters: {data_summary['unique_submitters']}"
        ]
    
    return insights

def generate_word_report(form, analysis, submissions):
    """Generate Word document with analysis and charts"""
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Automated Report: {form.title}', 0)
    title.alignment = 1  # Center alignment
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run(f"This report analyzes {analysis['summary']['total_submissions']} submissions ")
    summary.add_run(f"from {analysis['summary']['date_range']} for the form '{form.title}'.")
    
    # Key Metrics
    doc.add_heading('Key Metrics', level=1)
    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.style = 'Table Grid'
    
    # Add header row
    header_cells = metrics_table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    
    # Add data rows
    metrics_data = [
        ('Total Submissions', str(analysis['summary']['total_submissions'])),
        ('Unique Submitters', str(analysis['summary']['unique_submitters'])),
        ('Submission Rate', f"{analysis['trends']['submission_rate']:.1f} per day"),
        ('Date Range', analysis['summary']['date_range'])
    ]
    
    for metric, value in metrics_data:
        row_cells = metrics_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    # Charts
    if analysis['charts']:
        doc.add_heading('Visual Analysis', level=1)
        
        for chart_name, chart_path in analysis['charts'].items():
            if os.path.exists(chart_path):
                doc.add_paragraph(f'{chart_name.replace("_", " ").title()}:')
                doc.add_picture(chart_path, width=Inches(6))
                doc.add_paragraph()  # Add spacing
    
    # AI Insights
    if analysis['insights']:
        doc.add_heading('AI-Generated Insights', level=1)
        for insight in analysis['insights']:
            doc.add_paragraph(insight, style='List Bullet')
    
    # Recent Submissions Sample
    doc.add_heading('Recent Submissions Sample', level=1)
    recent_submissions = submissions[:10]  # Show last 10 submissions
    
    if recent_submissions:
        submissions_table = doc.add_table(rows=1, cols=3)
        submissions_table.style = 'Table Grid'
        
        # Header
        header_cells = submissions_table.rows[0].cells
        header_cells[0].text = 'Date'
        header_cells[1].text = 'Submitter'
        header_cells[2].text = 'Source'
        
        # Data
        for submission in recent_submissions:
            row_cells = submissions_table.add_row().cells
            row_cells[0].text = submission.submitted_at.strftime('%Y-%m-%d %H:%M')
            row_cells[1].text = submission.submitter_email or 'Anonymous'
            row_cells[2].text = submission.submission_source
    
    # Save document
    doc_path = f"static/reports/automated_report_{form.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    doc.save(doc_path)
    
    return doc_path

@shared_task
def schedule_automated_reports():
    """Scheduled task to generate automated reports for forms with auto-report enabled"""
    try:
        # Find forms with auto-report enabled
        forms = Form.query.filter(
            Form.form_settings.contains({'auto_generate_reports': True})
        ).all()
        
        results = []
        for form in forms:
            # Check if it's time to generate a report based on schedule
            schedule = form.form_settings.get('report_schedule', 'weekly')
            last_report = Report.query.filter_by(
                template_id='automated_report',
                user_id=form.creator_id
            ).order_by(Report.created_at.desc()).first()
            
            should_generate = False
            if not last_report:
                should_generate = True
            else:
                days_since_last = (datetime.utcnow() - last_report.created_at).days
                if schedule == 'daily' and days_since_last >= 1:
                    should_generate = True
                elif schedule == 'weekly' and days_since_last >= 7:
                    should_generate = True
                elif schedule == 'monthly' and days_since_last >= 30:
                    should_generate = True
            
            if should_generate:
                task = generate_automated_report_task.delay(
                    form_id=form.id,
                    report_type='scheduled',
                    date_range='last_30_days',
                    trigger_type='scheduled',
                    user_id=form.creator_id
                )
                results.append({
                    'form_id': form.id,
                    'form_title': form.title,
                    'task_id': task.id
                })
        
        return {
            'status': 'success',
            'reports_generated': len(results),
            'results': results
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e)
        }
