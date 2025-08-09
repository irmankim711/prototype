"""
Enhanced Report Template Generator with Malay Language Support
Aligned with Temp1.docx formatting requirements including fonts, margins, and structure
"""

import os
import tempfile
from datetime import datetime
from typing import Dict, List, Any, Optional
import base64
from io import BytesIO
import logging

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.enum.section import WD_ORIENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)

class MalayTemplateGenerator:
    """
    Enhanced report generator specifically designed for Malaysian standards
    Supports Malay language formatting and Temp1.docx compliance
    """
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        self.charts_dir = os.path.join(self.temp_dir, 'charts')
        os.makedirs(self.charts_dir, exist_ok=True)
        
        # Malaysian standard fonts
        self.malay_fonts = {
            'primary': 'Times New Roman',  # Standard for Malaysian official documents
            'secondary': 'Arial',
            'headers': 'Arial Black',
            'fallback': 'Calibri'
        }
        
        # Color scheme for Malaysian government/corporate documents
        self.color_scheme = {
            'primary': '1f4e79',     # Dark blue
            'secondary': '2f75b5',   # Medium blue  
            'accent': 'c5504b',      # Red accent
            'text': '000000',        # Black
            'light_blue': 'dae3f3'  # Light blue for tables
        }
    
    def create_temp1_compliant_report(self, data: Dict[str, Any], output_path: Optional[str] = None) -> str:
        """
        Create report following exact Temp1.docx structure and formatting
        
        Args:
            data: Report data including submissions, analytics, metadata
            output_path: Optional output file path
            
        Returns:
            Path to generated report
        """
        if not HAS_DOCX:
            raise ImportError("python-docx required for document generation")
        
        # Load base template or create new document
        doc = Document()
        
        # Setup Malaysian document standards
        self._setup_malaysian_document_format(doc)
        self._setup_malay_styles(doc)
        
        # Generate report sections in Temp1.docx order
        self._add_cover_page(doc, data)
        self._add_page_break(doc)
        
        self._add_table_of_contents_malay(doc)
        self._add_page_break(doc)
        
        self._add_program_information_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_program_objectives_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_program_schedule_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_program_evaluation_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_evaluation_summary_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_pre_post_analysis_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_attendance_analysis_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_individual_evaluations_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_program_gallery_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_conclusions_recommendations_malay(doc, data)
        self._add_page_break(doc)
        
        self._add_signatures_malay(doc, data)
        
        # Save document
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = os.path.join(self.temp_dir, f'laporan_program_{timestamp}.docx')
        
        doc.save(output_path)
        logger.info(f"Temp1.docx compliant report generated: {output_path}")
        return output_path
    
    def _setup_malaysian_document_format(self, doc: Document):
        """Setup document format according to Malaysian standards"""
        # Setup page layout - A4 with standard margins
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A4 height
        section.page_width = Cm(21.0)   # A4 width
        
        # Malaysian standard margins (usually 2.5cm all around)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        
        # Header and footer margins
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
    
    def _setup_malay_styles(self, doc: Document):
        """Create styles optimized for Malay language documents"""
        styles = doc.styles
        
        # Title style (for cover page)
        try:
            title_style = styles.add_style('Malay Title', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            title_style = styles['Malay Title']
        
        title_font = title_style.font
        title_font.name = self.malay_fonts['headers']
        title_font.size = Pt(18)
        title_font.bold = True
        title_font.color.rgb = self._hex_to_rgb(self.color_scheme['primary'])
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1 style (for main sections)
        try:
            h1_style = styles.add_style('Malay Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            h1_style = styles['Malay Heading 1']
        
        h1_font = h1_style.font
        h1_font.name = self.malay_fonts['primary']
        h1_font.size = Pt(14)
        h1_font.bold = True
        h1_font.color.rgb = self._hex_to_rgb(self.color_scheme['primary'])
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
        
        # Heading 2 style (for subsections)
        try:
            h2_style = styles.add_style('Malay Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            h2_style = styles['Malay Heading 2']
        
        h2_font = h2_style.font
        h2_font.name = self.malay_fonts['primary']
        h2_font.size = Pt(12)
        h2_font.bold = True
        h2_font.color.rgb = self._hex_to_rgb(self.color_scheme['secondary'])
        h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2_style.paragraph_format.space_before = Pt(6)
        h2_style.paragraph_format.space_after = Pt(3)
        
        # Body text style
        try:
            body_style = styles.add_style('Malay Body', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            body_style = styles['Malay Body']
        
        body_font = body_style.font
        body_font.name = self.malay_fonts['primary']
        body_font.size = Pt(11)
        body_font.color.rgb = self._hex_to_rgb(self.color_scheme['text'])
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.line_spacing = 1.15
        body_style.paragraph_format.space_after = Pt(6)
        
        # Table header style
        try:
            table_header_style = styles.add_style('Malay Table Header', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            table_header_style = styles['Malay Table Header']
        
        table_header_font = table_header_style.font
        table_header_font.name = self.malay_fonts['primary']
        table_header_font.size = Pt(10)
        table_header_font.bold = True
        table_header_font.color.rgb = self._hex_to_rgb('ffffff')
        table_header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _hex_to_rgb(self, hex_color: str):
        """Convert hex color to RGBColor object"""
        hex_color = hex_color.lstrip('#')
        r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        return RGBColor(r, g, b)
    
    def _add_page_break(self, doc: Document):
        """Add page break"""
        doc.add_page_break()
    
    def _add_cover_page(self, doc: Document, data: Dict[str, Any]):
        """Add Malay-style cover page"""
        # Organization logo placeholder
        cover_p = doc.add_paragraph()
        cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Main title
        title = data.get('title', 'LAPORAN PROGRAM')
        title_p = doc.add_paragraph(title)
        title_p.style = 'Malay Title'
        
        # Subtitle
        subtitle = data.get('subtitle', 'Program Pembangunan Kemahiran')
        subtitle_p = doc.add_paragraph(subtitle)
        subtitle_p.style = 'Malay Heading 1'
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Program details
        program_name = data.get('program_name', 'Nama Program')
        program_p = doc.add_paragraph(f'Program: {program_name}')
        program_p.style = 'Malay Body'
        program_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        report_date = data.get('date', datetime.now().strftime('%d %B %Y'))
        date_p = doc.add_paragraph(f'Tarikh: {report_date}')
        date_p.style = 'Malay Body'
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Organization details
        org_p = doc.add_paragraph('\n\n')
        org_details = data.get('organization', {})
        org_name = org_details.get('name', 'Nama Organisasi')
        org_address = org_details.get('address', 'Alamat Organisasi')
        
        org_info = doc.add_paragraph(f'{org_name}\n{org_address}')
        org_info.style = 'Malay Body'
        org_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_table_of_contents_malay(self, doc: Document):
        """Add table of contents in Malay"""
        toc_title = doc.add_paragraph('KANDUNGAN')
        toc_title.style = 'Malay Heading 1'
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # TOC entries
        toc_items = [
            ('1.0', 'Maklumat Program', '3'),
            ('2.0', 'Objektif Program', '4'),
            ('3.0', 'Jadual Tentatif Program', '5'),
            ('4.0', 'Penilaian Program', '6'),
            ('5.0', 'Rumusan Penilaian', '7'),
            ('6.0', 'Analisis Pra dan Pasca Program', '8'),
            ('7.0', 'Laporan Kehadiran', '9'),
            ('8.0', 'Penilaian Individu', '10'),
            ('9.0', 'Galeri Program', '11'),
            ('10.0', 'Kesimpulan dan Cadangan', '12'),
            ('11.0', 'Tandatangan', '13')
        ]
        
        for section_num, section_title, page_num in toc_items:
            toc_entry = doc.add_paragraph()
            toc_entry.style = 'Malay Body'
            
            # Add section number and title
            run1 = toc_entry.add_run(f'{section_num}\t{section_title}')
            
            # Add dots
            run2 = toc_entry.add_run('\t' + '.' * 50)
            
            # Add page number
            run3 = toc_entry.add_run(f'\t{page_num}')
    
    def _add_program_information_malay(self, doc: Document, data: Dict[str, Any]):
        """Add program information section in Malay"""
        section_title = doc.add_paragraph('1.0 MAKLUMAT PROGRAM')
        section_title.style = 'Malay Heading 1'
        
        # Create information table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Program information data
        program_info = [
            ('Nama Program', data.get('program_name', 'Program Pembangunan Kemahiran')),
            ('Anjuran', data.get('organizer', 'Bahagian Pembangunan Sumber Manusia')),
            ('Tarikh', data.get('date_range', '1 Januari 2025 - 31 Januari 2025')),
            ('Masa', data.get('time', '9:00 AM - 5:00 PM')),
            ('Tempat', data.get('venue', 'Dewan Utama, Tingkat 3')),
            ('Bilangan Peserta', str(data.get('participant_count', 25))),
            ('Fasilitator', data.get('facilitator', 'En. Ahmad bin Ali')),
            ('Objektif Utama', data.get('main_objective', 'Meningkatkan kemahiran teknikal peserta'))
        ]
        
        for label, value in program_info:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value
            
            # Style the cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Malay Body'
                    if cell == row_cells[0]:  # Label column
                        for run in paragraph.runs:
                            run.font.bold = True
    
    def _add_program_objectives_malay(self, doc: Document, data: Dict[str, Any]):
        """Add program objectives section"""
        section_title = doc.add_paragraph('2.0 OBJEKTIF PROGRAM')
        section_title.style = 'Malay Heading 1'
        
        objectives = data.get('objectives', [
            'Meningkatkan kemahiran teknikal peserta dalam bidang yang berkaitan',
            'Memupuk semangat kerjasama dan kepimpinan dalam kalangan peserta',
            'Menyediakan platform untuk berkongsi pengalaman dan best practices',
            'Membangunkan rangkaian profesional yang kukuh'
        ])
        
        for i, objective in enumerate(objectives, 1):
            obj_p = doc.add_paragraph(f'{i}. {objective}')
            obj_p.style = 'Malay Body'
    
    def _add_program_schedule_malay(self, doc: Document, data: Dict[str, Any]):
        """Add program schedule section"""
        section_title = doc.add_paragraph('3.0 JADUAL TENTATIF PROGRAM')
        section_title.style = 'Malay Heading 1'
        
        # Create schedule table
        schedule_table = doc.add_table(rows=1, cols=3)
        schedule_table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = schedule_table.rows[0].cells
        headers = ['Masa', 'Aktiviti', 'Fasilitator']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                paragraph.style = 'Malay Table Header'
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Schedule data
        schedule_items = data.get('schedule', [
            ('9:00 - 9:30', 'Pendaftaran dan Sarapan Pagi', 'Sekretariat'),
            ('9:30 - 10:00', 'Taklimat Program dan Perkenalan', 'Pengerusi Program'),
            ('10:00 - 12:00', 'Sesi 1: Pengenalan Konsep Asas', 'Fasilitator Utama'),
            ('12:00 - 13:00', 'Rehat Tengah Hari', '-'),
            ('13:00 - 15:00', 'Sesi 2: Aktiviti Berkumpulan', 'Fasilitator Utama'),
            ('15:00 - 15:30', 'Rehat Petang', '-'),
            ('15:30 - 17:00', 'Sesi 3: Pembentangan dan Perbincangan', 'Semua Fasilitator'),
            ('17:00 - 17:30', 'Penutup dan Penyampaian Sijil', 'Pengerusi Program')
        ])
        
        for time_slot, activity, facilitator in schedule_items:
            row_cells = schedule_table.add_row().cells
            row_cells[0].text = time_slot
            row_cells[1].text = activity
            row_cells[2].text = facilitator
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Malay Body'
    
    def _add_program_evaluation_malay(self, doc: Document, data: Dict[str, Any]):
        """Add program evaluation methodology section"""
        section_title = doc.add_paragraph('4.0 PENILAIAN PROGRAM')
        section_title.style = 'Malay Heading 1'
        
        eval_intro = doc.add_paragraph(
            'Penilaian program dilakukan menggunakan pendekatan komprehensif yang merangkumi '
            'beberapa kaedah untuk memastikan keberkesanan program dapat diukur dengan tepat.'
        )
        eval_intro.style = 'Malay Body'
        
        # Evaluation methods
        methods_title = doc.add_paragraph('4.1 Kaedah Penilaian')
        methods_title.style = 'Malay Heading 2'
        
        methods = [
            'Borang Penilaian Pra-Program (Pre-Assessment)',
            'Borang Penilaian Pasca-Program (Post-Assessment)',
            'Borang Maklum Balas Peserta (Participant Feedback)',
            'Pemerhatian dan Rekod Kehadiran',
            'Penilaian Aktiviti dan Projek Berkumpulan'
        ]
        
        for method in methods:
            method_p = doc.add_paragraph(f'• {method}')
            method_p.style = 'Malay Body'
        
        # Evaluation criteria
        criteria_title = doc.add_paragraph('4.2 Kriteria Penilaian')
        criteria_title.style = 'Malay Heading 2'
        
        criteria_table = doc.add_table(rows=1, cols=2)
        criteria_table.style = 'Light Grid Accent 1'
        
        # Criteria table headers
        criteria_headers = criteria_table.rows[0].cells
        criteria_headers[0].text = 'Aspek Penilaian'
        criteria_headers[1].text = 'Skala Penilaian'
        
        for cell in criteria_headers:
            for paragraph in cell.paragraphs:
                paragraph.style = 'Malay Table Header'
        
        # Criteria data
        criteria_data = [
            ('Kandungan Program', '1-5 (Sangat Tidak Bersetuju - Sangat Bersetuju)'),
            ('Kaedah Penyampaian', '1-5 (Sangat Tidak Bersetuju - Sangat Bersetuju)'),
            ('Kemudahan dan Peralatan', '1-5 (Sangat Tidak Bersetuju - Sangat Bersetuju)'),
            ('Keberkesanan Fasilitator', '1-5 (Sangat Tidak Bersetuju - Sangat Bersetuju)'),
            ('Peningkatan Pengetahuan', '1-5 (Sangat Tidak Bersetuju - Sangat Bersetuju)')
        ]
        
        for aspect, scale in criteria_data:
            row_cells = criteria_table.add_row().cells
            row_cells[0].text = aspect
            row_cells[1].text = scale
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Malay Body'
    
    def _add_evaluation_summary_malay(self, doc: Document, data: Dict[str, Any]):
        """Add evaluation summary with analytics"""
        section_title = doc.add_paragraph('5.0 RUMUSAN PENILAIAN')
        section_title.style = 'Malay Heading 1'
        
        # Overall satisfaction
        overall_title = doc.add_paragraph('5.1 Kepuasan Keseluruhan')
        overall_title.style = 'Malay Heading 2'
        
        # Get analytics data
        analytics = data.get('analytics', {})
        avg_satisfaction = analytics.get('average_satisfaction', 4.2)
        response_rate = analytics.get('response_rate', 85)
        
        satisfaction_text = doc.add_paragraph(
            f'Berdasarkan maklum balas yang diterima daripada {response_rate}% peserta, '
            f'skor purata kepuasan keseluruhan ialah {avg_satisfaction:.1f}/5.0. '
            'Ini menunjukkan tahap kepuasan yang tinggi terhadap program yang dijalankan.'
        )
        satisfaction_text.style = 'Malay Body'
        
        # Key findings
        findings_title = doc.add_paragraph('5.2 Dapatan Utama')
        findings_title.style = 'Malay Heading 2'
        
        findings = data.get('key_findings', [
            'Majoriti peserta (95%) bersetuju bahawa objektif program telah tercapai',
            'Kandungan program dinilai relevan dan praktikal (skor purata: 4.3/5)',
            'Kaedah penyampaian fasilitator mendapat penilaian yang positif (skor purata: 4.1/5)',
            'Kemudahan dan peralatan yang disediakan adalah mencukupi (skor purata: 4.0/5)'
        ])
        
        for finding in findings:
            finding_p = doc.add_paragraph(f'• {finding}')
            finding_p.style = 'Malay Body'
    
    def _add_pre_post_analysis_malay(self, doc: Document, data: Dict[str, Any]):
        """Add pre-post program analysis"""
        section_title = doc.add_paragraph('6.0 ANALISIS PRA DAN PASCA PROGRAM')
        section_title.style = 'Malay Heading 1'
        
        analysis_intro = doc.add_paragraph(
            'Analisis perbandingan antara tahap pengetahuan dan kemahiran peserta '
            'sebelum dan selepas program untuk mengukur keberkesanan program.'
        )
        analysis_intro.style = 'Malay Body'
        
        # Knowledge improvement table
        improvement_table = doc.add_table(rows=1, cols=4)
        improvement_table.style = 'Light Grid Accent 1'
        
        headers = ['Aspek', 'Pra-Program', 'Pasca-Program', 'Peningkatan (%)']
        header_cells = improvement_table.rows[0].cells
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                paragraph.style = 'Malay Table Header'
        
        # Sample improvement data
        improvement_data = data.get('improvement_data', [
            ('Pengetahuan Asas', '2.8', '4.2', '50%'),
            ('Kemahiran Praktikal', '2.5', '4.0', '60%'),
            ('Keyakinan Diri', '3.0', '4.3', '43%'),
            ('Kebolehan Kerja Berpasukan', '3.2', '4.5', '41%')
        ])
        
        for aspect, pre_score, post_score, improvement in improvement_data:
            row_cells = improvement_table.add_row().cells
            row_cells[0].text = aspect
            row_cells[1].text = pre_score
            row_cells[2].text = post_score
            row_cells[3].text = improvement
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Malay Body'
    
    def _add_attendance_analysis_malay(self, doc: Document, data: Dict[str, Any]):
        """Add attendance analysis section"""
        section_title = doc.add_paragraph('7.0 LAPORAN KEHADIRAN')
        section_title.style = 'Malay Heading 1'
        
        # Attendance summary
        total_participants = data.get('total_participants', 25)
        attendance_rate = data.get('attendance_rate', 96)
        
        summary_text = doc.add_paragraph(
            f'Daripada {total_participants} peserta yang berdaftar, {attendance_rate}% '
            'hadir sepanjang program. Ini menunjukkan komitmen yang tinggi daripada peserta.'
        )
        summary_text.style = 'Malay Body'
        
        # Daily attendance breakdown
        daily_title = doc.add_paragraph('7.1 Pecahan Kehadiran Harian')
        daily_title.style = 'Malay Heading 2'
        
        attendance_table = doc.add_table(rows=1, cols=3)
        attendance_table.style = 'Light Grid Accent 1'
        
        att_headers = ['Hari', 'Bilangan Hadir', 'Peratus Kehadiran']
        att_header_cells = attendance_table.rows[0].cells
        
        for i, header in enumerate(att_headers):
            att_header_cells[i].text = header
            for paragraph in att_header_cells[i].paragraphs:
                paragraph.style = 'Malay Table Header'
        
        # Sample attendance data
        attendance_data = data.get('daily_attendance', [
            ('Hari 1', '24', '96%'),
            ('Hari 2', '25', '100%'),
            ('Hari 3', '23', '92%')
        ])
        
        for day, count, percentage in attendance_data:
            row_cells = attendance_table.add_row().cells
            row_cells[0].text = day
            row_cells[1].text = count
            row_cells[2].text = percentage
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Malay Body'
    
    def _add_individual_evaluations_malay(self, doc: Document, data: Dict[str, Any]):
        """Add individual participant evaluations"""
        section_title = doc.add_paragraph('8.0 PENILAIAN INDIVIDU')
        section_title.style = 'Malay Heading 1'
        
        eval_intro = doc.add_paragraph(
            'Ringkasan penilaian prestasi individu peserta berdasarkan penyertaan, '
            'penglibatan, dan pencapaian objektif pembelajaran.'
        )
        eval_intro.style = 'Malay Body'
        
        # Performance categories
        categories_title = doc.add_paragraph('8.1 Kategori Prestasi')
        categories_title.style = 'Malay Heading 2'
        
        performance_data = data.get('performance_categories', {
            'Cemerlang (4.5-5.0)': '40%',
            'Baik (3.5-4.4)': '45%',
            'Memuaskan (2.5-3.4)': '15%',
            'Perlu Diperbaiki (<2.5)': '0%'
        })
        
        for category, percentage in performance_data.items():
            cat_p = doc.add_paragraph(f'• {category}: {percentage} peserta')
            cat_p.style = 'Malay Body'
    
    def _add_program_gallery_malay(self, doc: Document, data: Dict[str, Any]):
        """Add program photo gallery section"""
        section_title = doc.add_paragraph('9.0 GALERI PROGRAM')
        section_title.style = 'Malay Heading 1'
        
        gallery_intro = doc.add_paragraph(
            'Koleksi foto-foto aktiviti yang telah dijalankan sepanjang program.'
        )
        gallery_intro.style = 'Malay Body'
        
        # Placeholder for images
        image_sections = [
            'Majlis Perasmian dan Taklimat',
            'Sesi Pembelajaran dan Aktiviti',
            'Kerja Berkumpulan dan Pembentangan',
            'Majlis Penutup dan Penyampaian Sijil'
        ]
        
        for section in image_sections:
            img_title = doc.add_paragraph(section)
            img_title.style = 'Malay Heading 2'
            
            # Placeholder text for images
            img_placeholder = doc.add_paragraph('[Ruang untuk gambar akan dimasukkan di sini]')
            img_placeholder.style = 'Malay Body'
            img_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_conclusions_recommendations_malay(self, doc: Document, data: Dict[str, Any]):
        """Add conclusions and recommendations"""
        section_title = doc.add_paragraph('10.0 KESIMPULAN DAN CADANGAN')
        section_title.style = 'Malay Heading 1'
        
        # Conclusions
        conclusion_title = doc.add_paragraph('10.1 Kesimpulan')
        conclusion_title.style = 'Malay Heading 2'
        
        conclusions = data.get('conclusions', [
            'Program telah berjaya mencapai objektif yang ditetapkan dengan jayanya',
            'Peserta menunjukkan peningkatan yang ketara dalam pengetahuan dan kemahiran',
            'Tahap kepuasan peserta adalah tinggi dan melebihi jangkaan',
            'Metodologi yang digunakan terbukti berkesan untuk jenis program ini'
        ])
        
        for conclusion in conclusions:
            conc_p = doc.add_paragraph(f'• {conclusion}')
            conc_p.style = 'Malay Body'
        
        # Recommendations
        recommendation_title = doc.add_paragraph('10.2 Cadangan Penambahbaikan')
        recommendation_title.style = 'Malay Heading 2'
        
        recommendations = data.get('recommendations', [
            'Memanjangkan tempoh program untuk lebih banyak hands-on activities',
            'Menambah bilangan fasilitator untuk nisbah yang lebih baik',
            'Menyediakan bahan rujukan tambahan untuk pembelajaran berterusan',
            'Mengadakan program susulan untuk pemantauan dan sokongan'
        ])
        
        for recommendation in recommendations:
            rec_p = doc.add_paragraph(f'• {recommendation}')
            rec_p.style = 'Malay Body'
    
    def _add_signatures_malay(self, doc: Document, data: Dict[str, Any]):
        """Add signature section"""
        section_title = doc.add_paragraph('11.0 TANDATANGAN')
        section_title.style = 'Malay Heading 1'
        
        # Signature table
        sig_table = doc.add_table(rows=4, cols=2)
        
        # Prepared by
        sig_table.cell(0, 0).text = 'Disediakan oleh:'
        sig_table.cell(1, 0).text = '\n\n\n________________________'
        sig_table.cell(2, 0).text = data.get('prepared_by', 'Nama Penyedia Laporan')
        sig_table.cell(3, 0).text = data.get('prepared_by_title', 'Jawatan')
        
        # Approved by
        sig_table.cell(0, 1).text = 'Diluluskan oleh:'
        sig_table.cell(1, 1).text = '\n\n\n________________________'
        sig_table.cell(2, 1).text = data.get('approved_by', 'Nama Pelulus')
        sig_table.cell(3, 1).text = data.get('approved_by_title', 'Jawatan')
        
        # Style signature table
        for row in sig_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Malay Body'
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_p = doc.add_paragraph(f'\nTarikh: {datetime.now().strftime("%d %B %Y")}')
        date_p.style = 'Malay Body'
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
