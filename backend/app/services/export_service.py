"""
Unified export service for generating reports in PDF, DOCX, and HTML.
Uses ReportLab for PDF and docxtpl for DOCX template rendering.
"""
from __future__ import annotations

import os
import logging
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, List, Optional

from flask import current_app

logger = logging.getLogger(__name__)


try:
    from docxtpl import DocxTemplate  # docxtemplater
    from docx import Document  # used to create a default template if missing
    HAS_DOCXTPL = True
except Exception:
    HAS_DOCXTPL = False

try:
    # We will build a minimal PDF using reportlab directly if available;
    # alternatively, we can delegate to existing report_generator for richer PDFs.
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    HAS_REPORTLAB = True
except Exception:
    HAS_REPORTLAB = False

try:
    # Optional: leverage existing generator for richer formatting when possible
    from .report_generator import create_pdf_report as legacy_create_pdf
except Exception:
    legacy_create_pdf = None


ALLOWED_FORMATS = {"pdf", "docx", "html"}


@dataclass
class ExportResult:
    status: str
    urls: Dict[str, Optional[str]]
    filenames: Dict[str, Optional[str]]


class ExportService:
    def __init__(self):
        # No cached state to avoid cross-request/test contamination
        self.upload_folder = None

    def _ensure_dirs(self):
        # Always resolve from current app config to honor per-test/per-request upload folders
        base = current_app.config.get("UPLOAD_FOLDER", "uploads")
        upload_folder = os.path.abspath(os.path.join(current_app.root_path, "..", base))
        reports_dir = os.path.join(upload_folder, "reports")
        os.makedirs(reports_dir, exist_ok=True)
        return reports_dir

    def _templates_dir(self) -> str:
        return os.path.abspath(os.path.join(current_app.root_path, "..", "templates"))

    def _timestamp_name(self, template_id: str, ext: str) -> str:
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe_tpl = "".join(c for c in (template_id or "template") if c.isalnum() or c in ("-", "_"))
        return f"report_{safe_tpl}_{ts}.{ext}"

    def _build_url(self, filename: str) -> str:
        # Served by /api/reports/download/<filename>
        return f"/api/reports/download/{filename}"

    # ---------------------- Public API ----------------------
    def export(self, template_id: str, data_source: Dict, formats: List[str]) -> ExportResult:
        if not template_id or not isinstance(template_id, str):
            raise ValueError("template_id is required and must be a string")
        if not isinstance(data_source, dict):
            raise ValueError("data_source must be an object")
        if not formats or any(fmt.lower() not in ALLOWED_FORMATS for fmt in formats):
            raise ValueError("formats must be a non-empty list with any of: pdf, docx, html")

        reports_dir = self._ensure_dirs()
        urls: Dict[str, Optional[str]] = {k: None for k in ALLOWED_FORMATS}
        filenames: Dict[str, Optional[str]] = {k: None for k in ALLOWED_FORMATS}

        # Normalize a basic context from data_source
        context = self._build_context(template_id, data_source)

        for fmt in sorted(set(f.lower() for f in formats)):
            try:
                if fmt == "pdf":
                    filenames["pdf"] = self._export_pdf(template_id, context, reports_dir)
                    urls["pdf"] = self._build_url(filenames["pdf"])
                elif fmt == "docx":
                    filenames["docx"] = self._export_docx(template_id, context, reports_dir)
                    urls["docx"] = self._build_url(filenames["docx"])
                elif fmt == "html":
                    filenames["html"] = self._export_html(template_id, context, reports_dir)
                    urls["html"] = self._build_url(filenames["html"])
            except Exception as e:
                logger.error(f"Failed to export {fmt}: {e}")
                raise

        return ExportResult(status="success", urls=urls, filenames=filenames)

    # ---------------------- Helpers ----------------------
    def _build_context(self, template_id: str, data_source: Dict) -> Dict:
        # Basic pass-through; ensure a few known keys for templates
        ctx = dict(data_source or {})
        ctx.setdefault("title", ctx.get("title") or f"Report for {template_id}")
        ctx.setdefault("generated_at", datetime.utcnow().isoformat())
        return ctx

    def _export_pdf(self, template_id: str, context: Dict, reports_dir: str) -> str:
        if HAS_REPORTLAB:
            # If the legacy PDF creator exists, try to use it for richer output
            if legacy_create_pdf is not None:
                filename = self._timestamp_name(template_id, "pdf")
                output_path = os.path.join(reports_dir, filename)
                # The legacy function expects a template id string and data; it will write to path
                try:
                    legacy_create_pdf(template_id, context, output_path)
                    return filename
                except Exception:
                    # Fallback to minimal PDF if legacy fails
                    logger.warning("Falling back to minimal ReportLab PDF")

            # Minimal ReportLab PDF
            filename = self._timestamp_name(template_id, "pdf")
            output_path = os.path.join(reports_dir, filename)
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            story.append(Paragraph(str(context.get("title", "Report")), styles["Title"]))
            story.append(Spacer(1, 0.2 * inch))

            # Dump key-values into a simple table for visibility
            rows = [["Field", "Value"]]
            for k, v in list(context.items())[:20]:
                rows.append([str(k), str(v)])
            table = Table(rows, hAlign='LEFT')
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ]))
            story.append(table)

            doc.build(story)
            return filename

        raise RuntimeError("ReportLab is not available; cannot generate PDF")

    def _export_docx(self, template_id: str, context: Dict, reports_dir: str) -> str:
        if not HAS_DOCXTPL:
            raise RuntimeError("docxtpl (docxtemplater) is not available; cannot generate DOCX")

        templates_dir = self._templates_dir()
        # Resolve template path: prefer a matching file; fallback to TestTemplate.docx; else create a minimal template
        candidates = [
            os.path.join(templates_dir, f"{template_id}.docx"),
            os.path.join(templates_dir, "TestTemplate.docx"),
        ]
        template_path = next((p for p in candidates if os.path.exists(p)), None)
        temp_created = False
        if template_path is None:
            # Create a minimal template on the fly containing basic placeholders
            temp_created = True
            template_path = os.path.join(templates_dir, f"_auto_template_{template_id}.docx")
            try:
                os.makedirs(templates_dir, exist_ok=True)
                doc = Document()
                doc.add_heading("{{ title }}", 0)
                doc.add_paragraph("Generated at: {{ generated_at }}")
                doc.add_paragraph("Summary: {{ summary or 'N/A' }}")
                doc.save(template_path)
            except Exception as e:
                logger.error(f"Failed to create default DOCX template: {e}")
                raise

        try:
            tpl = DocxTemplate(template_path)
            try:
                tpl.render(context)
                filename = self._timestamp_name(template_id, "docx")
                output_path = os.path.join(reports_dir, filename)
                tpl.save(output_path)
                return filename
            except Exception as e:
                # Fallback: create a minimal DOCX without templating
                logger.warning(f"docxtpl render failed, falling back to plain DOCX: {e}")
                filename = self._timestamp_name(template_id, "docx")
                output_path = os.path.join(reports_dir, filename)
                from docx import Document as PlainDocument
                doc = PlainDocument()
                doc.add_heading(str(context.get('title', 'Report')), 0)
                doc.add_paragraph(f"Generated at: {context.get('generated_at')}")
                # Dump a few fields
                for k, v in list(context.items())[:20]:
                    if k not in ('title', 'generated_at'):
                        p = doc.add_paragraph()
                        p.add_run(f"{k}: ").bold = True
                        p.add_run(str(v))
                doc.save(output_path)
                return filename
        finally:
            # Keep auto template for reuse; do not delete to avoid races
            pass

    def _export_html(self, template_id: str, context: Dict, reports_dir: str) -> str:
        filename = self._timestamp_name(template_id, "html")
        output_path = os.path.join(reports_dir, filename)
        html = f"""
<!doctype html>
<html>
  <head>
    <meta charset=\"utf-8\" />
    <title>{context.get('title','Report')}</title>
    <style>
      body {{ font-family: Arial, sans-serif; margin: 24px; }}
      h1 {{ color: #2E4057; }}
      table {{ border-collapse: collapse; }}
      td, th {{ border: 1px solid #ccc; padding: 6px 8px; }}
      th {{ background: #eee; }}
    </style>
  </head>
  <body>
    <h1>{context.get('title','Report')}</h1>
    <p><strong>Generated at:</strong> {context.get('generated_at')}</p>
    <h2>Data</h2>
    <table>
      <tr><th>Field</th><th>Value</th></tr>
      {''.join(f'<tr><td>{k}</td><td>{v}</td></tr>' for k, v in list(context.items())[:50])}
    </table>
  </body>
  </html>
        """.strip()
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)
        return filename


export_service = ExportService()
