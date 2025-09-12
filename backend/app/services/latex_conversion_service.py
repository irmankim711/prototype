"""
LaTeX Conversion Service
Handles conversion of LaTeX files to PDF and DOCX formats
"""

import os
import subprocess
import tempfile
import shutil
import logging
from pathlib import Path
from typing import Tuple, Optional, Dict, Any
from datetime import datetime
import uuid

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

class LaTeXConversionService:
    """Service for converting LaTeX files to various formats"""
    
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or os.path.join(os.getcwd(), 'static', 'generated')
        self.temp_dir = os.path.join(os.getcwd(), 'temp', 'latex_conversion')
        self.ensure_directories()
    
    def ensure_directories(self):
        """Ensure required directories exist"""
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
        Path(self.temp_dir).mkdir(parents=True, exist_ok=True)
    
    def convert_latex_to_pdf(self, latex_file_path: str, output_filename: str = None) -> Tuple[str, int]:
        """
        Convert LaTeX file to PDF using pdflatex
        
        Args:
            latex_file_path: Path to the LaTeX file
            output_filename: Optional output filename
            
        Returns:
            Tuple of (file_path, file_size)
        """
        try:
            if not os.path.exists(latex_file_path):
                raise FileNotFoundError(f"LaTeX file not found: {latex_file_path}")
            
            # Generate output filename if not provided
            if not output_filename:
                base_name = Path(latex_file_path).stem
                output_filename = f"{base_name}.pdf"
            
            # Create temporary working directory
            with tempfile.TemporaryDirectory(dir=self.temp_dir) as temp_dir:
                # Copy LaTeX file to temp directory
                temp_latex_path = os.path.join(temp_dir, Path(latex_file_path).name)
                shutil.copy2(latex_file_path, temp_latex_path)
                
                # Change to temp directory for compilation
                original_dir = os.getcwd()
                os.chdir(temp_dir)
                
                try:
                    # Run pdflatex compilation
                    latex_filename = Path(latex_file_path).name
                    result = subprocess.run([
                        'pdflatex',
                        '-interaction=nonstopmode',
                        '-output-directory=' + temp_dir,
                        latex_filename
                    ], capture_output=True, text=True, timeout=60)
                    
                    if result.returncode != 0:
                        logger.warning(f"pdflatex compilation warnings: {result.stderr}")
                    
                    # Look for generated PDF
                    pdf_path = os.path.join(temp_dir, f"{Path(latex_filename).stem}.pdf")
                    
                    if not os.path.exists(pdf_path):
                        # Try alternative compilation with multiple passes
                        logger.info("First compilation failed, trying multiple passes...")
                        
                        # First pass
                        subprocess.run([
                            'pdflatex',
                            '-interaction=nonstopmode',
                            '-output-directory=' + temp_dir,
                            latex_filename
                        ], capture_output=True, text=True, timeout=60)
                        
                        # Second pass for references
                        subprocess.run([
                            'pdflatex',
                            '-interaction=nonstopmode',
                            '-output-directory=' + temp_dir,
                            latex_filename
                        ], capture_output=True, text=True, timeout=60)
                        
                        pdf_path = os.path.join(temp_dir, f"{Path(latex_filename).stem}.pdf")
                    
                    if not os.path.exists(pdf_path):
                        raise RuntimeError("PDF generation failed after multiple compilation attempts")
                    
                    # Copy PDF to output directory
                    final_pdf_path = os.path.join(self.output_dir, output_filename)
                    shutil.copy2(pdf_path, final_pdf_path)
                    
                    # Get file size
                    file_size = os.path.getsize(final_pdf_path)
                    
                    logger.info(f"Successfully converted LaTeX to PDF: {final_pdf_path}")
                    return final_pdf_path, file_size
                    
                finally:
                    os.chdir(original_dir)
                    
        except subprocess.TimeoutExpired:
            logger.error("LaTeX compilation timed out")
            raise RuntimeError("LaTeX compilation timed out")
        except subprocess.CalledProcessError as e:
            logger.error(f"LaTeX compilation failed: {e}")
            raise RuntimeError(f"LaTeX compilation failed: {e}")
        except Exception as e:
            logger.error(f"Error converting LaTeX to PDF: {str(e)}")
            raise RuntimeError(f"Failed to convert LaTeX to PDF: {str(e)}")
    
    def convert_latex_to_docx(self, latex_file_path: str, output_filename: str = None) -> Tuple[str, int]:
        """
        Convert LaTeX file to DOCX by parsing content and creating Word document
        
        Args:
            latex_file_path: Path to the LaTeX file
            output_filename: Optional output filename
            
        Returns:
            Tuple of (file_path, file_size)
        """
        try:
            if not os.path.exists(latex_file_path):
                raise FileNotFoundError(f"LaTeX file not found: {latex_file_path}")
            
            # Generate output filename if not provided
            if not output_filename:
                base_name = Path(latex_file_path).stem
                output_filename = f"{base_name}.docx"
            
            # Read LaTeX content
            with open(latex_file_path, 'r', encoding='utf-8') as f:
                latex_content = f.read()
            
            # Create Word document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = Path(latex_file_path).stem
            doc.core_properties.author = "Automated Report System"
            doc.core_properties.created = datetime.now()
            
            # Parse LaTeX content and convert to Word
            self._parse_latex_to_docx(latex_content, doc)
            
            # Save document
            final_docx_path = os.path.join(self.output_dir, output_filename)
            doc.save(final_docx_path)
            
            # Get file size
            file_size = os.path.getsize(final_docx_path)
            
            logger.info(f"Successfully converted LaTeX to DOCX: {final_docx_path}")
            return final_docx_path, file_size
            
        except Exception as e:
            logger.error(f"Error converting LaTeX to DOCX: {str(e)}")
            raise RuntimeError(f"Failed to convert LaTeX to DOCX: {str(e)}")
    
    def _parse_latex_to_docx(self, latex_content: str, doc: Document):
        """Parse LaTeX content and add it to Word document"""
        try:
            # Split content into lines
            lines = latex_content.split('\n')
            
            current_section = None
            in_environment = False
            environment_content = []
            
            for line in lines:
                line = line.strip()
                
                # Skip empty lines and comments
                if not line or line.startswith('%'):
                    continue
                
                # Handle document class and packages
                if line.startswith('\\documentclass') or line.startswith('\\usepackage'):
                    continue
                
                # Handle title
                if line.startswith('\\title{'):
                    title = line[7:-1]  # Remove \title{ and }
                    doc.add_heading(title, 0)
                    continue
                
                # Handle author
                if line.startswith('\\author{'):
                    author = line[8:-1]  # Remove \author{ and }
                    doc.add_paragraph(f"Author: {author}")
                    continue
                
                # Handle date
                if line.startswith('\\date{'):
                    date = line[6:-1]  # Remove \date{ and }
                    doc.add_paragraph(f"Date: {date}")
                    continue
                
                # Handle sections
                if line.startswith('\\section{'):
                    section_title = line[9:-1]  # Remove \section{ and }
                    doc.add_heading(section_title, 1)
                    current_section = section_title
                    continue
                
                if line.startswith('\\subsection{'):
                    subsection_title = line[12:-1]  # Remove \subsection{ and }
                    doc.add_heading(subsection_title, 2)
                    continue
                
                # Handle environments
                if line.startswith('\\begin{'):
                    env_name = line[7:-1]  # Remove \begin{ and }
                    in_environment = True
                    environment_content = []
                    continue
                
                if line.startswith('\\end{'):
                    in_environment = False
                    # Process collected environment content
                    if environment_content:
                        self._process_environment_content(env_name, environment_content, doc)
                    environment_content = []
                    continue
                
                if in_environment:
                    environment_content.append(line)
                    continue
                
                # Handle regular text
                if line and not line.startswith('\\'):
                    # Clean up LaTeX commands
                    clean_text = self._clean_latex_commands(line)
                    if clean_text.strip():
                        doc.add_paragraph(clean_text)
            
            # Add generation info
            doc.add_paragraph(f"\nGenerated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
        except Exception as e:
            logger.warning(f"Error parsing LaTeX content: {str(e)}")
            # Add fallback content
            doc.add_paragraph("LaTeX content could not be fully parsed. Please refer to the original LaTeX file.")
    
    def _clean_latex_commands(self, text: str) -> str:
        """Clean up LaTeX commands from text"""
        import re
        
        # Remove LaTeX commands
        text = re.sub(r'\\[a-zA-Z]+(\{[^}]*\})?', '', text)
        
        # Remove LaTeX symbols
        text = text.replace('\\&', '&')
        text = text.replace('\\%', '%')
        text = text.replace('\\$', '$')
        text = text.replace('\\#', '#')
        text = text.replace('\\_', '_')
        text = text.replace('\\{', '{')
        text = text.replace('\\}', '}')
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def _process_environment_content(self, env_name: str, content: list, doc: Document):
        """Process content from LaTeX environments"""
        if env_name == 'itemize':
            for item in content:
                if item.strip().startswith('\\item'):
                    item_text = item[6:].strip()  # Remove \item
                    if item_text:
                        doc.add_paragraph(item_text, style='List Bullet')
        
        elif env_name == 'enumerate':
            for item in content:
                if item.strip().startswith('\\item'):
                    item_text = item[6:].strip()  # Remove \item
                    if item_text:
                        doc.add_paragraph(item_text, style='List Number')
        
        elif env_name == 'table':
            # Simple table handling
            table_data = []
            for line in content:
                if '&' in line and not line.startswith('\\'):
                    # Split by & and clean up
                    row = [cell.strip().replace('\\\\', '') for cell in line.split('&')]
                    table_data.append(row)
            
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
        
        else:
            # Generic environment handling
            for line in content:
                if line.strip() and not line.startswith('\\'):
                    clean_text = self._clean_latex_commands(line)
                    if clean_text:
                        doc.add_paragraph(clean_text)
    
    def cleanup_temp_files(self, max_age_hours: int = 24):
        """Clean up temporary files older than specified age"""
        try:
            cutoff_time = datetime.now().timestamp() - (max_age_hours * 3600)
            
            for item in os.listdir(self.temp_dir):
                item_path = os.path.join(self.temp_dir, item)
                if os.path.isfile(item_path):
                    if os.path.getmtime(item_path) < cutoff_time:
                        os.remove(item_path)
                        logger.info(f"Cleaned up temp file: {item_path}")
                elif os.path.isdir(item_path):
                    if os.path.getmtime(item_path) < cutoff_time:
                        shutil.rmtree(item_path)
                        logger.info(f"Cleaned up temp directory: {item_path}")
                        
        except Exception as e:
            logger.error(f"Error cleaning up temp files: {str(e)}")

# Global instance
latex_conversion_service = LaTeXConversionService()
