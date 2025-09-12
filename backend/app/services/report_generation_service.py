"""
Enhanced Report Generation Service
Handles generation of PDF, DOCX, and Excel reports with production-ready features
"""

import os
import json
import uuid
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import logging

# Report generation libraries
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
import pandas as pd

from .. import db
from ..models import Report, Form, FormSubmission, User
from ..core.exceptions import ReportGenerationError
from .latex_conversion_service import latex_conversion_service

logger = logging.getLogger(__name__)

class ReportGenerationService:
    """Service for generating various types of reports"""
    
    def __init__(self, upload_folder: str = None):
        self.upload_folder = upload_folder or os.path.join(os.getcwd(), 'uploads', 'reports')
        self.static_dir = os.path.join(os.getcwd(), 'static', 'generated')
        self.ensure_upload_directory()
    
    def ensure_upload_directory(self):
        """Ensure the upload directory exists"""
        Path(self.upload_folder).mkdir(parents=True, exist_ok=True)
        Path(self.static_dir).mkdir(parents=True, exist_ok=True)
    
    def generate_timestamp_filename(self, prefix: str, extension: str) -> str:
        """Generate a unique filename with timestamp"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        return f"{prefix}_{timestamp}_{unique_id}.{extension}"
    
    def generate_pdf_report(self, data: Dict[str, Any], config: Dict[str, Any]) -> Tuple[str, int]:
        """
        Generate a PDF report using reportlab
        
        Args:
            data: Data to include in the report
            config: Report configuration (title, style, etc.)
            
        Returns:
            Tuple of (file_path, file_size)
        """
        try:
            filename = self.generate_timestamp_filename("report", "pdf")
            file_path = os.path.join(self.upload_folder, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            story = []
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.darkblue
            )
            
            # Add title
            title = Paragraph(config.get('title', 'Generated Report'), title_style)
            story.append(title)
            story.append(Spacer(1, 20))
            
            # Add generation info
            info_style = styles['Normal']
            generation_info = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            story.append(Paragraph(generation_info, info_style))
            story.append(Spacer(1, 20))
            
            # Add data content
            if isinstance(data, dict):
                story.extend(self._add_dict_to_pdf(data, styles))
            elif isinstance(data, list):
                story.extend(self._add_list_to_pdf(data, styles))
            else:
                story.append(Paragraph(str(data), info_style))
            
            # Build PDF
            doc.build(story)
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            logger.info(f"PDF report generated successfully: {file_path} ({file_size} bytes)")
            return file_path, file_size
            
        except Exception as e:
            logger.error(f"Error generating PDF report: {str(e)}")
            raise ReportGenerationError(f"Failed to generate PDF report: {str(e)}")
    
    def _add_dict_to_pdf(self, data: Dict[str, Any], styles) -> List:
        """Add dictionary data to PDF story"""
        story = []
        normal_style = styles['Normal']
        
        for key, value in data.items():
            if isinstance(value, (dict, list)):
                # Add key as subheading
                subheading = Paragraph(f"<b>{key}:</b>", normal_style)
                story.append(subheading)
                story.append(Spacer(1, 10))
                
                if isinstance(value, dict):
                    story.extend(self._add_dict_to_pdf(value, styles))
                else:
                    story.extend(self._add_list_to_pdf(value, styles))
            else:
                # Simple key-value pair
                content = f"<b>{key}:</b> {str(value)}"
                story.append(Paragraph(content, normal_style))
                story.append(Spacer(1, 5))
        
        return story
    
    def _add_list_to_pdf(self, data: List[Any], styles) -> List:
        """Add list data to PDF story"""
        story = []
        normal_style = styles['Normal']
        
        for i, item in enumerate(data):
            if isinstance(item, dict):
                story.append(Paragraph(f"<b>Item {i+1}:</b>", normal_style))
                story.append(Spacer(1, 5))
                story.extend(self._add_dict_to_pdf(item, styles))
            else:
                story.append(Paragraph(f"• {str(item)}", normal_style))
                story.append(Spacer(1, 3))
        
        return story
    
    def generate_docx_report(self, data: Dict[str, Any], config: Dict[str, Any]) -> Tuple[str, int]:
        """
        Generate a DOCX report using python-docx
        
        Args:
            data: Data to include in the report
            config: Report configuration (title, style, etc.)
            
        Returns:
            Tuple of (file_path, file_size)
        """
        try:
            filename = self.generate_timestamp_filename("report", "docx")
            file_path = os.path.join(self.upload_folder, filename)
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(config.get('title', 'Generated Report'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add generation info
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()
            
            # Add data content
            if isinstance(data, dict):
                self._add_dict_to_docx(doc, data)
            elif isinstance(data, list):
                self._add_list_to_docx(doc, data)
            else:
                doc.add_paragraph(str(data))
            
            # Save document
            doc.save(file_path)
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            logger.info(f"DOCX report generated successfully: {file_path} ({file_size} bytes)")
            return file_path, file_size
            
        except Exception as e:
            logger.error(f"Error generating DOCX report: {str(e)}")
            raise ReportGenerationError(f"Failed to generate DOCX report: {str(e)}")
    
    def _add_dict_to_docx(self, doc: Document, data: Dict[str, Any]):
        """Add dictionary data to DOCX document"""
        for key, value in data.items():
            if isinstance(value, (dict, list)):
                # Add key as subheading
                doc.add_heading(key, level=2)
                
                if isinstance(value, dict):
                    self._add_dict_to_docx(doc, value)
                else:
                    self._add_list_to_docx(doc, value)
            else:
                # Simple key-value pair
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
    
    def _add_list_to_docx(self, doc: Document, data: List[Any]):
        """Add list data to DOCX document"""
        for i, item in enumerate(data):
            if isinstance(item, dict):
                doc.add_heading(f"Item {i+1}", level=3)
                self._add_dict_to_docx(doc, item)
            else:
                doc.add_paragraph(f"• {str(item)}", style='List Bullet')
    
    def generate_excel_report(self, data: List[Dict[str, Any]], config: Dict[str, Any]) -> Tuple[str, int]:
        """
        Generate an Excel report using openpyxl
        
        Args:
            data: List of dictionaries representing rows
            config: Report configuration (sheet name, headers, etc.)
            
        Returns:
            Tuple of (file_path, file_size)
        """
        try:
            filename = self.generate_timestamp_filename("report", "xlsx")
            file_path = os.path.join(self.upload_folder, filename)
            
            # Create workbook and worksheet
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = config.get('sheet_name', 'Report Data')
            
            if not data:
                # Empty report
                ws['A1'] = "No data available"
                wb.save(file_path)
                file_size = os.path.getsize(file_path)
                return file_path, file_size
            
            # Get headers from first row
            headers = list(data[0].keys())
            
            # Style for headers
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")
            
            # Add headers
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
            
            # Add data rows
            for row_idx, row_data in enumerate(data, 2):
                for col_idx, header in enumerate(headers, 1):
                    value = row_data.get(header, '')
                    ws.cell(row=row_idx, column=col_idx, value=value)
            
            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # Add borders
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            for row in ws.iter_rows(min_row=1, max_row=len(data)+1, min_col=1, max_col=len(headers)):
                for cell in row:
                    cell.border = thin_border
            
            # Save workbook
            wb.save(file_path)
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            logger.info(f"Excel report generated successfully: {file_path} ({file_size} bytes)")
            return file_path, file_size
            
        except Exception as e:
            logger.error(f"Error generating Excel report: {str(e)}")
            raise ReportGenerationError(f"Failed to generate Excel report: {str(e)}")
    
    def generate_latex_based_report(self, report_id: int, latex_file_path: str, config: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a report from LaTeX template with automatic conversion to PDF and DOCX
        
        Args:
            report_id: ID of the report record
            latex_file_path: Path to the LaTeX file
            config: Report configuration
            
        Returns:
            Dictionary with file paths and metadata
        """
        try:
            # Update report status to generating
            report = Report.query.get(report_id)
            if not report:
                raise ReportGenerationError(f"Report {report_id} not found")
            
            report.update_status('generating', progress=10)
            db.session.commit()
            
            # Generate base filename
            base_filename = Path(latex_file_path).stem
            
            # Convert LaTeX to PDF
            logger.info(f"Converting LaTeX to PDF for report {report_id}")
            pdf_filename = f"{base_filename}.pdf"
            pdf_path, pdf_size = latex_conversion_service.convert_latex_to_pdf(
                latex_file_path, pdf_filename
            )
            report.update_status('generating', progress=50)
            db.session.commit()
            
            # Convert LaTeX to DOCX
            logger.info(f"Converting LaTeX to DOCX for report {report_id}")
            docx_filename = f"{base_filename}.docx"
            docx_path, docx_size = latex_conversion_service.convert_latex_to_docx(
                latex_file_path, docx_filename
            )
            report.update_status('generating', progress=80)
            db.session.commit()
            
            # Generate Excel report from data if available
            excel_path = None
            excel_size = None
            if config.get('include_excel', True) and config.get('data'):
                try:
                    excel_filename = f"{base_filename}.xlsx"
                    excel_path, excel_size = self.generate_excel_report(
                        config.get('data', {}).get('submissions', []),
                        config.get('excel_config', {})
                    )
                    # Rename to match base filename
                    if excel_path:
                        new_excel_path = os.path.join(self.static_dir, excel_filename)
                        if os.path.exists(excel_path):
                            os.rename(excel_path, new_excel_path)
                            excel_path = new_excel_path
                except Exception as e:
                    logger.warning(f"Excel generation failed for report {report_id}: {str(e)}")
            
            report.update_status('generating', progress=90)
            db.session.commit()
            
            # Update report with file information
            report.pdf_file_path = pdf_path
            report.docx_file_path = docx_path
            report.excel_file_path = excel_path
            report.pdf_file_size = pdf_size
            report.docx_file_size = docx_size
            report.excel_file_size = excel_size
            
            # Generate download URLs
            base_url = config.get('base_url', 'http://localhost:5000')
            report.pdf_download_url = f"{base_url}/api/reports/{report_id}/download/pdf"
            report.docx_download_url = f"{base_url}/api/reports/{report_id}/download/docx"
            if excel_path:
                report.excel_download_url = f"{base_url}/api/reports/{report_id}/download/excel"
            
            # Mark as completed
            report.update_status('completed', progress=100)
            report.generated_data = config.get('data', {})
            report.report_config = config
            db.session.commit()
            
            logger.info(f"LaTeX-based report {report_id} generated successfully")
            
            return {
                'report_id': report_id,
                'status': 'completed',
                'latex_source': latex_file_path,
                'pdf_file_path': pdf_path,
                'docx_file_path': docx_path,
                'excel_file_path': excel_path,
                'pdf_download_url': report.pdf_download_url,
                'docx_download_url': report.docx_download_url,
                'excel_download_url': report.excel_download_url,
                'file_sizes': {
                    'pdf': pdf_size,
                    'docx': docx_size,
                    'excel': excel_size
                }
            }
            
        except Exception as e:
            # Update report status to failed
            report = Report.query.get(report_id)
            if report:
                report.update_status('failed', error_message=str(e))
                db.session.commit()
            
            logger.error(f"Error generating LaTeX-based report {report_id}: {str(e)}")
            raise ReportGenerationError(f"Failed to generate LaTeX-based report: {str(e)}")

    def generate_comprehensive_report(self, report_id: int, data: Dict[str, Any], config: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a comprehensive report with all formats (PDF, DOCX, Excel)
        Now includes LaTeX conversion if LaTeX source is provided
        
        Args:
            report_id: ID of the report record
            data: Data to include in the report
            config: Report configuration
            
        Returns:
            Dictionary with file paths and metadata
        """
        try:
            # Check if this is a LaTeX-based report
            latex_source = config.get('latex_source') or config.get('latex_file_path')
            
            if latex_source and os.path.exists(latex_source):
                logger.info(f"Generating LaTeX-based comprehensive report {report_id}")
                return self.generate_latex_based_report(report_id, latex_source, config)
            
            # Update report status to generating
            report = Report.query.get(report_id)
            if not report:
                raise ReportGenerationError(f"Report {report_id} not found")
            
            report.update_status('generating', progress=10)
            db.session.commit()
            
            # Generate Excel report first (usually fastest)
            excel_path, excel_size = self.generate_excel_report(
                data.get('submissions', []), 
                config.get('excel_config', {})
            )
            report.update_status('generating', progress=40)
            db.session.commit()
            
            # Generate PDF report
            pdf_path, pdf_size = self.generate_pdf_report(data, config)
            report.update_status('generating', progress=70)
            db.session.commit()
            
            # Generate DOCX report
            docx_path, docx_size = self.generate_docx_report(data, config)
            report.update_status('generating', progress=90)
            db.session.commit()
            
            # Update report with file information
            report.pdf_file_path = pdf_path
            report.docx_file_path = docx_path
            report.excel_file_path = excel_path
            report.pdf_file_size = pdf_size
            report.docx_file_size = docx_size
            report.excel_file_size = excel_size
            
            # Generate download URLs
            base_url = config.get('base_url', 'http://localhost:5000')
            report.pdf_download_url = f"{base_url}/api/reports/{report_id}/download/pdf"
            report.docx_download_url = f"{base_url}/api/reports/{report_id}/download/docx"
            report.excel_download_url = f"{base_url}/api/reports/{report_id}/download/excel"
            
            # Mark as completed
            report.update_status('completed', progress=100)
            report.generated_data = data
            report.report_config = config
            db.session.commit()
            
            logger.info(f"Comprehensive report {report_id} generated successfully")
            
            return {
                'report_id': report_id,
                'status': 'completed',
                'pdf_file_path': pdf_path,
                'docx_file_path': docx_path,
                'excel_file_path': excel_path,
                'pdf_download_url': report.pdf_download_url,
                'docx_download_url': report.docx_download_url,
                'excel_download_url': report.excel_download_url,
                'file_sizes': {
                    'pdf': pdf_size,
                    'docx': docx_size,
                    'excel': excel_size
                }
            }
            
        except Exception as e:
            # Update report status to failed
            report = Report.query.get(report_id)
            if report:
                report.update_status('failed', error_message=str(e))
                db.session.commit()
            
            logger.error(f"Error generating comprehensive report {report_id}: {str(e)}")
            raise ReportGenerationError(f"Failed to generate comprehensive report: {str(e)}")
    
    def cleanup_old_files(self, days_old: int = 30):
        """Clean up old report files"""
        try:
            cutoff_date = datetime.now() - timedelta(days=days_old)
            old_reports = Report.query.filter(
                Report.created_at < cutoff_date,
                Report.status.in_(['completed', 'failed'])
            ).all()
            
            for report in old_reports:
                # Remove PDF file
                if report.pdf_file_path and os.path.exists(report.pdf_file_path):
                    os.remove(report.pdf_file_path)
                    report.pdf_file_path = None
                    report.pdf_file_size = None
                
                # Remove DOCX file
                if report.docx_file_path and os.path.exists(report.docx_file_path):
                    os.remove(report.docx_file_path)
                    report.docx_file_path = None
                    report.docx_file_size = None
                
                # Remove Excel file
                if report.excel_file_path and os.path.exists(report.excel_file_path):
                    os.remove(report.excel_file_path)
                    report.excel_file_path = None
                    report.excel_file_size = None
                
                # Update report
                report.status = 'archived'
                db.session.commit()
            
            logger.info(f"Cleaned up {len(old_reports)} old report files")
            
        except Exception as e:
            logger.error(f"Error cleaning up old files: {str(e)}")

# Global instance
report_generation_service = ReportGenerationService()
