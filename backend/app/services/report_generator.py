"""
Report Generator Service for creating Word documents with charts and AI analysis
Handles the creation of comprehensive reports from form data
"""

import os
import tempfile
from datetime import datetime
from typing import Dict, List, Any, Optional
import base64
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import matplotlib.pyplot as plt
    import pandas as pd
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False


def create_word_report(template_id: str, data: Dict[str, Any], output_path: Optional[str] = None) -> str:
    """
    Create a Word document report from form data
    
    Args:
        template_id: Template identifier (e.g., 'form_analysis')
        data: Report data including form submissions, AI analysis, charts
        output_path: Optional output path for the document
        
    Returns:
        Path to the generated Word document
    """
    if not HAS_DOCX:
        raise ImportError("python-docx is required for Word document generation")
    
    # Create document
    doc = Document()
    
    # Set document styles
    _setup_document_styles(doc)
    
    # Generate report based on template
    if template_id == 'form_analysis':
        _create_form_analysis_report(doc, data)
    else:
        _create_generic_report(doc, data)
    
    # Save document
    if not output_path:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"report_{template_id}_{timestamp}.docx"
        output_dir = os.path.join(os.getenv('UPLOAD_FOLDER', 'uploads'), 'reports')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
    
    doc.save(output_path)
    return output_path


def _setup_document_styles(doc: Document):
    """Setup document styles for professional appearance"""
    styles = doc.styles
    
    # Title style
    if 'Title' not in [style.name for style in styles]:
        title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(20)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Heading styles
    for level in [1, 2, 3]:
        style_name = f'Heading {level}'
        if style_name in [style.name for style in styles]:
            heading_style = styles[style_name]
        else:
            heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.size = Pt(16 - level * 2)
        heading_font.bold = True
        heading_style.paragraph_format.space_after = Pt(6)


def _create_form_analysis_report(doc: Document, data: Dict[str, Any]):
    """Create a comprehensive form analysis report"""
    
    # Title page
    title = doc.add_heading(data.get('title', 'Form Analysis Report'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with form info
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Analysis of {data.get('form_title', 'Form')}")
    run.font.size = Pt(14)
    run.font.italic = True
    
    # Generation date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_run.font.size = Pt(12)
    
    # Page break
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    ai_analysis = data.get('ai_analysis', {})
    summary = ai_analysis.get('summary', 'No summary available')
    doc.add_paragraph(summary)
    
    # Key Metrics
    doc.add_heading('Key Metrics', level=1)
    
    statistics = data.get('statistics', {})
    key_metrics = ai_analysis.get('key_metrics', {})
    
    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.style = 'Light Grid Accent 1'
    
    hdr_cells = metrics_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    # Add metrics rows
    metrics_data = [
        ('Total Submissions', statistics.get('total_submissions', 0)),
        ('Overall Completion Rate', f"{key_metrics.get('overall_completion_rate', 0)}%"),
        ('Data Quality Score', f"{key_metrics.get('data_quality_score', 0)}%"),
        ('Fields Analyzed', key_metrics.get('fields_count', 0)),
    ]
    
    if statistics.get('date_range'):
        date_range = statistics['date_range']
        metrics_data.extend([
            ('First Submission', date_range.get('first_submission', 'N/A')),
            ('Last Submission', date_range.get('last_submission', 'N/A')),
            ('Collection Span', f"{date_range.get('span_days', 0)} days")
        ])
    
    for metric, value in metrics_data:
        row_cells = metrics_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)
    
    # AI Insights
    doc.add_heading('AI-Powered Insights', level=1)
    
    insights = ai_analysis.get('insights', [])
    if insights:
        for i, insight in enumerate(insights, 1):
            para = doc.add_paragraph()
            para.add_run(f"{i}. ").bold = True
            para.add_run(insight)
    else:
        doc.add_paragraph("No insights available.")
    
    # Data Visualization
    charts = data.get('charts', [])
    if charts:
        doc.add_heading('Data Visualization', level=1)
        
        for chart in charts:
            # Add chart title
            doc.add_heading(chart.get('title', 'Chart'), level=2)
            
            # Add chart description
            if chart.get('description'):
                doc.add_paragraph(chart['description'])
            
            # Add chart image
            if chart.get('image_base64'):
                try:
                    image_data = base64.b64decode(chart['image_base64'])
                    image_stream = BytesIO(image_data)
                    doc.add_picture(image_stream, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"Error loading chart: {str(e)}")
            
            doc.add_paragraph()  # Add space after chart
    
    # Field Analysis
    doc.add_heading('Field-by-Field Analysis', level=1)
    
    field_completion = statistics.get('field_completion', {})
    top_responses = statistics.get('top_responses', {})
    
    if field_completion:
        # Create field analysis table
        field_table = doc.add_table(rows=1, cols=3)
        field_table.style = 'Light Grid Accent 1'
        
        hdr_cells = field_table.rows[0].cells
        hdr_cells[0].text = 'Field Name'
        hdr_cells[1].text = 'Completion Rate'
        hdr_cells[2].text = 'Top Response'
        
        for field_name, completion_rate in field_completion.items():
            if field_name not in ['submission_id', 'submitted_at', 'submitter_email']:
                row_cells = field_table.add_row().cells
                row_cells[0].text = field_name
                row_cells[1].text = f"{completion_rate}%"
                
                # Add top response if available
                top_response = top_responses.get(field_name, {})
                if top_response:
                    response_text = f"{top_response.get('value', 'N/A')} ({top_response.get('percentage', 0)}%)"
                    row_cells[2].text = response_text
                else:
                    row_cells[2].text = 'N/A'
    
    # Trends Analysis
    trends = ai_analysis.get('trends', [])
    if trends:
        doc.add_heading('Trends Analysis', level=1)
        
        for trend in trends:
            para = doc.add_paragraph()
            
            # Add trend type indicator
            direction = trend.get('direction', 'neutral')
            if direction == 'up':
                indicator = '📈 '
            elif direction == 'down':
                indicator = '📉 '
            else:
                indicator = '📊 '
            
            para.add_run(indicator).font.size = Pt(14)
            para.add_run(trend.get('description', '')).font.size = Pt(11)
    
    # Recommendations
    recommendations = ai_analysis.get('recommendations', [])
    if recommendations:
        doc.add_heading('Recommendations', level=1)
        
        for i, recommendation in enumerate(recommendations, 1):
            para = doc.add_paragraph()
            para.add_run(f"{i}. ").bold = True
            para.add_run(recommendation)
    
    # Data Quality Assessment
    doc.add_heading('Data Quality Assessment', level=1)
    
    quality_score = key_metrics.get('data_quality_score', 0)
    
    if quality_score >= 80:
        quality_assessment = "Excellent - Data quality is high with good completion rates and consistency."
    elif quality_score >= 60:
        quality_assessment = "Good - Data quality is acceptable with room for improvement in some areas."
    elif quality_score >= 40:
        quality_assessment = "Fair - Data quality needs attention. Consider form optimization."
    else:
        quality_assessment = "Poor - Significant data quality issues detected. Form redesign recommended."
    
    doc.add_paragraph(f"Overall Data Quality Score: {quality_score}%")
    doc.add_paragraph(quality_assessment)
    
    # Technical Details
    doc.add_heading('Technical Details', level=1)
    
    tech_para = doc.add_paragraph()
    tech_para.add_run("Analysis Method: ").bold = True
    tech_para.add_run("AI-powered analysis using OpenAI GPT models combined with statistical analysis")
    
    tech_para = doc.add_paragraph()
    tech_para.add_run("Data Processing: ").bold = True
    tech_para.add_run("Automated data normalization and quality assessment")
    
    tech_para = doc.add_paragraph()
    tech_para.add_run("Report Generation: ").bold = True
    tech_para.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d at %H:%M:%S')}")


def _create_generic_report(doc: Document, data: Dict[str, Any]):
    """Create a generic report template"""
    
    # Title
    doc.add_heading(data.get('title', 'Report'), 0)
    
    # Add data sections
    for key, value in data.items():
        if key != 'title' and not key.startswith('_'):
            doc.add_heading(key.replace('_', ' ').title(), level=1)
            
            if isinstance(value, str):
                doc.add_paragraph(value)
            elif isinstance(value, (list, tuple)):
                for item in value:
                    doc.add_paragraph(str(item), style='List Bullet')
            elif isinstance(value, dict):
                for sub_key, sub_value in value.items():
                    para = doc.add_paragraph()
                    para.add_run(f"{sub_key}: ").bold = True
                    para.add_run(str(sub_value))
            else:
                doc.add_paragraph(str(value))


def create_pdf_report(template_id: str, data: Dict[str, Any], output_path: Optional[str] = None) -> str:
    """
    Create a PDF report using ReportLab
    
    Args:
        template_id: Template identifier
        data: Report data including form fields and metadata
        output_path: Optional output path
        
    Returns:
        Path to generated PDF
    """
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    except ImportError:
        raise ImportError("ReportLab is required for PDF generation. Install with: pip install reportlab")
    
    # Generate output path if not provided
    if not output_path:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"report_{template_id}_{timestamp}.pdf"
        output_dir = os.path.join(os.getenv('UPLOAD_FOLDER', 'uploads'), 'pdfs')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
    
    # Create PDF document
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#2E4057')
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=12,
        spaceBefore=20,
        textColor=colors.HexColor('#2E4057')
    )
    
    # Generate report based on template
    if template_id == 'form_analysis':
        _create_form_analysis_pdf(story, data, styles, title_style, heading_style)
    else:
        _create_generic_pdf(story, data, styles, title_style, heading_style)
    
    # Build PDF
    doc.build(story)
    return output_path


def _create_form_analysis_pdf(story, data: Dict[str, Any], styles, title_style, heading_style):
    """Create form analysis PDF content"""
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    
    # Title
    form_name = data.get('form_name', 'Form Analysis Report')
    title = Paragraph(f"Form Analysis: {form_name}", title_style)
    story.append(title)
    story.append(Spacer(1, 20))
    
    # Metadata section
    metadata_data = [
        ['Report Information', ''],
        ['Generated Date:', data.get('generated_date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))],
        ['Form Name:', data.get('form_name', 'N/A')],
        ['Generated By:', data.get('user_name', 'System')],
        ['Total Entries:', str(data.get('entry_count', 0))],
        ['Response Rate:', f"{data.get('response_rate', 0):.1f}%"],
    ]
    
    metadata_table = Table(metadata_data, colWidths=[2*inch, 3*inch])
    metadata_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E4057')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    story.append(metadata_table)
    story.append(Spacer(1, 30))
    
    # Form data section
    if 'table_data' in data and data['table_data']:
        story.append(Paragraph("Form Responses", heading_style))
        story.append(Spacer(1, 12))
        
        table_data = data['table_data']
        if len(table_data) > 0:
            # Create table with proper column widths
            col_count = len(table_data[0]) if table_data[0] else 1
            col_width = (7 * inch) / col_count  # Distribute width evenly
            
            data_table = Table(table_data, colWidths=[col_width] * col_count)
            data_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4A90A4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            
            story.append(data_table)
            story.append(Spacer(1, 30))
    
    # Summary statistics
    if 'statistics' in data:
        story.append(Paragraph("Summary Statistics", heading_style))
        story.append(Spacer(1, 12))
        
        stats = data['statistics']
        stats_data = [
            ['Statistic', 'Value'],
        ]
        
        for key, value in stats.items():
            stats_data.append([key.replace('_', ' ').title(), str(value)])
        
        stats_table = Table(stats_data, colWidths=[2.5*inch, 2.5*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4A90A4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        story.append(stats_table)
    
    # AI insights section
    if 'ai_insights' in data and data['ai_insights']:
        story.append(Spacer(1, 30))
        story.append(Paragraph("AI-Generated Insights", heading_style))
        story.append(Spacer(1, 12))
        
        insights_text = data['ai_insights']
        if isinstance(insights_text, list):
            for insight in insights_text:
                story.append(Paragraph(f"• {insight}", styles['Normal']))
                story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(insights_text, styles['Normal']))


def _create_generic_pdf(story, data: Dict[str, Any], styles, title_style, heading_style):
    """Create generic PDF content"""
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    
    # Title
    title = Paragraph(data.get('title', 'Report'), title_style)
    story.append(title)
    story.append(Spacer(1, 20))
    
    # Description
    if 'description' in data:
        story.append(Paragraph("Description", heading_style))
        story.append(Paragraph(data['description'], styles['Normal']))
        story.append(Spacer(1, 20))
    
    # Content sections
    if 'content' in data:
        content = data['content']
        if isinstance(content, dict):
            for section_title, section_content in content.items():
                story.append(Paragraph(section_title.replace('_', ' ').title(), heading_style))
                if isinstance(section_content, (list, tuple)):
                    for item in section_content:
                        story.append(Paragraph(f"• {item}", styles['Normal']))
                        story.append(Spacer(1, 6))
                else:
                    story.append(Paragraph(str(section_content), styles['Normal']))
                story.append(Spacer(1, 15))
        else:
            story.append(Paragraph(str(content), styles['Normal']))
    
    # Data table if available
    if 'table_data' in data and data['table_data']:
        story.append(Spacer(1, 20))
        story.append(Paragraph("Data", heading_style))
        story.append(Spacer(1, 12))
        
        table_data = data['table_data']
        if len(table_data) > 0:
            col_count = len(table_data[0]) if table_data[0] else 1
            col_width = (7 * inch) / col_count
            
            data_table = Table(table_data, colWidths=[col_width] * col_count)
            data_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            
            story.append(data_table)


def create_excel_export(data: Dict[str, Any], output_path: Optional[str] = None) -> str:
    """
    Create Excel export of form data
    
    Args:
        data: Form data to export
        output_path: Optional output path
        
    Returns:
        Path to generated Excel file
    """
    if not HAS_MATPLOTLIB:  # Using this as a proxy for pandas availability
        raise ImportError("pandas is required for Excel export")
    
    import pandas as pd
    
    # Create output path
    if not output_path:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"form_data_{timestamp}.xlsx"
        output_dir = os.path.join(os.getenv('UPLOAD_FOLDER', 'uploads'), 'exports')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
    
    # Create Excel writer
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        
        # Export submissions data
        submissions = data.get('submissions', [])
        if submissions:
            df = pd.DataFrame(submissions)
            df.to_excel(writer, sheet_name='Submissions', index=False)
        
        # Export statistics
        statistics = data.get('statistics', {})
        if statistics:
            stats_data = []
            for key, value in statistics.items():
                if isinstance(value, dict):
                    for sub_key, sub_value in value.items():
                        stats_data.append({'Category': key, 'Metric': sub_key, 'Value': sub_value})
                else:
                    stats_data.append({'Category': 'General', 'Metric': key, 'Value': value})
            
            if stats_data:
                stats_df = pd.DataFrame(stats_data)
                stats_df.to_excel(writer, sheet_name='Statistics', index=False)
        
        # Export AI insights
        ai_analysis = data.get('ai_analysis', {})
        if ai_analysis:
            insights_data = []
            
            for insight in ai_analysis.get('insights', []):
                insights_data.append({'Type': 'Insight', 'Content': insight})
            
            for recommendation in ai_analysis.get('recommendations', []):
                insights_data.append({'Type': 'Recommendation', 'Content': recommendation})
            
            if insights_data:
                insights_df = pd.DataFrame(insights_data)
                insights_df.to_excel(writer, sheet_name='AI Analysis', index=False)
    
    return output_path


def validate_report_data(data: Dict[str, Any]) -> List[str]:
    """
    Validate report data structure
    
    Args:
        data: Report data to validate
        
    Returns:
        List of validation errors
    """
    errors = []
    
    # Check required fields
    if not data.get('title'):
        errors.append("Report title is required")
    
    # Check submissions data
    submissions = data.get('submissions', [])
    if not submissions:
        errors.append("No submissions data provided")
    elif not isinstance(submissions, list):
        errors.append("Submissions must be a list")
    
    # Check AI analysis structure
    ai_analysis = data.get('ai_analysis', {})
    if ai_analysis and not isinstance(ai_analysis, dict):
        errors.append("AI analysis must be a dictionary")
    
    # Check charts data
    charts = data.get('charts', [])
    if charts and not isinstance(charts, list):
        errors.append("Charts must be a list")
    
    for i, chart in enumerate(charts):
        if not isinstance(chart, dict):
            errors.append(f"Chart {i+1} must be a dictionary")
        elif not chart.get('title'):
            errors.append(f"Chart {i+1} missing title")
    
    return errors


# Convenience function for the existing codebase
def generate_report_suggestions(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Generate report suggestions (backward compatibility)
    
    Args:
        data: Form data
        
    Returns:
        Dictionary with suggestions
    """
    # This function maintains compatibility with existing code
    # while providing basic suggestions
    
    suggestions = {
        'insights': [],
        'recommendations': [],
        'summary': 'Basic report analysis completed'
    }
    
    if isinstance(data, dict) and 'submissions' in data:
        submissions = data['submissions']
        if isinstance(submissions, list):
            suggestions['insights'].append(f"Analyzed {len(submissions)} form submissions")
            
            if len(submissions) > 100:
                suggestions['recommendations'].append("Consider implementing data archival for better performance")
            elif len(submissions) < 10:
                suggestions['recommendations'].append("Increase form promotion to gather more responses")
    
    return suggestions


def format_report_data_for_pdf(report) -> List[List[str]]:
    """
    Format report data for PDF table generation
    
    Args:
        report: Report model instance with form data
        
    Returns:
        List of lists representing table data with headers
    """
    try:
        # Initialize table data with headers
        table_data = []
        
        # Get form submissions data
        if hasattr(report, 'data') and report.data:
            # If report.data contains processed form data
            data = report.data
            
            if isinstance(data, dict):
                # Handle different data structures
                if 'submissions' in data:
                    submissions = data['submissions']
                    if submissions and len(submissions) > 0:
                        # Extract headers from first submission
                        first_submission = submissions[0]
                        if isinstance(first_submission, dict):
                            headers = list(first_submission.keys())
                            table_data.append(headers)
                            
                            # Add data rows
                            for submission in submissions:
                                row = [str(submission.get(header, '')) for header in headers]
                                table_data.append(row)
                
                elif 'fields' in data:
                    # Handle field-based data structure
                    fields = data['fields']
                    if fields:
                        table_data.append(['Field', 'Value'])
                        for field_name, field_value in fields.items():
                            table_data.append([str(field_name), str(field_value)])
                
                else:
                    # Generic key-value structure
                    table_data.append(['Property', 'Value'])
                    for key, value in data.items():
                        if not isinstance(value, (dict, list)):
                            table_data.append([str(key), str(value)])
        
        # If no data found, check if report has form with submissions
        if not table_data and hasattr(report, 'form') and report.form:
            form = report.form
            if hasattr(form, 'submissions') and form.submissions:
                submissions = form.submissions
                if submissions:
                    # Get field names from form schema or first submission
                    headers = []
                    if hasattr(form, 'schema') and form.schema:
                        schema = form.schema
                        if isinstance(schema, dict) and 'fields' in schema:
                            headers = [field.get('name', field.get('id', '')) 
                                     for field in schema['fields'] if isinstance(field, dict)]
                    
                    # If no headers from schema, extract from first submission
                    if not headers and submissions[0].data:
                        headers = list(submissions[0].data.keys())
                    
                    if headers:
                        table_data.append(headers)
                        
                        # Add submission data
                        for submission in submissions:
                            if submission.data:
                                row = [str(submission.data.get(header, '')) for header in headers]
                                table_data.append(row)
        
        # Fallback: create minimal data if nothing found
        if not table_data:
            table_data = [
                ['Report ID', 'Title', 'Status', 'Created'],
                [str(report.id), report.title or 'N/A', report.status or 'N/A', 
                 report.created_at.strftime('%Y-%m-%d') if report.created_at else 'N/A']
            ]
        
        return table_data
        
    except Exception as e:
        # Return error information in case of issues
        return [
            ['Error', 'Details'],
            ['Data Processing Failed', str(e)],
            ['Report ID', str(report.id) if hasattr(report, 'id') else 'Unknown']
        ]
