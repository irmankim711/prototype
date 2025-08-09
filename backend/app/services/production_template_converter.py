"""
Production Template Converter Service - ZERO MOCK DATA
Converts real database data to template placeholders for Temp1.docx and other templates
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional, Union
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from dataclasses import dataclass
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

@dataclass
class TemplateMapping:
    """Template field mapping configuration"""
    placeholder: str
    data_source: str
    data_type: str = "string"
    default_value: str = ""
    transformation: str = None

class ProductionTemplateConverter:
    """Production Template Converter - NO MOCK DATA"""
    
    def __init__(self, db_session: Session):
        self.db_session = db_session
        self.template_mappings = self._load_template_mappings()
        self.templates_dir = os.getenv('TEMPLATES_STORAGE_PATH', './templates')
        self.output_dir = os.getenv('REPORTS_STORAGE_PATH', './reports')
        
        # Ensure directories exist
        os.makedirs(self.templates_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)

    def _load_template_mappings(self) -> Dict[str, Dict[str, TemplateMapping]]:
        """Load real template mappings for different templates"""
        return {
            'Temp1.docx': {
                # Program Information Mappings
                '{{program.title}}': TemplateMapping(
                    placeholder='{{program.title}}',
                    data_source='program.title',
                    data_type='string',
                    default_value='Program Title Not Available'
                ),
                '{{program.date}}': TemplateMapping(
                    placeholder='{{program.date}}',
                    data_source='program.start_date',
                    data_type='date',
                    transformation='format_date_range'
                ),
                '{{program.time}}': TemplateMapping(
                    placeholder='{{program.time}}',
                    data_source='program.start_date,program.end_date',
                    data_type='time',
                    transformation='format_time_range'
                ),
                '{{program.location}}': TemplateMapping(
                    placeholder='{{program.location}}',
                    data_source='program.location',
                    data_type='string',
                    default_value='Location To Be Confirmed'
                ),
                '{{program.organizer}}': TemplateMapping(
                    placeholder='{{program.organizer}}',
                    data_source='program.organizer',
                    data_type='string',
                    default_value='Organizing Body'
                ),
                '{{program.speaker}}': TemplateMapping(
                    placeholder='{{program.speaker}}',
                    data_source='program.speaker',
                    data_type='string',
                    default_value='Speaker To Be Announced'
                ),
                '{{program.trainer}}': TemplateMapping(
                    placeholder='{{program.trainer}}',
                    data_source='program.trainer',
                    data_type='string',
                    default_value='Trainer To Be Assigned'
                ),
                '{{program.facilitator}}': TemplateMapping(
                    placeholder='{{program.facilitator}}',
                    data_source='program.facilitator',
                    data_type='string',
                    default_value='Facilitator To Be Assigned'
                ),
                '{{program.background}}': TemplateMapping(
                    placeholder='{{program.background}}',
                    data_source='program.background',
                    data_type='text',
                    default_value='Program background information will be provided.'
                ),
                '{{program.objectives}}': TemplateMapping(
                    placeholder='{{program.objectives}}',
                    data_source='program.objectives',
                    data_type='text',
                    default_value='Program objectives will be defined.'
                ),
                
                # Participant Count Mappings
                '{{program.male_participants}}': TemplateMapping(
                    placeholder='{{program.male_participants}}',
                    data_source='participants.gender',
                    data_type='count',
                    transformation='count_male_participants'
                ),
                '{{program.female_participants}}': TemplateMapping(
                    placeholder='{{program.female_participants}}',
                    data_source='participants.gender',
                    data_type='count',
                    transformation='count_female_participants'
                ),
                '{{program.total_participants}}': TemplateMapping(
                    placeholder='{{program.total_participants}}',
                    data_source='participants.id',
                    data_type='count',
                    transformation='count_total_participants'
                ),
                
                # Attendance Mappings
                '{{attendance.total_attended}}': TemplateMapping(
                    placeholder='{{attendance.total_attended}}',
                    data_source='attendance_records.day_1_status',
                    data_type='count',
                    transformation='count_attended'
                ),
                '{{attendance.total_absent}}': TemplateMapping(
                    placeholder='{{attendance.total_absent}}',
                    data_source='attendance_records.day_1_status',
                    data_type='count',
                    transformation='count_absent'
                ),
                
                # Staff Signature Mappings
                '{{signature.consultant}}': TemplateMapping(
                    placeholder='{{signature.consultant}}',
                    data_source='program.consultant_name',
                    data_type='string',
                    default_value='Consultant Name'
                ),
                '{{signature.executive}}': TemplateMapping(
                    placeholder='{{signature.executive}}',
                    data_source='program.executive_name',
                    data_type='string',
                    default_value='Executive Name'
                ),
                '{{signature.head}}': TemplateMapping(
                    placeholder='{{signature.head}}',
                    data_source='program.head_name',
                    data_type='string',
                    default_value='Department Head'
                )
            }
        }

    def generate_report_from_template(
        self, 
        template_name: str, 
        program_id: int, 
        output_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """Generate real report from template using database data"""
        try:
            # Get real program data
            program_data = self._get_program_data(program_id)
            if not program_data:
                raise Exception(f"Program with ID {program_id} not found")
            
            # Get real participant data
            participants_data = self._get_participants_data(program_id)
            
            # Get real attendance data
            attendance_data = self._get_attendance_data(program_id)
            
            # Load template
            template_path = os.path.join(self.templates_dir, template_name)
            if not os.path.exists(template_path):
                raise Exception(f"Template {template_name} not found")
            
            # Generate context from real data
            template_context = self._generate_template_context(
                program_data, 
                participants_data, 
                attendance_data
            )
            
            # Process template
            output_path = self._process_template(
                template_path, 
                template_context, 
                output_filename or f"report_{program_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            
            return {
                'status': 'success',
                'output_path': output_path,
                'program_id': program_id,
                'template_used': template_name,
                'generated_at': datetime.utcnow().isoformat(),
                'context_summary': {
                    'program_title': template_context.get('program', {}).get('title'),
                    'participants_count': len(participants_data),
                    'attendance_records': len(attendance_data)
                }
            }
            
        except Exception as e:
            logger.error(f"Error generating report from template: {str(e)}")
            return {
                'status': 'error',
                'message': str(e),
                'program_id': program_id,
                'template_name': template_name
            }

    def _get_program_data(self, program_id: int) -> Optional[Dict[str, Any]]:
        """Get real program data from database"""
        try:
            from ..models.production_models import Program
            
            program = self.db_session.query(Program).filter_by(id=program_id).first()
            if not program:
                return None
            
            return {
                'id': program.id,
                'title': program.title,
                'description': program.description,
                'start_date': program.start_date,
                'end_date': program.end_date,
                'location': program.location,
                'organizer': program.organizer,
                'speaker': program.speaker,
                'trainer': program.trainer,
                'facilitator': program.facilitator,
                'background': program.background,
                'objectives': program.objectives,
                'capacity': program.capacity,
                'status': program.status,
                'created_at': program.created_at
            }
            
        except Exception as e:
            logger.error(f"Error fetching program data: {str(e)}")
            return None

    def _get_participants_data(self, program_id: int) -> List[Dict[str, Any]]:
        """Get real participants data from database"""
        try:
            from ..models.production_models import Participant
            
            participants = self.db_session.query(Participant).filter_by(
                program_id=program_id
            ).all()
            
            return [
                {
                    'id': p.id,
                    'full_name': p.full_name,
                    'identification_number': p.identification_number,
                    'email': p.email,
                    'phone': p.phone,
                    'gender': p.gender,
                    'age': p.age,
                    'organization': p.organization,
                    'position': p.position,
                    'department': p.department,
                    'registration_date': p.registration_date,
                    'status': p.status
                }
                for p in participants
            ]
            
        except Exception as e:
            logger.error(f"Error fetching participants data: {str(e)}")
            return []

    def _get_attendance_data(self, program_id: int) -> List[Dict[str, Any]]:
        """Get real attendance data from database"""
        try:
            from ..models.production_models import AttendanceRecord, Participant
            
            attendance_records = self.db_session.query(
                AttendanceRecord, Participant
            ).join(
                Participant, AttendanceRecord.participant_id == Participant.id
            ).filter(
                AttendanceRecord.program_id == program_id
            ).all()
            
            return [
                {
                    'participant_id': record.AttendanceRecord.participant_id,
                    'participant_name': record.Participant.full_name,
                    'day_1_status': record.AttendanceRecord.day_1_status,
                    'day_2_status': record.AttendanceRecord.day_2_status,
                    'day_3_status': record.AttendanceRecord.day_3_status,
                    'total_hours_attended': float(record.AttendanceRecord.total_hours_attended or 0),
                    'attendance_percentage': float(record.AttendanceRecord.attendance_percentage or 0),
                    'recorded_at': record.AttendanceRecord.recorded_at
                }
                for record in attendance_records
            ]
            
        except Exception as e:
            logger.error(f"Error fetching attendance data: {str(e)}")
            return []

    def _generate_template_context(
        self, 
        program_data: Dict, 
        participants_data: List[Dict], 
        attendance_data: List[Dict]
    ) -> Dict[str, Any]:
        """Generate template context from real database data"""
        
        context = {
            'program': {
                'title': program_data.get('title', 'Program Title Not Available'),
                'date': self._format_date_range(program_data.get('start_date'), program_data.get('end_date')),
                'time': self._format_time_range(program_data.get('start_date'), program_data.get('end_date')),
                'location': program_data.get('location', 'Location To Be Confirmed'),
                'organizer': program_data.get('organizer', 'Organizing Body'),
                'speaker': program_data.get('speaker', 'Speaker To Be Announced'),
                'trainer': program_data.get('trainer', 'Trainer To Be Assigned'),
                'facilitator': program_data.get('facilitator', 'Facilitator To Be Assigned'),
                'background': program_data.get('background', 'Program background information will be provided.'),
                'objectives': program_data.get('objectives', 'Program objectives will be defined.'),
                'male_participants': self._count_male_participants(participants_data),
                'female_participants': self._count_female_participants(participants_data),
                'total_participants': len(participants_data)
            },
            'participants': [
                {
                    'bil': i + 1,
                    'name': participant['full_name'],
                    'ic': participant['identification_number'],
                    'attendance_day1': self._get_participant_attendance(participant['id'], attendance_data, 'day_1_status'),
                    'attendance_day2': self._get_participant_attendance(participant['id'], attendance_data, 'day_2_status'),
                    'organization': participant.get('organization', ''),
                    'position': participant.get('position', '')
                }
                for i, participant in enumerate(participants_data)
            ],
            'attendance': {
                'total_attended': self._count_attended(attendance_data),
                'total_absent': self._count_absent(attendance_data),
                'attendance_percentage': self._calculate_attendance_percentage(attendance_data)
            },
            'signature': {
                'consultant': {'name': program_data.get('consultant_name', 'Consultant Name')},
                'executive': {'name': program_data.get('executive_name', 'Executive Name')},
                'head': {'name': program_data.get('head_name', 'Department Head')}
            },
            'generated_date': datetime.now().strftime('%d/%m/%Y'),
            'generated_time': datetime.now().strftime('%H:%M')
        }
        
        return context

    def _format_date_range(self, start_date, end_date) -> str:
        """Format date range for template"""
        if not start_date:
            return 'Date To Be Confirmed'
        
        if isinstance(start_date, str):
            try:
                start_date = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
            except:
                return 'Date To Be Confirmed'
        
        if not end_date or start_date.date() == end_date.date():
            return start_date.strftime('%d/%m/%Y')
        else:
            if isinstance(end_date, str):
                try:
                    end_date = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
                except:
                    return start_date.strftime('%d/%m/%Y')
            
            return f"{start_date.strftime('%d/%m/%Y')} - {end_date.strftime('%d/%m/%Y')}"

    def _format_time_range(self, start_date, end_date) -> str:
        """Format time range for template"""
        if not start_date:
            return 'Time To Be Confirmed'
        
        if isinstance(start_date, str):
            try:
                start_date = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
            except:
                return 'Time To Be Confirmed'
        
        if not end_date:
            return start_date.strftime('%H:%M')
        
        if isinstance(end_date, str):
            try:
                end_date = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
            except:
                return start_date.strftime('%H:%M')
        
        return f"{start_date.strftime('%H:%M')} - {end_date.strftime('%H:%M')}"

    def _count_male_participants(self, participants_data: List[Dict]) -> str:
        """Count male participants from real data"""
        count = sum(1 for p in participants_data if p.get('gender', '').lower() in ['male', 'lelaki', 'm'])
        return str(count)

    def _count_female_participants(self, participants_data: List[Dict]) -> str:
        """Count female participants from real data"""
        count = sum(1 for p in participants_data if p.get('gender', '').lower() in ['female', 'perempuan', 'f'])
        return str(count)

    def _get_participant_attendance(self, participant_id: int, attendance_data: List[Dict], day_field: str) -> str:
        """Get participant attendance status for specific day"""
        for record in attendance_data:
            if record['participant_id'] == participant_id:
                status = record.get(day_field, 'absent')
                return 'Hadir' if status == 'present' else 'Tidak Hadir'
        return 'Tidak Hadir'

    def _count_attended(self, attendance_data: List[Dict]) -> str:
        """Count total attended from real attendance data"""
        count = sum(1 for record in attendance_data if record.get('day_1_status') == 'present')
        return str(count)

    def _count_absent(self, attendance_data: List[Dict]) -> str:
        """Count total absent from real attendance data"""
        total = len(attendance_data)
        attended = int(self._count_attended(attendance_data))
        return str(total - attended)

    def _calculate_attendance_percentage(self, attendance_data: List[Dict]) -> str:
        """Calculate attendance percentage from real data"""
        if not attendance_data:
            return "0%"
        
        total = len(attendance_data)
        attended = int(self._count_attended(attendance_data))
        percentage = (attended / total) * 100 if total > 0 else 0
        return f"{percentage:.1f}%"

    def _process_template(self, template_path: str, context: Dict, output_filename: str) -> str:
        """Process template with real data context"""
        try:
            # Load Word document
            doc = Document(template_path)
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_placeholders_in_paragraph(paragraph, context)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(paragraph, context)
            
            # Handle special sections (participants table, etc.)
            self._process_participants_table(doc, context)
            
            # Save processed document
            output_path = os.path.join(self.output_dir, output_filename)
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error processing template: {str(e)}")
            raise

    def _replace_placeholders_in_paragraph(self, paragraph, context: Dict):
        """Replace placeholders in paragraph with real data"""
        for run in paragraph.runs:
            text = run.text
            
            # Replace all placeholders
            for placeholder_pattern in re.findall(r'\{\{[^}]+\}\}', text):
                value = self._get_context_value(placeholder_pattern, context)
                text = text.replace(placeholder_pattern, str(value))
            
            run.text = text

    def _get_context_value(self, placeholder: str, context: Dict) -> str:
        """Get value from context for placeholder"""
        try:
            # Remove braces and split by dots
            key_path = placeholder.strip('{}').split('.')
            
            value = context
            for key in key_path:
                if isinstance(value, dict) and key in value:
                    value = value[key]
                else:
                    return placeholder  # Return original if not found
            
            return str(value) if value is not None else ""
            
        except Exception:
            return placeholder

    def _process_participants_table(self, doc: Document, context: Dict):
        """Process participants table with real data"""
        participants = context.get('participants', [])
        
        # Find table that contains participant headers
        for table in doc.tables:
            if len(table.rows) > 0:
                header_text = ""
                for cell in table.rows[0].cells:
                    header_text += cell.text.lower()
                
                # Check if this is participants table
                if any(keyword in header_text for keyword in ['bil', 'nama', 'participant', 'peserta']):
                    # Clear existing rows except header
                    while len(table.rows) > 1:
                        table._element.remove(table.rows[-1]._element)
                    
                    # Add participant rows
                    for participant in participants:
                        row = table.add_row()
                        cells = row.cells
                        
                        # Populate cells based on table structure
                        if len(cells) >= 3:
                            cells[0].text = str(participant.get('bil', ''))
                            cells[1].text = participant.get('name', '')
                            cells[2].text = participant.get('ic', '')
                        
                        if len(cells) >= 5:
                            cells[3].text = participant.get('attendance_day1', '')
                            cells[4].text = participant.get('attendance_day2', '')

    def validate_template_data(self, program_id: int) -> Dict[str, Any]:
        """Validate that all required data is available for template generation"""
        try:
            # Check program data
            program_data = self._get_program_data(program_id)
            if not program_data:
                return {
                    'valid': False,
                    'missing_data': ['program_data'],
                    'message': 'Program data not found'
                }
            
            # Check participants data
            participants_data = self._get_participants_data(program_id)
            
            # Check attendance data
            attendance_data = self._get_attendance_data(program_id)
            
            missing_fields = []
            warnings = []
            
            # Validate critical program fields
            required_program_fields = ['title', 'location', 'organizer']
            for field in required_program_fields:
                if not program_data.get(field):
                    missing_fields.append(f'program.{field}')
            
            # Check if we have participants
            if not participants_data:
                warnings.append('No participants registered for this program')
            
            # Check if we have attendance data
            if not attendance_data:
                warnings.append('No attendance records found for this program')
            
            return {
                'valid': len(missing_fields) == 0,
                'missing_data': missing_fields,
                'warnings': warnings,
                'data_summary': {
                    'program_fields': len([k for k, v in program_data.items() if v]),
                    'participants_count': len(participants_data),
                    'attendance_records': len(attendance_data)
                }
            }
            
        except Exception as e:
            logger.error(f"Error validating template data: {str(e)}")
            return {
                'valid': False,
                'error': str(e),
                'message': 'Validation failed due to system error'
            }
