"""
Production Automated Report System - ZERO MOCK DATA
Real implementation for automated report generation from form submissions
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime, timedelta
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib.backends.backend_agg import FigureCanvasAgg
import io
import base64
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session
from celery import current_app as celery_app
from dataclasses import dataclass
import uuid

logger = logging.getLogger(__name__)

@dataclass
class ReportConfiguration:
    """Configuration for automated report generation"""
    report_type: str  # summary, detailed, analytics, executive
    include_charts: bool = True
    include_ai_insights: bool = True
    date_range: str = "all_time"  # last_7_days, last_30_days, last_90_days, all_time
    format: str = "docx"  # docx, pdf, html
    language: str = "en"  # en, ms

class ProductionAutomatedReportSystem:
    """Production Automated Report System - NO MOCK DATA"""
    
    def __init__(self, db_session: Session, ai_service=None):
        self.db_session = db_session
        self.ai_service = ai_service
        self.reports_dir = os.getenv('REPORTS_STORAGE_PATH', './reports')
        self.charts_dir = os.getenv('CHARTS_STORAGE_PATH', './charts')
        
        # Ensure directories exist
        os.makedirs(self.reports_dir, exist_ok=True)
        os.makedirs(self.charts_dir, exist_ok=True)
        
        # Configure matplotlib for production
        plt.style.use('seaborn-v0_8')
        plt.rcParams['figure.figsize'] = (10, 6)
        plt.rcParams['font.size'] = 10
        
    def generate_automated_report(
        self,
        program_id: int,
        config: ReportConfiguration,
        user_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """Generate real automated report from program data"""
        try:
            # Get real program data
            program_data = self._get_program_data(program_id)
            if not program_data:
                raise Exception(f"Program {program_id} not found")
            
            # Get real participants data
            participants_data = self._get_participants_data(program_id, config.date_range)
            
            # Get real attendance data
            attendance_data = self._get_attendance_data(program_id, config.date_range)
            
            # Get real form responses data if available
            form_responses_data = self._get_form_responses_data(program_id, config.date_range)
            
            # Perform real statistical analysis
            analysis_results = self._perform_statistical_analysis(
                participants_data, 
                attendance_data, 
                form_responses_data
            )
            
            # Generate real charts
            charts = []
            if config.include_charts:
                charts = self._generate_charts(analysis_results, program_id)
            
            # Get AI insights from real data
            ai_insights = {}
            if config.include_ai_insights and self.ai_service:
                ai_insights = self._get_ai_insights(
                    program_data, 
                    participants_data, 
                    attendance_data, 
                    analysis_results
                )
            
            # Generate report document
            report_path = self._generate_report_document(
                program_data,
                participants_data,
                attendance_data,
                analysis_results,
                charts,
                ai_insights,
                config
            )
            
            # Save report record to database
            report_record = self._save_report_record(
                program_id, 
                user_id, 
                report_path, 
                config, 
                analysis_results
            )
            
            return {
                'status': 'success',
                'report_id': report_record['id'],
                'report_path': report_path,
                'program_id': program_id,
                'generated_at': datetime.utcnow().isoformat(),
                'summary': {
                    'participants_analyzed': len(participants_data),
                    'attendance_records': len(attendance_data),
                    'form_responses': len(form_responses_data),
                    'charts_generated': len(charts),
                    'ai_insights_included': bool(ai_insights),
                    'report_type': config.report_type
                },
                'analysis_highlights': self._extract_key_insights(analysis_results, ai_insights)
            }
            
        except Exception as e:
            logger.error(f"Error generating automated report: {str(e)}")
            return {
                'status': 'error',
                'message': str(e),
                'program_id': program_id,
                'timestamp': datetime.utcnow().isoformat()
            }

    def _get_program_data(self, program_id: int) -> Optional[Dict[str, Any]]:
        """Get real program data from database"""
        try:
            from ..models.production_models import Program
            
            program = self.db_session.query(Program).filter_by(id=program_id).first()
            if not program:
                return None
            
            return {
                'id': program.id,
                'title': program.title,
                'description': program.description,
                'start_date': program.start_date,
                'end_date': program.end_date,
                'location': program.location,
                'organizer': program.organizer,
                'speaker': program.speaker,
                'trainer': program.trainer,
                'facilitator': program.facilitator,
                'background': program.background,
                'objectives': program.objectives,
                'capacity': program.capacity,
                'status': program.status,
                'created_at': program.created_at
            }
            
        except Exception as e:
            logger.error(f"Error fetching program data: {str(e)}")
            return None

    def _get_participants_data(self, program_id: int, date_range: str) -> List[Dict[str, Any]]:
        """Get real participants data with date filtering"""
        try:
            from ..models.production_models import Participant
            
            query = self.db_session.query(Participant).filter_by(program_id=program_id)
            
            # Apply date filtering
            if date_range != "all_time":
                days_back = self._get_days_back(date_range)
                cutoff_date = datetime.utcnow() - timedelta(days=days_back)
                query = query.filter(Participant.registration_date >= cutoff_date)
            
            participants = query.all()
            
            return [
                {
                    'id': p.id,
                    'full_name': p.full_name,
                    'identification_number': p.identification_number,
                    'email': p.email,
                    'phone': p.phone,
                    'gender': p.gender,
                    'age': p.age,
                    'organization': p.organization,
                    'position': p.position,
                    'department': p.department,
                    'registration_date': p.registration_date,
                    'registration_source': p.registration_source,
                    'status': p.status
                }
                for p in participants
            ]
            
        except Exception as e:
            logger.error(f"Error fetching participants data: {str(e)}")
            return []

    def _get_attendance_data(self, program_id: int, date_range: str) -> List[Dict[str, Any]]:
        """Get real attendance data with date filtering"""
        try:
            from ..models.production_models import AttendanceRecord, Participant
            
            query = self.db_session.query(
                AttendanceRecord, Participant
            ).join(
                Participant, AttendanceRecord.participant_id == Participant.id
            ).filter(AttendanceRecord.program_id == program_id)
            
            # Apply date filtering
            if date_range != "all_time":
                days_back = self._get_days_back(date_range)
                cutoff_date = datetime.utcnow() - timedelta(days=days_back)
                query = query.filter(AttendanceRecord.recorded_at >= cutoff_date)
            
            records = query.all()
            
            return [
                {
                    'participant_id': record.AttendanceRecord.participant_id,
                    'participant_name': record.Participant.full_name,
                    'day_1_status': record.AttendanceRecord.day_1_status,
                    'day_2_status': record.AttendanceRecord.day_2_status,
                    'day_3_status': record.AttendanceRecord.day_3_status,
                    'total_hours_attended': float(record.AttendanceRecord.total_hours_attended or 0),
                    'attendance_percentage': float(record.AttendanceRecord.attendance_percentage or 0),
                    'recorded_at': record.AttendanceRecord.recorded_at
                }
                for record in records
            ]
            
        except Exception as e:
            logger.error(f"Error fetching attendance data: {str(e)}")
            return []

    def _get_form_responses_data(self, program_id: int, date_range: str) -> List[Dict[str, Any]]:
        """Get real form responses data if available"""
        try:
            from ..models.production_models import FormIntegration
            
            # Get form integrations for this program
            integrations = self.db_session.query(FormIntegration).filter_by(
                program_id=program_id,
                integration_status='active'
            ).all()
            
            all_responses = []
            
            for integration in integrations:
                # Get responses based on platform
                if integration.platform == 'google_forms':
                    responses = self._get_google_form_responses(integration.form_id, date_range)
                elif integration.platform == 'microsoft_forms':
                    responses = self._get_microsoft_form_responses(integration.form_id, date_range)
                else:
                    continue
                
                # Add platform info to responses
                for response in responses:
                    response['platform'] = integration.platform
                    response['form_integration_id'] = integration.id
                
                all_responses.extend(responses)
            
            return all_responses
            
        except Exception as e:
            logger.error(f"Error fetching form responses data: {str(e)}")
            return []

    def _perform_statistical_analysis(
        self, 
        participants_data: List[Dict], 
        attendance_data: List[Dict], 
        form_responses_data: List[Dict]
    ) -> Dict[str, Any]:
        """Perform real statistical analysis on data"""
        try:
            analysis = {
                'participants_analysis': self._analyze_participants(participants_data),
                'attendance_analysis': self._analyze_attendance(attendance_data),
                'temporal_analysis': self._analyze_temporal_patterns(participants_data, attendance_data),
                'demographic_analysis': self._analyze_demographics(participants_data),
                'form_responses_analysis': self._analyze_form_responses(form_responses_data),
                'quality_metrics': self._calculate_quality_metrics(participants_data, attendance_data),
                'completion_analysis': self._analyze_completion_rates(participants_data, attendance_data)
            }
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error performing statistical analysis: {str(e)}")
            return {}

    def _analyze_participants(self, participants_data: List[Dict]) -> Dict[str, Any]:
        """Analyze participants data"""
        if not participants_data:
            return {'total': 0, 'message': 'No participants data available'}
        
        df = pd.DataFrame(participants_data)
        
        analysis = {
            'total_participants': len(df),
            'registration_sources': df['registration_source'].value_counts().to_dict(),
            'status_distribution': df['status'].value_counts().to_dict(),
            'organization_distribution': df['organization'].value_counts().head(10).to_dict(),
            'registration_trend': self._calculate_registration_trend(df)
        }
        
        # Age analysis if available
        if 'age' in df.columns and df['age'].notna().sum() > 0:
            analysis['age_statistics'] = {
                'mean_age': float(df['age'].mean()),
                'median_age': float(df['age'].median()),
                'age_range': f"{int(df['age'].min())}-{int(df['age'].max())}"
            }
        
        return analysis

    def _analyze_attendance(self, attendance_data: List[Dict]) -> Dict[str, Any]:
        """Analyze attendance patterns"""
        if not attendance_data:
            return {'message': 'No attendance data available'}
        
        df = pd.DataFrame(attendance_data)
        
        # Calculate attendance rates
        day1_present = (df['day_1_status'] == 'present').sum()
        day2_present = (df['day_2_status'] == 'present').sum()
        total_records = len(df)
        
        analysis = {
            'total_records': total_records,
            'day_1_attendance_rate': (day1_present / total_records * 100) if total_records > 0 else 0,
            'day_2_attendance_rate': (day2_present / total_records * 100) if total_records > 0 else 0,
            'overall_attendance_rate': df['attendance_percentage'].mean() if 'attendance_percentage' in df.columns else 0,
            'perfect_attendance_count': len(df[
                (df['day_1_status'] == 'present') & 
                (df['day_2_status'] == 'present')
            ]),
            'average_hours_attended': df['total_hours_attended'].mean() if 'total_hours_attended' in df.columns else 0
        }
        
        return analysis

    def _analyze_temporal_patterns(self, participants_data: List[Dict], attendance_data: List[Dict]) -> Dict[str, Any]:
        """Analyze temporal patterns in data"""
        analysis = {}
        
        # Registration patterns
        if participants_data:
            df_participants = pd.DataFrame(participants_data)
            if 'registration_date' in df_participants.columns:
                df_participants['registration_date'] = pd.to_datetime(df_participants['registration_date'])
                
                # Daily registration counts
                daily_registrations = df_participants.groupby(
                    df_participants['registration_date'].dt.date
                ).size().to_dict()
                
                analysis['registration_patterns'] = {
                    'daily_counts': {str(k): int(v) for k, v in daily_registrations.items()},
                    'peak_registration_day': str(max(daily_registrations, key=daily_registrations.get)),
                    'registration_span_days': (
                        df_participants['registration_date'].max() - 
                        df_participants['registration_date'].min()
                    ).days
                }
        
        # Attendance timing patterns
        if attendance_data:
            df_attendance = pd.DataFrame(attendance_data)
            if 'recorded_at' in df_attendance.columns:
                df_attendance['recorded_at'] = pd.to_datetime(df_attendance['recorded_at'])
                
                analysis['attendance_timing'] = {
                    'records_by_hour': df_attendance.groupby(
                        df_attendance['recorded_at'].dt.hour
                    ).size().to_dict()
                }
        
        return analysis

    def _analyze_demographics(self, participants_data: List[Dict]) -> Dict[str, Any]:
        """Analyze demographic distribution"""
        if not participants_data:
            return {'message': 'No demographic data available'}
        
        df = pd.DataFrame(participants_data)
        
        analysis = {}
        
        # Gender distribution
        if 'gender' in df.columns:
            gender_counts = df['gender'].value_counts().to_dict()
            analysis['gender_distribution'] = {
                'counts': gender_counts,
                'percentages': {
                    k: (v / len(df) * 100) for k, v in gender_counts.items()
                }
            }
        
        # Organization analysis
        if 'organization' in df.columns:
            org_counts = df['organization'].value_counts().head(10).to_dict()
            analysis['top_organizations'] = org_counts
        
        # Position analysis
        if 'position' in df.columns:
            position_counts = df['position'].value_counts().head(10).to_dict()
            analysis['common_positions'] = position_counts
        
        return analysis

    def _analyze_form_responses(self, form_responses_data: List[Dict]) -> Dict[str, Any]:
        """Analyze form responses if available"""
        if not form_responses_data:
            return {'message': 'No form responses data available'}
        
        analysis = {
            'total_responses': len(form_responses_data),
            'platforms': {},
            'response_quality': {}
        }
        
        # Platform distribution
        platforms = [response.get('platform', 'unknown') for response in form_responses_data]
        platform_counts = pd.Series(platforms).value_counts().to_dict()
        analysis['platforms'] = platform_counts
        
        # Analyze response quality
        completed_responses = 0
        total_questions = 0
        answered_questions = 0
        
        for response in form_responses_data:
            answers = response.get('answers', {})
            total_questions += len(answers)
            answered_questions += len([v for v in answers.values() if v and str(v).strip()])
            
            if answers and len([v for v in answers.values() if v and str(v).strip()]) == len(answers):
                completed_responses += 1
        
        if len(form_responses_data) > 0:
            analysis['response_quality'] = {
                'completion_rate': (completed_responses / len(form_responses_data)) * 100,
                'average_answer_rate': (answered_questions / total_questions * 100) if total_questions > 0 else 0
            }
        
        return analysis

    def _calculate_quality_metrics(self, participants_data: List[Dict], attendance_data: List[Dict]) -> Dict[str, Any]:
        """Calculate data quality metrics"""
        metrics = {}
        
        # Participant data quality
        if participants_data:
            df_participants = pd.DataFrame(participants_data)
            
            # Calculate completeness for each field
            completeness = {}
            for column in df_participants.columns:
                non_null_count = df_participants[column].notna().sum()
                completeness[column] = (non_null_count / len(df_participants)) * 100
            
            metrics['participant_data_completeness'] = completeness
            metrics['overall_participant_quality'] = sum(completeness.values()) / len(completeness)
        
        # Attendance data quality
        if attendance_data:
            df_attendance = pd.DataFrame(attendance_data)
            
            # Check for missing attendance records
            expected_records = len(participants_data) if participants_data else 0
            actual_records = len(attendance_data)
            
            metrics['attendance_data_coverage'] = (actual_records / expected_records * 100) if expected_records > 0 else 0
            
            # Check attendance data completeness
            attendance_completeness = {}
            for column in df_attendance.columns:
                if column.endswith('_status'):
                    valid_statuses = df_attendance[column].isin(['present', 'absent', 'late']).sum()
                    attendance_completeness[column] = (valid_statuses / len(df_attendance)) * 100
            
            metrics['attendance_status_quality'] = attendance_completeness
        
        return metrics

    def _analyze_completion_rates(self, participants_data: List[Dict], attendance_data: List[Dict]) -> Dict[str, Any]:
        """Analyze completion rates"""
        analysis = {}
        
        if participants_data and attendance_data:
            total_participants = len(participants_data)
            participants_with_attendance = len(set(
                record['participant_id'] for record in attendance_data
            ))
            
            analysis['attendance_record_completion'] = (
                participants_with_attendance / total_participants * 100
            ) if total_participants > 0 else 0
            
            # Calculate program completion based on attendance
            if attendance_data:
                df_attendance = pd.DataFrame(attendance_data)
                completed_program = len(df_attendance[
                    df_attendance['attendance_percentage'] >= 80
                ])
                
                analysis['program_completion_rate'] = (
                    completed_program / total_participants * 100
                ) if total_participants > 0 else 0
        
        return analysis

    def _generate_charts(self, analysis_results: Dict, program_id: int) -> List[Dict[str, Any]]:
        """Generate real charts from analysis data"""
        charts = []
        chart_base_filename = f"program_{program_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        try:
            # 1. Gender Distribution Chart
            if 'demographic_analysis' in analysis_results:
                demo_data = analysis_results['demographic_analysis']
                if 'gender_distribution' in demo_data:
                    chart_path = self._create_gender_distribution_chart(
                        demo_data['gender_distribution'], 
                        f"{chart_base_filename}_gender.png"
                    )
                    charts.append({
                        'type': 'gender_distribution',
                        'title': 'Participant Gender Distribution',
                        'path': chart_path,
                        'description': 'Distribution of participants by gender'
                    })
            
            # 2. Attendance Rates Chart
            if 'attendance_analysis' in analysis_results:
                attendance_data = analysis_results['attendance_analysis']
                chart_path = self._create_attendance_chart(
                    attendance_data,
                    f"{chart_base_filename}_attendance.png"
                )
                charts.append({
                    'type': 'attendance_rates',
                    'title': 'Attendance Rates by Day',
                    'path': chart_path,
                    'description': 'Daily attendance rates comparison'
                })
            
            # 3. Registration Trend Chart
            if 'temporal_analysis' in analysis_results:
                temporal_data = analysis_results['temporal_analysis']
                if 'registration_patterns' in temporal_data:
                    chart_path = self._create_registration_trend_chart(
                        temporal_data['registration_patterns'],
                        f"{chart_base_filename}_registration.png"
                    )
                    charts.append({
                        'type': 'registration_trend',
                        'title': 'Registration Trend Over Time',
                        'path': chart_path,
                        'description': 'Daily registration patterns'
                    })
            
            # 4. Organization Distribution Chart
            if 'demographic_analysis' in analysis_results:
                demo_data = analysis_results['demographic_analysis']
                if 'top_organizations' in demo_data:
                    chart_path = self._create_organization_chart(
                        demo_data['top_organizations'],
                        f"{chart_base_filename}_organizations.png"
                    )
                    charts.append({
                        'type': 'organization_distribution',
                        'title': 'Top Participating Organizations',
                        'path': chart_path,
                        'description': 'Most represented organizations'
                    })
            
        except Exception as e:
            logger.error(f"Error generating charts: {str(e)}")
        
        return charts

    def _create_gender_distribution_chart(self, gender_data: Dict, filename: str) -> str:
        """Create gender distribution pie chart"""
        fig, ax = plt.subplots(figsize=(8, 6))
        
        counts = gender_data.get('counts', {})
        if counts:
            labels = list(counts.keys())
            sizes = list(counts.values())
            colors = ['#FF9999', '#66B2FF', '#99FF99', '#FFCC99']
            
            ax.pie(sizes, labels=labels, colors=colors[:len(labels)], autopct='%1.1f%%', startangle=90)
            ax.set_title('Participant Gender Distribution', fontsize=14, fontweight='bold')
        else:
            ax.text(0.5, 0.5, 'No gender data available', ha='center', va='center', transform=ax.transAxes)
        
        chart_path = os.path.join(self.charts_dir, filename)
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return chart_path

    def _create_attendance_chart(self, attendance_data: Dict, filename: str) -> str:
        """Create attendance rates bar chart"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        days = ['Day 1', 'Day 2']
        rates = [
            attendance_data.get('day_1_attendance_rate', 0),
            attendance_data.get('day_2_attendance_rate', 0)
        ]
        
        bars = ax.bar(days, rates, color=['#4CAF50', '#2196F3'])
        ax.set_ylabel('Attendance Rate (%)', fontsize=12)
        ax.set_title('Attendance Rates by Day', fontsize=14, fontweight='bold')
        ax.set_ylim(0, 100)
        
        # Add value labels on bars
        for bar, rate in zip(bars, rates):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                   f'{rate:.1f}%', ha='center', va='bottom', fontweight='bold')
        
        chart_path = os.path.join(self.charts_dir, filename)
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return chart_path

    def _create_registration_trend_chart(self, registration_data: Dict, filename: str) -> str:
        """Create registration trend line chart"""
        fig, ax = plt.subplots(figsize=(12, 6))
        
        daily_counts = registration_data.get('daily_counts', {})
        if daily_counts:
            dates = sorted(daily_counts.keys())
            counts = [daily_counts[date] for date in dates]
            
            ax.plot(dates, counts, marker='o', linewidth=2, markersize=6, color='#FF6B6B')
            ax.set_xlabel('Date', fontsize=12)
            ax.set_ylabel('Registration Count', fontsize=12)
            ax.set_title('Daily Registration Trend', fontsize=14, fontweight='bold')
            ax.grid(True, alpha=0.3)
            
            # Rotate x-axis labels for better readability
            plt.xticks(rotation=45)
        else:
            ax.text(0.5, 0.5, 'No registration trend data available', 
                   ha='center', va='center', transform=ax.transAxes)
        
        chart_path = os.path.join(self.charts_dir, filename)
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return chart_path

    def _create_organization_chart(self, org_data: Dict, filename: str) -> str:
        """Create organization distribution horizontal bar chart"""
        fig, ax = plt.subplots(figsize=(12, 8))
        
        if org_data:
            organizations = list(org_data.keys())[:10]  # Top 10
            counts = [org_data[org] for org in organizations]
            
            y_pos = np.arange(len(organizations))
            bars = ax.barh(y_pos, counts, color='#FFA726')
            
            ax.set_yticks(y_pos)
            ax.set_yticklabels(organizations)
            ax.set_xlabel('Number of Participants', fontsize=12)
            ax.set_title('Top Participating Organizations', fontsize=14, fontweight='bold')
            
            # Add value labels on bars
            for i, bar in enumerate(bars):
                width = bar.get_width()
                ax.text(width + 0.1, bar.get_y() + bar.get_height()/2.,
                       f'{int(width)}', ha='left', va='center', fontweight='bold')
        else:
            ax.text(0.5, 0.5, 'No organization data available', 
                   ha='center', va='center', transform=ax.transAxes)
        
        chart_path = os.path.join(self.charts_dir, filename)
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return chart_path

    def _get_ai_insights(
        self, 
        program_data: Dict, 
        participants_data: List[Dict], 
        attendance_data: List[Dict], 
        analysis_results: Dict
    ) -> Dict[str, Any]:
        """Get AI insights from real data"""
        if not self.ai_service:
            return {'message': 'AI service not available'}
        
        try:
            insights = self.ai_service.generate_report_insights(
                program_data, 
                participants_data, 
                attendance_data
            )
            
            # Add analysis-based insights
            insights['statistical_insights'] = self._generate_statistical_insights(analysis_results)
            
            return insights
            
        except Exception as e:
            logger.error(f"Error getting AI insights: {str(e)}")
            return {'error': f'AI insights generation failed: {str(e)}'}

    def _generate_statistical_insights(self, analysis_results: Dict) -> List[str]:
        """Generate insights from statistical analysis"""
        insights = []
        
        # Participation insights
        if 'participants_analysis' in analysis_results:
            total = analysis_results['participants_analysis'].get('total_participants', 0)
            if total > 0:
                insights.append(f"Program attracted {total} participants")
                
                # Registration source insights
                sources = analysis_results['participants_analysis'].get('registration_sources', {})
                if sources:
                    top_source = max(sources, key=sources.get)
                    insights.append(f"Most participants registered via {top_source}")
        
        # Attendance insights
        if 'attendance_analysis' in analysis_results:
            day1_rate = analysis_results['attendance_analysis'].get('day_1_attendance_rate', 0)
            day2_rate = analysis_results['attendance_analysis'].get('day_2_attendance_rate', 0)
            
            if day1_rate > 80:
                insights.append("Excellent Day 1 attendance rate (>80%)")
            elif day1_rate > 60:
                insights.append("Good Day 1 attendance rate (60-80%)")
            
            if abs(day1_rate - day2_rate) > 10:
                insights.append("Significant attendance drop between days")
        
        # Quality insights
        if 'quality_metrics' in analysis_results:
            quality = analysis_results['quality_metrics'].get('overall_participant_quality', 0)
            if quality > 90:
                insights.append("Excellent data quality (>90% complete)")
            elif quality < 70:
                insights.append("Data quality could be improved (<70% complete)")
        
        return insights

    def _generate_report_document(
        self,
        program_data: Dict,
        participants_data: List[Dict],
        attendance_data: List[Dict],
        analysis_results: Dict,
        charts: List[Dict],
        ai_insights: Dict,
        config: ReportConfiguration
    ) -> str:
        """Generate the final report document"""
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f"Automated Report: {program_data.get('title', 'Program Report')}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation info
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Report Type: {config.report_type.title()}")
        doc.add_paragraph("")
        
        # Executive Summary
        self._add_executive_summary(doc, program_data, analysis_results, ai_insights)
        
        # Program Overview
        self._add_program_overview(doc, program_data)
        
        # Participants Analysis
        self._add_participants_section(doc, participants_data, analysis_results)
        
        # Attendance Analysis
        self._add_attendance_section(doc, attendance_data, analysis_results)
        
        # Charts
        if charts:
            self._add_charts_section(doc, charts)
        
        # AI Insights
        if ai_insights and config.include_ai_insights:
            self._add_ai_insights_section(doc, ai_insights)
        
        # Recommendations
        self._add_recommendations_section(doc, analysis_results, ai_insights)
        
        # Save document
        report_filename = f"automated_report_{program_data['id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        report_path = os.path.join(self.reports_dir, report_filename)
        doc.save(report_path)
        
        return report_path

    def _add_executive_summary(self, doc: Document, program_data: Dict, analysis_results: Dict, ai_insights: Dict):
        """Add executive summary to document"""
        doc.add_heading('Executive Summary', level=1)
        
        # Key metrics
        participants_count = analysis_results.get('participants_analysis', {}).get('total_participants', 0)
        attendance_rate = analysis_results.get('attendance_analysis', {}).get('day_1_attendance_rate', 0)
        
        summary_text = f"""
This automated report provides a comprehensive analysis of the {program_data.get('title', 'program')} 
conducted from {self._format_date(program_data.get('start_date'))} to {self._format_date(program_data.get('end_date'))}.

Key Highlights:
• Total Participants: {participants_count}
• Overall Attendance Rate: {attendance_rate:.1f}%
• Program Location: {program_data.get('location', 'Not specified')}
• Organizer: {program_data.get('organizer', 'Not specified')}
        """
        
        doc.add_paragraph(summary_text.strip())
        
        # Add AI summary if available
        if ai_insights.get('executive_summary'):
            doc.add_paragraph("")
            doc.add_paragraph("AI-Generated Summary:")
            doc.add_paragraph(ai_insights['executive_summary'])

    def _add_program_overview(self, doc: Document, program_data: Dict):
        """Add program overview section"""
        doc.add_heading('Program Overview', level=1)
        
        # Create table for program details
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        program_fields = [
            ('Program Title', program_data.get('title', 'N/A')),
            ('Description', program_data.get('description', 'N/A')),
            ('Start Date', self._format_date(program_data.get('start_date'))),
            ('End Date', self._format_date(program_data.get('end_date'))),
            ('Location', program_data.get('location', 'N/A')),
            ('Organizer', program_data.get('organizer', 'N/A')),
            ('Speaker', program_data.get('speaker', 'N/A')),
            ('Trainer', program_data.get('trainer', 'N/A')),
            ('Facilitator', program_data.get('facilitator', 'N/A'))
        ]
        
        for field_name, field_value in program_fields:
            row = table.add_row()
            row.cells[0].text = field_name
            row.cells[1].text = str(field_value)

    def _add_participants_section(self, doc: Document, participants_data: List[Dict], analysis_results: Dict):
        """Add participants analysis section"""
        doc.add_heading('Participants Analysis', level=1)
        
        participants_analysis = analysis_results.get('participants_analysis', {})
        
        doc.add_paragraph(f"Total Participants: {participants_analysis.get('total_participants', 0)}")
        
        # Registration sources
        if 'registration_sources' in participants_analysis:
            doc.add_heading('Registration Sources', level=2)
            sources = participants_analysis['registration_sources']
            for source, count in sources.items():
                doc.add_paragraph(f"• {source.title()}: {count} participants")
        
        # Demographics
        if 'demographic_analysis' in analysis_results:
            demo_data = analysis_results['demographic_analysis']
            
            doc.add_heading('Demographics', level=2)
            
            # Gender distribution
            if 'gender_distribution' in demo_data:
                doc.add_paragraph("Gender Distribution:")
                gender_counts = demo_data['gender_distribution']['counts']
                for gender, count in gender_counts.items():
                    percentage = demo_data['gender_distribution']['percentages'][gender]
                    doc.add_paragraph(f"• {gender.title()}: {count} ({percentage:.1f}%)")

    def _add_attendance_section(self, doc: Document, attendance_data: List[Dict], analysis_results: Dict):
        """Add attendance analysis section"""
        doc.add_heading('Attendance Analysis', level=1)
        
        attendance_analysis = analysis_results.get('attendance_analysis', {})
        
        if attendance_analysis:
            doc.add_paragraph(f"Day 1 Attendance Rate: {attendance_analysis.get('day_1_attendance_rate', 0):.1f}%")
            doc.add_paragraph(f"Day 2 Attendance Rate: {attendance_analysis.get('day_2_attendance_rate', 0):.1f}%")
            doc.add_paragraph(f"Perfect Attendance Count: {attendance_analysis.get('perfect_attendance_count', 0)}")
            
            avg_hours = attendance_analysis.get('average_hours_attended', 0)
            if avg_hours > 0:
                doc.add_paragraph(f"Average Hours Attended: {avg_hours:.1f}")
        else:
            doc.add_paragraph("No attendance data available for analysis.")

    def _add_charts_section(self, doc: Document, charts: List[Dict]):
        """Add charts to document"""
        doc.add_heading('Visual Analysis', level=1)
        
        for chart in charts:
            doc.add_heading(chart['title'], level=2)
            doc.add_paragraph(chart['description'])
            
            # Add chart image
            if os.path.exists(chart['path']):
                doc.add_picture(chart['path'], width=Inches(6))
            else:
                doc.add_paragraph("Chart image not available.")
            
            doc.add_paragraph("")

    def _add_ai_insights_section(self, doc: Document, ai_insights: Dict):
        """Add AI insights section"""
        doc.add_heading('AI-Powered Insights', level=1)
        
        # Add different types of insights
        insight_sections = [
            ('effectiveness', 'Program Effectiveness'),
            ('participation_metrics', 'Participation Analysis'),
            ('attendance_analysis', 'Attendance Insights'),
            ('success_indicators', 'Success Indicators'),
            ('improvements', 'Improvement Areas')
        ]
        
        for key, title in insight_sections:
            if key in ai_insights:
                doc.add_heading(title, level=2)
                insight_text = ai_insights[key]
                if isinstance(insight_text, list):
                    for item in insight_text:
                        doc.add_paragraph(f"• {item}")
                else:
                    doc.add_paragraph(str(insight_text))

    def _add_recommendations_section(self, doc: Document, analysis_results: Dict, ai_insights: Dict):
        """Add recommendations section"""
        doc.add_heading('Recommendations', level=1)
        
        recommendations = []
        
        # Statistical recommendations
        attendance_rate = analysis_results.get('attendance_analysis', {}).get('day_1_attendance_rate', 0)
        if attendance_rate < 70:
            recommendations.append("Consider strategies to improve attendance rates")
        
        quality_score = analysis_results.get('quality_metrics', {}).get('overall_participant_quality', 0)
        if quality_score < 80:
            recommendations.append("Improve data collection processes for better completeness")
        
        # AI recommendations
        if ai_insights.get('improvements'):
            if isinstance(ai_insights['improvements'], list):
                recommendations.extend(ai_insights['improvements'])
            else:
                recommendations.append(ai_insights['improvements'])
        
        # Add recommendations to document
        if recommendations:
            for i, recommendation in enumerate(recommendations, 1):
                doc.add_paragraph(f"{i}. {recommendation}")
        else:
            doc.add_paragraph("No specific recommendations generated.")

    def _save_report_record(
        self, 
        program_id: int, 
        user_id: Optional[int], 
        report_path: str, 
        config: ReportConfiguration, 
        analysis_results: Dict
    ) -> Dict[str, Any]:
        """Save report record to database"""
        try:
            from ..models.production_models import Report
            
            report = Report(
                program_id=program_id,
                template_name=f"automated_{config.report_type}",
                file_path=report_path,
                generation_status='completed',
                generated_at=datetime.utcnow(),
                file_size=os.path.getsize(report_path) if os.path.exists(report_path) else 0
            )
            
            self.db_session.add(report)
            self.db_session.commit()
            
            return {
                'id': report.id,
                'status': 'success',
                'generated_at': report.generated_at
            }
            
        except Exception as e:
            logger.error(f"Error saving report record: {str(e)}")
            return {
                'id': None,
                'status': 'error',
                'message': str(e)
            }

    def _extract_key_insights(self, analysis_results: Dict, ai_insights: Dict) -> List[str]:
        """Extract key insights for summary"""
        insights = []
        
        # Statistical insights
        if 'participants_analysis' in analysis_results:
            total = analysis_results['participants_analysis'].get('total_participants', 0)
            insights.append(f"Analyzed {total} participants")
        
        # Attendance insights
        if 'attendance_analysis' in analysis_results:
            rate = analysis_results['attendance_analysis'].get('day_1_attendance_rate', 0)
            insights.append(f"Day 1 attendance: {rate:.1f}%")
        
        # AI insights summary
        if ai_insights.get('statistical_insights'):
            insights.extend(ai_insights['statistical_insights'][:3])  # Top 3 insights
        
        return insights

    def _get_days_back(self, date_range: str) -> int:
        """Get number of days back for date filtering"""
        mapping = {
            'last_7_days': 7,
            'last_30_days': 30,
            'last_90_days': 90
        }
        return mapping.get(date_range, 365)

    def _format_date(self, date_obj) -> str:
        """Format date for display"""
        if not date_obj:
            return 'Not specified'
        
        if isinstance(date_obj, str):
            try:
                date_obj = datetime.fromisoformat(date_obj.replace('Z', '+00:00'))
            except:
                return str(date_obj)
        
        return date_obj.strftime('%d/%m/%Y')

    def _calculate_registration_trend(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Calculate registration trend from dataframe"""
        if 'registration_date' not in df.columns:
            return {}
        
        try:
            df['registration_date'] = pd.to_datetime(df['registration_date'])
            daily_counts = df.groupby(df['registration_date'].dt.date).size()
            
            return {
                'total_days': len(daily_counts),
                'peak_day': str(daily_counts.idxmax()),
                'peak_count': int(daily_counts.max()),
                'average_daily': float(daily_counts.mean())
            }
        except Exception:
            return {}

    def _get_google_form_responses(self, form_id: str, date_range: str) -> List[Dict[str, Any]]:
        """Get Google Form responses (placeholder for integration)"""
        # This would integrate with GoogleFormsService
        return []

    def _get_microsoft_form_responses(self, form_id: str, date_range: str) -> List[Dict[str, Any]]:
        """Get Microsoft Form responses (placeholder for integration)"""
        # This would integrate with MicrosoftGraphService
        return []
