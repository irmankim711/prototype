"""
Enhanced Report Service for generating comprehensive reports with charts and analysis
"""
import io
import base64
from io import BytesIO
from flask import current_app
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
# import matplotlib.pyplot as plt
# import seaborn as sns
import json
import os
import logging
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any
import uuid
from pathlib import Path
import tempfile

# Import Flask models with fallback for development  
try:
    from ..models import Report, ReportTemplate, User, FormSubmission, Form, db
    HAS_MODELS = True
except ImportError:
    # Fallback for testing or development mode
    Report = ReportTemplate = User = FormSubmission = Form = db = None
    HAS_MODELS = False

logger = logging.getLogger(__name__)

class EnhancedReportService:
    """Enhanced service for generating comprehensive reports with visualizations"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.custom_styles = self._create_custom_styles()
        # plt.style.use('seaborn-v0_8')  # Disabled due to matplotlib issue
        
        # Setup reports directory
        current_dir = Path(__file__).parent.parent.parent
        self.reports_dir = current_dir / 'reports'
        self.reports_dir.mkdir(exist_ok=True)
        
    def _create_custom_styles(self):
        """Create custom paragraph styles"""
        custom_styles = {}
        
        # Title style
        custom_styles['CustomTitle'] = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.darkblue,
            alignment=1  # Center alignment
        )
        
        # Heading style
        custom_styles['CustomHeading'] = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.darkblue,
            spaceBefore=20
        )
        
        # Body style
        custom_styles['CustomBody'] = ParagraphStyle(
            'CustomBody',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            textColor=colors.black
        )
        
        return custom_styles
    
    def get_templates(self) -> List[Dict]:
        """Get all active report templates"""
        if HAS_MODELS and ReportTemplate:
            try:
                templates = ReportTemplate.query.filter_by(is_active=True).all()
                return [
                    {
                        'id': str(t.id),
                        'name': t.name,
                        'description': t.description,
                        'schema': t.schema,
                        'isActive': t.is_active,
                        'createdAt': t.created_at.isoformat() if t.created_at else None
                    }
                    for t in templates
                ]
            except Exception as e:
                logger.error(f"Error fetching templates: {e}")
        
        # Return mock templates
        return [
            {
                'id': '1',
                'name': 'Form Submission Summary',
                'description': 'Summary report of form submissions',
                'schema': {'format': 'pdf'},
                'isActive': True,
                'createdAt': datetime.now().isoformat()
            },
            {
                'id': '2', 
                'name': 'Detailed Analytics Report',
                'description': 'Comprehensive analytics and insights',
                'schema': {'format': 'pdf'},
                'isActive': True,
                'createdAt': datetime.now().isoformat()
            },
            {
                'id': '3',
                'name': 'Executive Summary Report',
                'description': 'High-level executive summary',
                'schema': {'format': 'pdf'},
                'isActive': True,
                'createdAt': datetime.now().isoformat()
            }
        ]
    
    def generate_report(self, template_id: str, data: Dict, user_id: int) -> Dict:
        """Generate a report based on template and data"""
        try:
            # Get template info
            templates = self.get_templates()
            template = next((t for t in templates if t['id'] == template_id), None)
            
            if not template:
                raise ValueError(f"Template {template_id} not found")
            
            # Generate unique filename
            report_id = str(uuid.uuid4())
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            # Determine output format from template schema
            output_format = template['schema'].get('format', 'pdf')
            filename = f"report_{template_id}_{timestamp}_{report_id}.{output_format}"
            output_path = self.reports_dir / filename
            
            # Generate report content based on template type
            if template['name'] == 'Form Submission Summary':
                pdf_content = self._generate_form_submission_report(data)
            elif template['name'] == 'Detailed Analytics Report':
                pdf_content = self._generate_analytics_report(data)
            elif template['name'] == 'Executive Summary Report':
                pdf_content = self._generate_executive_report(data)
            else:
                pdf_content = self._generate_custom_report(data)
            
            # Save the report
            with open(output_path, 'wb') as f:
                f.write(pdf_content)
            
            return {
                'status': 'success',
                'output_path': str(output_path),
                'output_url': f'/api/reports/download/{filename}',
                'filename': filename,
                'format': output_format
            }
            
        except Exception as e:
            logger.error(f"Error generating report: {e}")
            raise

    def _generate_form_submission_report(self, data: Dict) -> bytes:
        """Generate form submission summary report"""
        report_data = {
            'title': 'Form Submission Summary Report',
            'metadata': data.get('metadata', {}),
            'summary': data.get('summary', {}),
            'charts': data.get('charts', []),
            'analysis': data.get('analysis', {}),
            'raw_data': data.get('responses', [])
        }
        return self.generate_enhanced_report(report_data, 'form_summary')
    
    def _generate_analytics_report(self, data: Dict) -> bytes:
        """Generate detailed analytics report"""
        report_data = {
            'title': 'Detailed Analytics Report',
            'metadata': data.get('metadata', {}),
            'summary': data.get('summary', {}),
            'charts': data.get('charts', []),
            'analysis': data.get('analysis', {}),
            'raw_data': data.get('responses', [])
        }
        return self.generate_enhanced_report(report_data, 'analytics')
    
    def _generate_executive_report(self, data: Dict) -> bytes:
        """Generate executive summary report"""
        report_data = {
            'title': 'Executive Summary Report',
            'metadata': data.get('metadata', {}),
            'summary': data.get('summary', {}),
            'charts': data.get('charts', []),
            'analysis': data.get('analysis', {}),
            'raw_data': []  # No raw data in executive summary
        }
        return self.generate_enhanced_report(report_data, 'executive')
    
    def _generate_custom_report(self, data: Dict) -> bytes:
        """Generate custom report"""
        report_data = {
            'title': 'Custom Report',
            'metadata': data.get('metadata', {}),
            'summary': data.get('summary', {}),
            'charts': data.get('charts', []),
            'analysis': data.get('analysis', {}),
            'raw_data': data.get('responses', [])
        }
        return self.generate_enhanced_report(report_data, 'custom')
        
    def generate_enhanced_report(self, data: Dict[str, Any], report_type: str = 'general') -> bytes:
        """Generate an enhanced PDF report with charts and analysis"""
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            
            # Add title
            title = data.get('title', 'Automated Report')
            story.append(Paragraph(title, self.custom_styles['CustomTitle']))
            story.append(Spacer(1, 12))
            
            # Add metadata
            metadata = data.get('metadata', {})
            if metadata:
                story.append(Paragraph("Report Information", self.custom_styles['CustomHeading']))
                
                info_data = [
                    ['Generated', metadata.get('generated_at', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))],
                    ['Report Type', report_type.title()],
                    ['Data Source', metadata.get('source', 'Google Forms')],
                    ['Total Responses', str(metadata.get('total_responses', 0))]
                ]
                
                info_table = Table(info_data, colWidths=[2*inch, 4*inch])
                info_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(info_table)
                story.append(Spacer(1, 20))
            
            # Add summary section
            summary = data.get('summary', {})
            if summary:
                story.append(Paragraph("Executive Summary", self.custom_styles['CustomHeading']))
                
                for key, value in summary.items():
                    if isinstance(value, (int, float)):
                        value = f"{value:,.2f}" if isinstance(value, float) else f"{value:,}"
                    story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", 
                                         self.custom_styles['CustomBody']))
                
                story.append(Spacer(1, 20))
            
            # Add charts if data is available
            charts_data = data.get('charts', [])
            if charts_data:
                story.append(Paragraph("Data Visualizations", self.custom_styles['CustomHeading']))
                
                for chart_data in charts_data:
                    chart_image = self._create_chart(chart_data)
                    if chart_image:
                        story.append(chart_image)
                        story.append(Spacer(1, 12))
            
            # Add detailed analysis
            analysis = data.get('analysis', {})
            if analysis:
                story.append(PageBreak())
                story.append(Paragraph("Detailed Analysis", self.custom_styles['CustomHeading']))
                
                for section, content in analysis.items():
                    story.append(Paragraph(section.replace('_', ' ').title(), 
                                         self.custom_styles['CustomHeading']))
                    
                    if isinstance(content, list):
                        for item in content:
                            story.append(Paragraph(f"• {item}", self.custom_styles['CustomBody']))
                    else:
                        story.append(Paragraph(str(content), self.custom_styles['CustomBody']))
                    
                    story.append(Spacer(1, 12))
            
            # Add raw data table if available
            raw_data = data.get('raw_data', [])
            if raw_data and len(raw_data) > 0:
                story.append(PageBreak())
                story.append(Paragraph("Response Data", self.custom_styles['CustomHeading']))
                
                # Create table from raw data
                if isinstance(raw_data[0], dict):
                    headers = list(raw_data[0].keys())
                    table_data = [headers]
                    
                    for row in raw_data[:20]:  # Limit to first 20 rows
                        table_data.append([str(row.get(header, '')) for header in headers])
                    
                    data_table = Table(table_data)
                    data_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 8)
                    ]))
                    
                    story.append(data_table)
            
            # Build PDF
            doc.build(story)
            
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating enhanced report: {e}")
            raise
    
    def _create_chart(self, chart_data: Dict[str, Any]) -> Optional[Image]:
        """Create a chart from data and return as ReportLab Image"""
        try:
            # Temporarily disabled due to matplotlib issues
            logger.warning("Chart generation temporarily disabled due to matplotlib issues")
            return None
            
            # Original chart code commented out:
            # chart_type = chart_data.get('type', 'bar')
            # data = chart_data.get('data', {})
            # title = chart_data.get('title', 'Chart')
            
            # Create figure
            # plt.figure(figsize=(8, 6))
            
            # if chart_type == 'bar':
            #     plt.bar(data.keys(), data.values())
            #     plt.xticks(rotation=45)
            # elif chart_type == 'pie':
            #     plt.pie(data.values(), labels=data.keys(), autopct='%1.1f%%')
            # elif chart_type == 'line':
            #     plt.plot(list(data.keys()), list(data.values()), marker='o')
            #     plt.xticks(rotation=45)
            # else:
            #     # Default to bar chart
            #     plt.bar(data.keys(), data.values())
            #     plt.xticks(rotation=45)
            
            # plt.title(title)
            # plt.tight_layout()
            
            # Save to bytes
            # img_buffer = BytesIO()
            # plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            # img_buffer.seek(0)
            # plt.close()
            
            # Create ReportLab Image
            # image = Image(img_buffer, width=6*inch, height=4*inch)
            # return image
            
        except Exception as e:
            logger.error(f"Error creating chart: {e}")
            # plt.close()  # Ensure plot is closed even on error
            return None
    
    def generate_word_report(self, data: Dict[str, Any]) -> bytes:
        """Generate a Word document report"""
        try:
            doc = Document()
            
            # Add title
            title = data.get('title', 'Automated Report')
            doc.add_heading(title, 0)
            
            # Add metadata
            metadata = data.get('metadata', {})
            if metadata:
                doc.add_heading('Report Information', level=1)
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Property'
                hdr_cells[1].text = 'Value'
                
                for key, value in metadata.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = key.replace('_', ' ').title()
                    row_cells[1].text = str(value)
            
            # Add summary
            summary = data.get('summary', {})
            if summary:
                doc.add_heading('Summary', level=1)
                
                for key, value in summary.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                    p.add_run(str(value))
            
            # Add analysis
            analysis = data.get('analysis', {})
            if analysis:
                doc.add_heading('Analysis', level=1)
                
                for section, content in analysis.items():
                    doc.add_heading(section.replace('_', ' ').title(), level=2)
                    
                    if isinstance(content, list):
                        for item in content:
                            doc.add_paragraph(f"• {item}")
                    else:
                        doc.add_paragraph(str(content))
            
            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating Word report: {e}")
            raise
    
    def create_comparison_report(self, datasets: List[Dict[str, Any]], 
                               comparison_type: str = 'side_by_side') -> bytes:
        """Create a comparison report between multiple datasets"""
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            
            # Title
            story.append(Paragraph("Comparison Report", self.custom_styles['CustomTitle']))
            story.append(Spacer(1, 20))
            
            # Comparison summary
            story.append(Paragraph("Comparison Summary", self.custom_styles['CustomHeading']))
            
            summary_data = [['Dataset', 'Records', 'Date Range']]
            for i, dataset in enumerate(datasets):
                metadata = dataset.get('metadata', {})
                summary_data.append([
                    f"Dataset {i+1}",
                    str(metadata.get('total_responses', 0)),
                    metadata.get('date_range', 'N/A')
                ])
            
            summary_table = Table(summary_data)
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(summary_table)
            story.append(Spacer(1, 20))
            
            # Individual dataset analysis
            for i, dataset in enumerate(datasets):
                story.append(PageBreak())
                story.append(Paragraph(f"Dataset {i+1} Analysis", self.custom_styles['CustomHeading']))
                
                # Add dataset-specific content
                summary = dataset.get('summary', {})
                for key, value in summary.items():
                    story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", 
                                         self.custom_styles['CustomBody']))
                
                # Add charts for this dataset
                charts = dataset.get('charts', [])
                for chart in charts:
                    chart_image = self._create_chart(chart)
                    if chart_image:
                        story.append(chart_image)
                        story.append(Spacer(1, 12))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating comparison report: {e}")
            raise

# Create global instance
enhanced_report_service = EnhancedReportService()
