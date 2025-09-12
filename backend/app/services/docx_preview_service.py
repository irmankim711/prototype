"""
DOCX Preview Service
Converts DOCX files to HTML for in-browser preview
"""

import os
import logging
import tempfile
from pathlib import Path
from typing import Tuple, Optional, Dict, Any
from docx import Document
from docx.shared import Inches
import base64
from io import BytesIO
import zipfile
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)

class DocxPreviewService:
    """Service for converting DOCX files to HTML for browser preview"""
    
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or os.path.join(os.getcwd(), 'static', 'previews')
        self.ensure_directories()
    
    def ensure_directories(self):
        """Ensure required directories exist"""
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
    
    def convert_docx_to_html(self, docx_path: str, output_filename: str = None) -> Tuple[str, str]:
        """
        Convert DOCX to HTML for browser preview
        
        Args:
            docx_path: Path to DOCX file
            output_filename: Optional output filename
            
        Returns:
            Tuple of (html_file_path, html_content)
        """
        try:
            if not os.path.exists(docx_path):
                raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            
            # Generate output filename if not provided
            if not output_filename:
                base_name = Path(docx_path).stem
                output_filename = f"{base_name}_preview.html"
            
            html_file_path = os.path.join(self.output_dir, output_filename)
            
            # Open and parse the DOCX document
            doc = Document(docx_path)
            
            # Extract images from the DOCX file
            images = self._extract_images_from_docx(docx_path)
            
            # Convert to HTML
            html_content = self._convert_document_to_html(doc, images)
            
            # Save HTML file
            with open(html_file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"Successfully converted DOCX to HTML: {html_file_path}")
            return html_file_path, html_content
            
        except Exception as e:
            logger.error(f"DOCX to HTML conversion failed: {str(e)}")
            raise
    
    def _extract_images_from_docx(self, docx_path: str) -> Dict[str, str]:
        """Extract images from DOCX file and convert to base64"""
        images = {}
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # Find all image files in the DOCX
                image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                for image_file in image_files:
                    # Read image data
                    image_data = docx_zip.read(image_file)
                    
                    # Determine image format
                    image_ext = Path(image_file).suffix.lower()
                    if image_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                        # Convert to base64
                        image_b64 = base64.b64encode(image_data).decode('utf-8')
                        mime_type = f"image/{image_ext[1:].replace('jpg', 'jpeg')}"
                        
                        # Store with filename as key
                        image_name = Path(image_file).name
                        images[image_name] = f"data:{mime_type};base64,{image_b64}"
        
        except Exception as e:
            logger.warning(f"Failed to extract images from DOCX: {e}")
        
        return images
    
    def _convert_document_to_html(self, doc: Document, images: Dict[str, str]) -> str:
        """Convert python-docx Document to HTML"""
        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="en">',
            '<head>',
            '<meta charset="UTF-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
            '<title>Document Preview</title>',
            '<style>',
            self._get_preview_styles(),
            '</style>',
            '</head>',
            '<body>',
            '<div class="docx-preview-container">',
        ]
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            html_parts.append(self._convert_paragraph_to_html(paragraph))
        
        # Process tables
        for table in doc.tables:
            html_parts.append(self._convert_table_to_html(table))
        
        html_parts.extend([
            '</div>',
            '</body>',
            '</html>'
        ])
        
        return '\n'.join(html_parts)
    
    def _convert_paragraph_to_html(self, paragraph) -> str:
        """Convert a paragraph to HTML"""
        if not paragraph.text.strip():
            return '<br>'
        
        # Determine paragraph style
        style_class = 'paragraph'
        if paragraph.style.name.startswith('Heading'):
            level = paragraph.style.name.replace('Heading ', '')
            if level.isdigit() and 1 <= int(level) <= 6:
                return f'<h{level} class="heading-{level}">{self._escape_html(paragraph.text)}</h{level}>'
        
        # Handle different paragraph styles
        if 'Title' in paragraph.style.name:
            style_class = 'title'
        elif 'Subtitle' in paragraph.style.name:
            style_class = 'subtitle'
        
        # Process runs for formatting
        html_content = ''
        for run in paragraph.runs:
            text = self._escape_html(run.text)
            
            # Apply formatting
            if run.bold:
                text = f'<strong>{text}</strong>'
            if run.italic:
                text = f'<em>{text}</em>'
            if run.underline:
                text = f'<u>{text}</u>'
            
            html_content += text
        
        return f'<p class="{style_class}">{html_content}</p>'
    
    def _convert_table_to_html(self, table) -> str:
        """Convert a table to HTML"""
        html_parts = ['<table class="docx-table">']
        
        for i, row in enumerate(table.rows):
            html_parts.append('<tr>')
            
            for cell in row.cells:
                tag = 'th' if i == 0 else 'td'  # First row as header
                cell_text = self._escape_html(cell.text)
                html_parts.append(f'<{tag} class="table-cell">{cell_text}</{tag}>')
            
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        return '\n'.join(html_parts)
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&#x27;'))
    
    def _get_preview_styles(self) -> str:
        """Get CSS styles for the preview"""
        return """
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                color: #333;
                background-color: #f5f5f5;
                margin: 0;
                padding: 20px;
            }
            
            .docx-preview-container {
                max-width: 800px;
                margin: 0 auto;
                background: white;
                padding: 40px;
                box-shadow: 0 0 20px rgba(0,0,0,0.1);
                border-radius: 8px;
            }
            
            .title {
                font-size: 2em;
                font-weight: bold;
                text-align: center;
                margin-bottom: 1em;
                color: #2c3e50;
                border-bottom: 3px solid #3498db;
                padding-bottom: 10px;
            }
            
            .subtitle {
                font-size: 1.3em;
                font-weight: 600;
                text-align: center;
                margin-bottom: 1.5em;
                color: #7f8c8d;
            }
            
            .heading-1 {
                font-size: 1.8em;
                color: #2c3e50;
                margin-top: 1.5em;
                margin-bottom: 0.8em;
                border-bottom: 2px solid #ecf0f1;
                padding-bottom: 5px;
            }
            
            .heading-2 {
                font-size: 1.5em;
                color: #34495e;
                margin-top: 1.3em;
                margin-bottom: 0.7em;
            }
            
            .heading-3 {
                font-size: 1.3em;
                color: #34495e;
                margin-top: 1.1em;
                margin-bottom: 0.6em;
            }
            
            .paragraph {
                margin-bottom: 1em;
                text-align: justify;
            }
            
            .docx-table {
                width: 100%;
                border-collapse: collapse;
                margin: 1.5em 0;
                background: white;
            }
            
            .table-cell {
                border: 1px solid #ddd;
                padding: 12px;
                text-align: left;
            }
            
            .docx-table th {
                background-color: #f8f9fa;
                font-weight: 600;
                color: #2c3e50;
            }
            
            .docx-table tr:nth-child(even) {
                background-color: #f8f9fa;
            }
            
            strong {
                color: #2c3e50;
            }
            
            em {
                color: #7f8c8d;
            }
            
            @media print {
                body {
                    background: white;
                    padding: 0;
                }
                
                .docx-preview-container {
                    box-shadow: none;
                    padding: 20px;
                }
            }
            
            @media (max-width: 768px) {
                .docx-preview-container {
                    padding: 20px;
                    margin: 10px;
                }
                
                .title {
                    font-size: 1.5em;
                }
                
                .docx-table {
                    font-size: 0.9em;
                }
                
                .table-cell {
                    padding: 8px;
                }
            }
        """
    
    def get_preview_url(self, docx_path: str) -> str:
        """
        Get URL for previewing a DOCX file
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            URL for preview
        """
        try:
            # Convert to HTML
            html_file_path, _ = self.convert_docx_to_html(docx_path)
            
            # Return relative URL
            filename = Path(html_file_path).name
            return f'/static/previews/{filename}'
            
        except Exception as e:
            logger.error(f"Failed to get preview URL: {str(e)}")
            raise
    
    def cleanup_old_previews(self, max_age_hours: int = 24):
        """Clean up old preview files"""
        try:
            import time
            current_time = time.time()
            max_age_seconds = max_age_hours * 3600
            
            for file_path in Path(self.output_dir).glob('*.html'):
                if current_time - file_path.stat().st_mtime > max_age_seconds:
                    file_path.unlink()
                    logger.info(f"Cleaned up old preview: {file_path}")
                    
        except Exception as e:
            logger.error(f"Failed to cleanup previews: {str(e)}")

# Global service instance
docx_preview_service = DocxPreviewService()