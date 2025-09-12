import openai
from openai import OpenAI
from flask import Blueprint, jsonify, request
from flask_jwt_extended import jwt_required, get_jwt_identity
from app import db
from app.models import Report
from sqlalchemy import desc
import os
import json
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import io
import base64
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from io import BytesIO

ai_reports_bp = Blueprint('ai_reports', __name__)

# Initialize OpenAI client conditionally
openai_api_key = os.getenv('OPENAI_API_KEY')
client = OpenAI(api_key=openai_api_key) if openai_api_key else None

@ai_reports_bp.route('/reports/<int:report_id>/ai-suggestions', methods=['POST'])
@jwt_required()
def generate_ai_suggestions(report_id):
    """Generate AI-powered suggestions for improving a report"""
    try:
        current_user_id = get_jwt_identity()
        
        # Get the report
        report = Report.query.filter_by(id=report_id, user_id=current_user_id).first()
        if not report:
            return jsonify({'error': 'Report not found'}), 404
        
        data = request.get_json()
        content = data.get('content', '')
        report_type = data.get('type', 'summary')
        data_source = data.get('data_source', 'unknown')
        
        # Prepare prompt for OpenAI
        prompt = f"""
        As an expert data analyst and report writer, analyze the following report and provide specific suggestions for improvement:

        Report Type: {report_type}
        Data Source: {data_source}
        Current Content:
        {content}

        Please provide 3-5 specific suggestions in the following categories:
        1. Content improvements (adding missing insights, better explanations)
        2. Structure improvements (organization, flow, readability)
        3. Visualization suggestions (charts, graphs, tables that would help)
        4. Key insights that might be missing
        
        Format your response as a JSON array with objects containing:
        - type: 'content' | 'structure' | 'visualization' | 'insight'
        - suggestion: detailed suggestion text
        - confidence: confidence score from 0.0 to 1.0
        - reasoning: explanation of why this suggestion would improve the report
        """
        
        # Check if OpenAI client is available
        if not client:
            return jsonify({
                'success': False,
                'error': 'AI service is not configured. Please set OPENAI_API_KEY environment variable.'
            }), 503
            
        # Call OpenAI API
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert data analyst and report writer who provides actionable suggestions for improving reports."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.7
        )
        
        # Parse the response
        ai_response = response.choices[0].message.content or ""
        
        # Try to parse as JSON, fallback to structured text if needed
        try:
            suggestions = json.loads(ai_response)
        except json.JSONDecodeError:
            # Fallback: create structured suggestions from text
            suggestions = [
                {
                    "type": "content",
                    "suggestion": ai_response[:200] + "..." if len(ai_response) > 200 else ai_response,
                    "confidence": 0.8,
                    "reasoning": "AI-generated content improvement suggestion"
                }
            ]
        
        # Update report with AI suggestions
        if not report.ai_suggestions:
            report.ai_suggestions = []
        
        report.ai_suggestions.extend(suggestions)
        report.generated_by_ai = True
        report.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify(suggestions)
        
    except Exception as e:
        print(f"Error generating AI suggestions: {str(e)}")
        return jsonify({'error': 'Failed to generate AI suggestions'}), 500

@ai_reports_bp.route('/reports/<int:report_id>/download', methods=['POST'])
@jwt_required()
def download_report(report_id):
    """Download report in specified format"""
    try:
        current_user_id = get_jwt_identity()
        
        # Get the report
        report = Report.query.filter_by(id=report_id, user_id=current_user_id).first()
        if not report:
            return jsonify({'error': 'Report not found'}), 404
        
        data = request.get_json()
        format_type = data.get('format', 'pdf').lower()
        
        if format_type == 'pdf':
            return generate_pdf_report(report)
        elif format_type == 'word':
            return generate_word_report(report)
        elif format_type == 'excel':
            return generate_excel_report(report)
        elif format_type == 'html':
            return generate_html_report(report)
        else:
            return jsonify({'error': 'Unsupported format'}), 400
            
    except Exception as e:
        print(f"Error downloading report: {str(e)}")
        return jsonify({'error': 'Failed to download report'}), 500

def generate_pdf_report(report):
    """Generate PDF version of the report"""
    buffer = BytesIO()
    
    # Create PDF document
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        textColor=colors.darkblue
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12,
        textColor=colors.darkblue
    )
    
    # Build document content
    content = []
    
    # Title
    content.append(Paragraph(report.title, title_style))
    content.append(Spacer(1, 12))
    
    # Description
    if report.description:
        content.append(Paragraph("Description", heading_style))
        content.append(Paragraph(report.description, styles['Normal']))
        content.append(Spacer(1, 12))
    
    # Report content
    content.append(Paragraph("Report Content", heading_style))
    
    # Split content into paragraphs and add them
    paragraphs = report.content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            content.append(Paragraph(para.strip(), styles['Normal']))
            content.append(Spacer(1, 6))
    
    # Metadata table
    content.append(Spacer(1, 20))
    content.append(Paragraph("Report Information", heading_style))
    
    metadata = [
        ['Type', report.type.title()],
        ['Status', report.status.title()],
        ['Data Source', report.data_source.replace('_', ' ').title()],
        ['Created', report.created_at.strftime('%Y-%m-%d %H:%M')],
        ['Updated', report.updated_at.strftime('%Y-%m-%d %H:%M')],
    ]
    
    table = Table(metadata, colWidths=[2*inch, 3*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    content.append(table)
    
    # Build PDF
    doc.build(content)
    
    # Return file
    buffer.seek(0)
    return buffer.getvalue(), 200, {
        'Content-Type': 'application/pdf',
        'Content-Disposition': f'attachment; filename="{report.title}.pdf"'
    }

def generate_word_report(report):
    """Generate Word document version of the report"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(report.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add description
    if report.description:
        doc.add_heading('Description', level=1)
        doc.add_paragraph(report.description)
    
    # Add content
    doc.add_heading('Report Content', level=1)
    
    # Split content and add paragraphs
    paragraphs = report.content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    # Add metadata
    doc.add_heading('Report Information', level=1)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    metadata = [
        ('Type', report.type.title()),
        ('Status', report.status.title()),
        ('Data Source', report.data_source.replace('_', ' ').title()),
        ('Created', report.created_at.strftime('%Y-%m-%d %H:%M')),
        ('Updated', report.updated_at.strftime('%Y-%m-%d %H:%M')),
    ]
    
    for i, (key, value) in enumerate(metadata):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = value
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue(), 200, {
        'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'Content-Disposition': f'attachment; filename="{report.title}.docx"'
    }

def generate_excel_report(report):
    """Generate Excel version of the report"""
    buffer = BytesIO()
    
    # Create workbook with pandas
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        # Report metadata
        metadata_df = pd.DataFrame([
            ['Title', report.title],
            ['Description', report.description or ''],
            ['Type', report.type.title()],
            ['Status', report.status.title()],
            ['Data Source', report.data_source.replace('_', ' ').title()],
            ['Created', report.created_at.strftime('%Y-%m-%d %H:%M')],
            ['Updated', report.updated_at.strftime('%Y-%m-%d %H:%M')],
        ], columns=['Field', 'Value'])
        
        metadata_df.to_excel(writer, sheet_name='Report Info', index=False)
        
        # Report content (split into lines)
        content_lines = report.content.split('\n')
        content_df = pd.DataFrame(content_lines, columns=['Content'])
        content_df.to_excel(writer, sheet_name='Content', index=False)
        
        # If there are metrics, add them
        if hasattr(report, 'metrics') and report.metrics:
            metrics_data = []
            for key, value in report.metrics.items():
                metrics_data.append([key.replace('_', ' ').title(), value])
            
            metrics_df = pd.DataFrame(metrics_data, columns=['Metric', 'Value'])
            metrics_df.to_excel(writer, sheet_name='Metrics', index=False)
    
    buffer.seek(0)
    
    return buffer.getvalue(), 200, {
        'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'Content-Disposition': f'attachment; filename="{report.title}.xlsx"'
    }

def generate_html_report(report):
    """Generate HTML version of the report"""
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{report.title}</title>
        <style>
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                line-height: 1.6;
                color: #333;
            }}
            .header {{
                border-bottom: 3px solid #007acc;
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            h1 {{
                color: #007acc;
                margin-bottom: 10px;
            }}
            .description {{
                font-style: italic;
                color: #666;
                margin-bottom: 20px;
            }}
            .content {{
                background: #f9f9f9;
                padding: 20px;
                border-radius: 5px;
                margin: 20px 0;
                white-space: pre-wrap;
            }}
            .metadata {{
                background: #e3f2fd;
                padding: 15px;
                border-radius: 5px;
                margin-top: 30px;
            }}
            .metadata table {{
                width: 100%;
                border-collapse: collapse;
            }}
            .metadata th, .metadata td {{
                padding: 8px;
                text-align: left;
                border-bottom: 1px solid #ddd;
            }}
            .metadata th {{
                background-color: #007acc;
                color: white;
            }}
            .status {{
                display: inline-block;
                padding: 4px 8px;
                border-radius: 3px;
                font-size: 0.9em;
                font-weight: bold;
            }}
            .status.completed {{ background: #4caf50; color: white; }}
            .status.draft {{ background: #ff9800; color: white; }}
            .status.generating {{ background: #2196f3; color: white; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{report.title}</h1>
            {f'<div class="description">{report.description}</div>' if report.description else ''}
        </div>
        
        <div class="content">
            {report.content}
        </div>
        
        <div class="metadata">
            <h3>Report Information</h3>
            <table>
                <tr><th>Field</th><th>Value</th></tr>
                <tr><td>Type</td><td>{report.type.title()}</td></tr>
                <tr><td>Status</td><td><span class="status {report.status}">{report.status.title()}</span></td></tr>
                <tr><td>Data Source</td><td>{report.data_source.replace('_', ' ').title()}</td></tr>
                <tr><td>Created</td><td>{report.created_at.strftime('%Y-%m-%d %H:%M')}</td></tr>
                <tr><td>Updated</td><td>{report.updated_at.strftime('%Y-%m-%d %H:%M')}</td></tr>
            </table>
        </div>
        
        <div style="margin-top: 30px; text-align: center; color: #666; font-size: 0.9em;">
            Generated by Automated Report System
        </div>
    </body>
    </html>
    """
    
    return html_content.encode('utf-8'), 200, {
        'Content-Type': 'text/html; charset=utf-8',
        'Content-Disposition': f'attachment; filename="{report.title}.html"'
    }

@ai_reports_bp.route('/reports/<int:report_id>/enhance', methods=['POST'])
@jwt_required()
def enhance_report_with_ai(report_id):
    """Enhance existing report content with AI"""
    try:
        current_user_id = get_jwt_identity()
        
        # Get the report
        report = Report.query.filter_by(id=report_id, user_id=current_user_id).first()
        if not report:
            return jsonify({'error': 'Report not found'}), 404
        
        data = request.get_json()
        enhancement_type = data.get('type', 'general')  # general, insights, formatting, summary
        
        # Prepare enhancement prompt
        if enhancement_type == 'insights':
            prompt = f"""
            Analyze the following report and add 3-5 key insights that readers should know:
            
            {report.content}
            
            Add an "Key Insights" section with bullet points highlighting:
            - Important trends or patterns
            - Significant findings
            - Actionable recommendations
            - Potential concerns or opportunities
            """
        elif enhancement_type == 'formatting':
            prompt = f"""
            Improve the formatting and structure of this report for better readability:
            
            {report.content}
            
            Please:
            - Add appropriate headings and subheadings
            - Structure the content logically
            - Add bullet points where appropriate
            - Ensure clear flow and organization
            """
        elif enhancement_type == 'summary':
            prompt = f"""
            Add an executive summary at the beginning of this report:
            
            {report.content}
            
            Create a concise 2-3 paragraph executive summary that:
            - Highlights the main purpose and scope
            - Summarizes key findings
            - States main conclusions or recommendations
            """
        else:
            prompt = f"""
            Enhance and improve the following report by:
            - Adding clarity and depth to existing content
            - Ensuring logical flow and structure
            - Adding relevant insights and analysis
            - Improving readability and professionalism
            
            Original content:
            {report.content}
            """
        
        # Check if OpenAI client is available
        if not client:
            return jsonify({
                'success': False,
                'error': 'AI service is not configured. Please set OPENAI_API_KEY environment variable.'
            }), 503
        
        # Call OpenAI API
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert report writer who enhances and improves professional reports."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.7
        )
        
        # Get enhanced content
        enhanced_content = response.choices[0].message.content or ""
        
        # Update report
        report.content = enhanced_content
        report.generated_by_ai = True
        report.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify({
            'enhanced_content': enhanced_content,
            'message': 'Report enhanced successfully'
        })
        
    except Exception as e:
        print(f"Error enhancing report: {str(e)}")
        return jsonify({'error': 'Failed to enhance report'}), 500
