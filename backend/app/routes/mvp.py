from flask import Blueprint, jsonify, request, send_from_directory, current_app
from docxtpl import DocxTemplate
import os
import uuid
from docx import Document
import base64
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from app.services.ai_service import ai_service
from app.services.template_optimizer import TemplateOptimizerService
import logging

# Set up logging
logger = logging.getLogger(__name__)

mvp = Blueprint('mvp', __name__)

TEMPLATE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../templates'))

@mvp.route('/ai/analyze', methods=['POST'])
def ai_analyze():
    """
    Enhanced AI analysis endpoint using OpenAI for comprehensive data analysis
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'success': False,
                'error': 'No data provided for analysis',
                'summary': 'Analysis failed - missing data',
                'insights': [],
                'suggestions': ['Please provide data to analyze']
            }), 400
        
        # Extract context and analysis type from request
        context = data.get('context', 'general')
        analysis_data = data.get('data', data)
        
        # Log the analysis request
        logger.info(f"AI analysis requested for context: {context}")
        
        # Call enhanced AI service
        result = ai_service.analyze_data(analysis_data, context)
        
        # Ensure the response includes all expected fields for frontend compatibility
        response = {
            'success': result.get('success', True),
            'summary': result.get('summary', 'Analysis completed'),
            'insights': result.get('insights', []),
            'patterns': result.get('patterns', []),
            'anomalies': result.get('anomalies', []),
            'recommendations': result.get('recommendations', []),
            'risks': result.get('risks', []),
            'opportunities': result.get('opportunities', []),
            'confidence_score': result.get('confidence_score', 0.8),
            'data_quality': result.get('data_quality', 'medium'),
            'analysis_type': context,
            'timestamp': result.get('timestamp'),
            'suggestions': result.get('recommendations', [])  # For backward compatibility
        }
        
        if not result.get('success', True):
            response['error'] = result.get('error', 'Unknown error occurred')
            return jsonify(response), 500
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error in AI analysis endpoint: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'summary': 'Analysis failed due to technical error',
            'insights': ['Technical error prevented analysis completion'],
            'suggestions': ['Please check your data format and try again'],
            'patterns': [],
            'anomalies': [],
            'recommendations': [],
            'risks': ['Analysis incomplete'],
            'opportunities': [],
            'confidence_score': 0.0
        }), 500

@mvp.route('/ai/report-suggestions', methods=['POST'])
def ai_report_suggestions():
    """
    Generate AI-powered suggestions for report content based on data
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'success': False,
                'error': 'No data provided for suggestions'
            }), 400
        
        report_type = data.get('report_type', 'general')
        report_data = data.get('data', data)
        
        logger.info(f"AI report suggestions requested for type: {report_type}")
        
        result = ai_service.generate_report_suggestions(report_data, report_type)
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error in AI report suggestions: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'suggestions': {}
        }), 500

@mvp.route('/ai/optimize-template', methods=['POST'])
def ai_optimize_template():
    """
    Use AI to analyze and optimize document templates
    """
    try:
        data = request.get_json()
        
        if not data or not data.get('template_content'):
            return jsonify({
                'success': False,
                'error': 'Template content is required'
            }), 400
        
        template_content = data.get('template_content', '')
        placeholders = data.get('placeholders', [])
        template_type = data.get('template_type', 'general')
        
        logger.info(f"AI template optimization requested for type: {template_type}")
        
        result = ai_service.optimize_template(template_content, placeholders, template_type)
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error in AI template optimization: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'optimization': {}
        }), 500

@mvp.route('/ai/optimize-template-with-excel', methods=['POST'])
def ai_optimize_template_with_excel():
    """
    Enhanced template optimization using Excel data for comprehensive analysis and mapping.
    Extracts all data from Excel files and maps to template placeholders.
    """
    try:
        # Check if files are present
        if 'template_file' not in request.files and 'excel_file' not in request.files:
            return jsonify({
                'success': False,
                'error': 'Both template and Excel files are required'
            }), 400
        
        template_file = request.files.get('template_file')
        excel_file = request.files.get('excel_file')
        
        if not template_file or not excel_file:
            return jsonify({
                'success': False,
                'error': 'Both template and Excel files must be provided'
            }), 400
        
        # Save files temporarily
        import tempfile
        temp_dir = tempfile.mkdtemp()
        
        template_path = os.path.join(temp_dir, 'template.tex')
        excel_path = os.path.join(temp_dir, 'data.xlsx')
        
        template_file.save(template_path)
        excel_file.save(excel_path)
        
        # Read template content
        with open(template_path, 'r', encoding='utf-8') as f:
            template_content = f.read()
        
        # Initialize template optimizer
        optimizer = TemplateOptimizerService()
        
        # Optimize template with Excel data
        result = optimizer.optimize_template_with_excel(template_content, excel_path)
        
        # Clean up temporary files
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        if result['success']:
            return jsonify({
                'success': True,
                'template_analysis': result['placeholders'],
                'data_extraction': result['data_analysis'],
                'context': result['enhanced_context'],
                'missing_fields': result['missing_fields'],
                'optimizations': result['optimizations'],
                'message': 'Template optimized successfully with Excel data'
            })
        else:
            return jsonify({
                'success': False,
                'error': result['error']
            }), 500
        
    except Exception as e:
        logger.error(f"Error in enhanced template optimization: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'message': 'Failed to optimize template with Excel data'
        }), 500

@mvp.route('/templates/generate-with-excel', methods=['POST'])
def generate_report_with_excel():
    """
    Generate report using template and Excel data with intelligent mapping.
    """
    try:
        # Log request payload for diagnostics
        try:
            current_app.logger.info(
                "generate-with-excel: files=%s form=%s",
                list(request.files.keys()),
                list(request.form.keys()),
            )
        except Exception:
            pass

        # Accept multiple parameter names for better UX
        template_file = request.files.get('template_file') or request.files.get('template') or request.files.get('file')
        excel_file = request.files.get('excel_file') or request.files.get('excel') or request.files.get('data')

        # Validate presence
        if not template_file:
            return jsonify({
                'success': False,
                'error_code': 'template_missing',
                'message': 'Template file is required',
                'details': {'expected_keys': ['template_file', 'template']},
                'suggestions': ['Attach a DOCX or TEX template as form-data with key "template_file".'],
                'retry_possible': True
            }), 400

        if not excel_file:
            return jsonify({
                'success': False,
                'error_code': 'excel_missing',
                'message': 'Excel file is required',
                'details': {'expected_keys': ['excel_file', 'excel']},
                'suggestions': ['Attach an Excel .xlsx file as form-data with key "excel_file".'],
                'retry_possible': True
            }), 400

        # Basic file-type validation
        template_name = (template_file.filename or '').lower()
        excel_name = (excel_file.filename or '').lower()

        if not template_name.endswith(('.docx', '.tex')):
            return jsonify({
                'success': False,
                'error_code': 'invalid_template',
                'message': 'Template must be DOCX or TEX format',
                'details': {'filename': template_file.filename},
                'suggestions': ['Use a .docx Word template with Jinja placeholders or a .tex LaTeX template.'],
                'retry_possible': True
            }), 400

        if not excel_name.endswith(('.xlsx', '.xls')):
            return jsonify({
                'success': False,
                'error_code': 'invalid_excel',
                'message': 'Excel file must be .xlsx or .xls',
                'details': {'filename': excel_file.filename},
                'suggestions': ['Export your sheet as .xlsx and try again.'],
                'retry_possible': True
            }), 400
        
        # Save files temporarily
        import tempfile
        temp_dir = tempfile.mkdtemp()
        
        template_filename = template_file.filename or 'template.tex'
        template_path = os.path.join(temp_dir, f'template{os.path.splitext(template_filename)[1]}')
        excel_path = os.path.join(temp_dir, 'data.xlsx')
        
        template_file.save(template_path)
        excel_file.save(excel_path)
        
        # Determine template type
        filename = template_file.filename or 'template.tex'
        file_ext = os.path.splitext(filename)[1].lower()
        
        # Initialize optimizer and extract data
        optimizer = TemplateOptimizerService()
        
        # Read template content (text-based only). For DOCX, skip reading as text.
        if file_ext == '.docx':
            template_content = ''
        else:
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()
        
        # Get optimized context
        optimization_result = optimizer.optimize_template_with_excel(template_content, excel_path)
        
        if not optimization_result['success']:
            return jsonify({
                'success': False,
                'error': f"Data extraction failed: {optimization_result['error']}"
            }), 500
        
        context = optimization_result['enhanced_context']
        
        # Generate output file
        output_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../static/generated'))
        os.makedirs(output_dir, exist_ok=True)
        
        if file_ext == '.tex':
            # LaTeX template processing
            output_filename = f"report_{uuid.uuid4().hex}.pdf"
            output_path = os.path.join(output_dir, output_filename)
            
            # Use Jinja2 to render LaTeX with tolerant undefined handling and safe fallbacks
            from jinja2 import Environment, StrictUndefined
            
            # Custom undefined class that provides helpful error messages
            class HelpfulUndefined(StrictUndefined):
                def __str__(self):
                    hint = ""
                    if hasattr(self, '_undefined_name') and self._undefined_name:
                        name = str(self._undefined_name)
                        if name == 'p' or name.startswith('p.'):
                            hint = " (Hint: 'p' variables should be inside {% for p in participants %} loops)"
                    return f"[UNDEFINED: {getattr(self, '_undefined_name', 'unknown')}{hint}]"
            
            env = Environment(
                autoescape=False, 
                trim_blocks=True, 
                lstrip_blocks=True, 
                undefined=HelpfulUndefined
            )
            
            try:
                template = env.from_string(template_content)
                rendered_latex = template.render(context)
            except Exception as render_err:
                current_app.logger.exception("LaTeX render failed; attempting syntax conversion and safe context")
                try:
                    # Convert Mustache syntax to Jinja2 if needed and prepare safe context
                    from app.services.template_converter import TemplateConverter

                    syntax_type = TemplateConverter.analyze_template_syntax(template_content)
                    if syntax_type in ['mustache', 'mixed']:
                        converted_content = TemplateConverter.mustache_to_jinja2(template_content)
                        prepared_context = TemplateConverter.prepare_context_for_mustache_conversion(context)
                    else:
                        converted_content = template_content
                        prepared_context = context

                    # Generate safe defaults for any missing placeholders
                    safe_defaults = TemplateConverter.generate_safe_context_from_template(converted_content)

                    # Deep-merge user/prepared context into safe defaults
                    def deep_merge(base, user):
                        for k, v in user.items():
                            if k in base and isinstance(base[k], dict) and isinstance(v, dict):
                                deep_merge(base[k], v)
                            else:
                                base[k] = v
                        return base

                    final_context = deep_merge(safe_defaults, prepared_context)

                    # Use more tolerant undefined for the retry
                    from jinja2 import Undefined
                    env_tolerant = Environment(
                        autoescape=False, 
                        trim_blocks=True, 
                        lstrip_blocks=True, 
                        undefined=Undefined
                    )
                    
                    template = env_tolerant.from_string(converted_content)
                    rendered_latex = template.render(final_context)
                except Exception as second_err:
                    return jsonify({
                        'success': False,
                        'error': f"Template rendering failed: {str(render_err)}",
                        'hint': 'Check for variables used outside loops (e.g., {{ p.* }} outside {% for p in participants %} ... {% endfor %}). The variable "p" should only be used inside participant loops.',
                        'debug_info': {
                            'original_error': str(render_err),
                            'retry_error': str(second_err),
                            'template_syntax': syntax_type if 'syntax_type' in locals() else 'unknown'
                        }
                    }), 400
            
            # Save rendered LaTeX for debugging
            latex_output_path = os.path.join(output_dir, f"debug_{uuid.uuid4().hex}.tex")
            with open(latex_output_path, 'w', encoding='utf-8') as f:
                f.write(rendered_latex)
            
            # For now, return the rendered LaTeX content
            # In production, you would compile this to PDF
            download_url = f"/mvp/static/generated/{os.path.basename(latex_output_path)}"
            
            # Clean up temporary files
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            return jsonify({
                'success': True,
                'downloadUrl': download_url,
                'message': 'Report generated successfully with Excel data!',
                'filename': os.path.basename(latex_output_path),
                'context_used': context,
                'optimizations': optimization_result.get('optimizations', {}),
                'missing_fields': optimization_result.get('missing_fields', [])
            })
            
        elif file_ext == '.docx':
            # Word template processing
            output_filename = f"report_{uuid.uuid4().hex}.docx"
            output_path = os.path.join(output_dir, output_filename)

            try:
                # Use docxtpl for Word templates
                from docxtpl import DocxTemplate

                # Log a concise context summary for diagnostics
                try:
                    current_app.logger.info(
                        "Excel->Context summary: program_title=%s, participants=%d, keys=%s",
                        (context.get('program', {}) or {}).get('title'),
                        len(context.get('participants', []) or []),
                        list(context.keys()),
                    )
                except Exception:
                    pass

                doc = DocxTemplate(template_path)
                # Use a tolerant Jinja2 environment so missing vars render as empty strings instead of raising
                from jinja2 import Environment, Undefined
                env = Environment(autoescape=False, trim_blocks=True, lstrip_blocks=True, undefined=Undefined)

                try:
                    doc.render(context, jinja_env=env)
                except Exception as render_err:
                    current_app.logger.exception("docxtpl rendering failed - retrying with sanitized context")
                    # Retry with sanitized context and tolerant env
                    try:
                        safe_context = context or {}
                        if 'participants' not in safe_context:
                            safe_context['participants'] = []
                        if 'program' not in safe_context:
                            safe_context['program'] = {}
                        doc.render(safe_context, jinja_env=env)
                    except Exception as second_err:
                        return jsonify({
                            'success': False,
                            'error': f"Template rendering failed: {str(render_err)}",
                            'hint': "Check for variables used outside their loops (e.g., using {{ p.* }} outside {% for p in participants %} ... {% endfor %}). Missing variables now render as empty strings.",
                        }), 400

                try:
                    doc.save(output_path)
                except Exception as save_err:
                    current_app.logger.exception("Saving rendered DOCX failed")
                    return jsonify({
                        'success': False,
                        'error': f"Failed to save generated DOCX: {str(save_err)}"
                    }), 500

                download_url = f"/mvp/static/generated/{output_filename}"

                # Clean up temporary files
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)

                return jsonify({
                    'success': True,
                    'downloadUrl': download_url,
                    'message': 'Report generated successfully with Excel data!',
                    'filename': output_filename,
                    'context_used': context,
                    'context_summary': {
                        'program_title': (context.get('program', {}) or {}).get('title'),
                        'participants_count': len(context.get('participants', []) or [])
                    },
                    'optimizations': optimization_result.get('optimizations', {}),
                    'missing_fields': optimization_result.get('missing_fields', [])
                })
            except Exception as e:
                current_app.logger.exception("Error generating DOCX from Excel")
                return jsonify({
                    'success': False,
                    'error': f"Server error during DOCX generation: {str(e)}"
                }), 500
        
        else:
            return jsonify({
                'success': False,
                'error': 'Unsupported template format. Use .tex or .docx files.'
            }), 400
        
    except Exception as e:
        logger.error(f"Error generating report with Excel: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'message': 'Failed to generate report with Excel data'
        }), 500

@mvp.route('/ai/validate-data', methods=['POST'])
def ai_validate_data():
    """
    AI-powered data quality validation and suggestions
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'success': False,
                'error': 'No data provided for validation'
            }), 400
        
        validation_data = data.get('data', data)
        
        logger.info("AI data validation requested")
        
        result = ai_service.validate_data_quality(validation_data)
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error in AI data validation: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'validation': {
                'overall_quality_score': 0.0,
                'data_readiness': 'unknown',
                'issues_found': []
            }
        }), 500

@mvp.route('/ai/smart-placeholders', methods=['POST'])
def ai_smart_placeholders():
    """
    Generate intelligent placeholder suggestions based on context and industry
    """
    try:
        data = request.get_json()
        
        context = data.get('context', 'general') if data else 'general'
        industry = data.get('industry', 'general') if data else 'general'
        
        logger.info(f"AI placeholder suggestions requested for context: {context}, industry: {industry}")
        
        result = ai_service.generate_smart_placeholders(context, industry)
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error generating smart placeholders: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'placeholders': {}
        }), 500

@mvp.route('/ai/health', methods=['GET'])
def ai_health_check():
    """
    Check AI service health and API key status
    """
    try:
        # Check if OpenAI API key is configured
        openai_configured = bool(ai_service.openai_api_key and ai_service.openai_api_key.startswith('sk-'))
        google_ai_configured = bool(ai_service.google_ai_api_key)
        
        return jsonify({
            'success': True,
            'status': 'healthy',
            'services': {
                'openai': {
                    'configured': openai_configured,
                    'status': 'ready' if openai_configured else 'not_configured'
                },
                'google_ai': {
                    'configured': google_ai_configured,
                    'status': 'ready' if google_ai_configured else 'not_configured'
                }
            },
            'features_available': [
                'data_analysis',
                'report_suggestions',
                'template_optimization',
                'data_validation',
                'smart_placeholders'
            ]
        })
        
    except Exception as e:
        logger.error(f"Error in AI health check: {str(e)}")
        return jsonify({
            'success': False,
            'status': 'unhealthy',
            'error': str(e)
        }), 500

@mvp.route('/templates/<template_name>/placeholders', methods=['GET'])
def extract_placeholders_from_stored(template_name):
    """
    Returns a JSON list of unique placeholders found in templates stored on the server.
    Supports both .docx and .tex files.
    Example: GET /mvp/templates/Temp2.tex/placeholders
    """
    # Sanitize filename to prevent directory traversal
    safe_name = os.path.basename(template_name)
    template_path = os.path.join(TEMPLATE_DIR, safe_name)

    if not os.path.isfile(template_path):
        return jsonify({'error': 'Template not found'}), 404

    try:
        # Determine file type and handle accordingly
        file_ext = os.path.splitext(safe_name)[1].lower()
        
        if file_ext == '.tex':
            # Handle LaTeX templates
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()
            
            # Use the LaTeX template analyzer
            from app.services.template_optimizer import LaTeXTemplateAnalyzer
            analyzer = LaTeXTemplateAnalyzer()
            placeholders_analysis = analyzer.extract_all_placeholders(template_content)
            
            # Flatten all placeholder types into a single list
            all_placeholders = []
            all_placeholders.extend(placeholders_analysis.get('simple', []))
            all_placeholders.extend(placeholders_analysis.get('nested', []))
            
            # Add loop variables
            for loop in placeholders_analysis.get('loops', []):
                all_placeholders.append(f"#{loop['variable']}")
                all_placeholders.extend(loop['inner_placeholders'])
            
            # Add table placeholders
            for table in placeholders_analysis.get('tables', []):
                all_placeholders.extend(table['placeholders'])
            
            return jsonify({
                'placeholders': list(set(all_placeholders)),
                'analysis': placeholders_analysis,
                'file_type': 'latex',
                'total_count': len(set(all_placeholders))
            })
            
        elif file_ext == '.docx':
            # Handle Word documents (existing logic)
            doc = Document(template_path)
            
            # Extract placeholders manually from the document content
            import re
            placeholder_pattern = r'\{\{([^}]+)\}\}'
            all_placeholders = set()
            
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    placeholders = re.findall(placeholder_pattern, paragraph.text)
                    all_placeholders.update(placeholders)
            
            # Also check table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text:
                                placeholders = re.findall(placeholder_pattern, paragraph.text)
                                all_placeholders.update(placeholders)
            
            return jsonify({
                'placeholders': list(all_placeholders),
                'file_type': 'word',
                'total_count': len(all_placeholders)
            })
        
        else:
            return jsonify({'error': f'Unsupported file type: {file_ext}'}), 400
        
    except Exception as e:
        logger.error(f"Error extracting placeholders from {safe_name}: {str(e)}")
        return jsonify({'error': f'Failed to process template: {str(e)}'}), 500

@mvp.route('/templates/<template_name>/content', methods=['GET'])
def extract_template_content(template_name):
    """
    Extracts the raw content from a Word template for frontend editing.
    Preserves formatting information and structure.
    """
    # Sanitize filename to prevent directory traversal
    safe_name = os.path.basename(template_name)
    template_path = os.path.join(TEMPLATE_DIR, safe_name)

    if not os.path.isfile(template_path):
        return jsonify({'error': 'Template not found'}), 404

    try:
        doc = Document(template_path)
        content = []
        
        import re
        placeholder_pattern = r'\{\{([^}]+)\}\}'
        all_placeholders = set()
        
        # Process paragraphs with formatting information
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Extract formatting information
                formatting = {
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'alignment': paragraph.alignment,
                    'font_size': None,
                    'font_name': None,
                    'bold': False,
                    'italic': False
                }
                
                # Try to extract font information from runs
                for run in paragraph.runs:
                    if run.font.size:
                        formatting['font_size'] = run.font.size.pt
                    if run.font.name:
                        formatting['font_name'] = run.font.name
                    if run.bold:
                        formatting['bold'] = True
                    if run.italic:
                        formatting['italic'] = True
                
                # Convert paragraph to HTML-like structure with formatting
                text = paragraph.text
                # Replace placeholders with editable spans
                text = re.sub(placeholder_pattern, r'<span class="placeholder" data-placeholder="\1">{{{\1}}}</span>', text)
                
                content.append({
                    'type': 'paragraph',
                    'text': text,
                    'formatting': formatting,
                    'style': paragraph.style.name if paragraph.style else 'Normal'
                })
                
                # Extract placeholders
                placeholders = re.findall(placeholder_pattern, paragraph.text)
                all_placeholders.update(placeholders)
        
        # Process tables with structure preservation
        for table_index, table in enumerate(doc.tables):
            table_content = {
                'type': 'table',
                'index': table_index,
                'rows': [],
                'formatting': {
                    'style': 'Table Grid',
                    'alignment': 'left'
                }
            }
            
            for row_index, row in enumerate(table.rows):
                row_data = []
                for cell_index, cell in enumerate(row.cells):
                    cell_text = ''
                    cell_formatting = {
                        'bold': False,
                        'italic': False,
                        'font_size': None,
                        'font_name': None
                    }
                    
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text = paragraph.text
                            # Replace placeholders with editable spans
                            text = re.sub(placeholder_pattern, r'<span class="placeholder" data-placeholder="\1">{{{\1}}}</span>', text)
                            cell_text += text + ' '
                            
                            # Extract formatting from runs
                            for run in paragraph.runs:
                                if run.font.size:
                                    cell_formatting['font_size'] = run.font.size.pt
                                if run.font.name:
                                    cell_formatting['font_name'] = run.font.name
                                if run.bold:
                                    cell_formatting['bold'] = True
                                if run.italic:
                                    cell_formatting['italic'] = True
                    
                    row_data.append({
                        'text': cell_text.strip(),
                        'formatting': cell_formatting,
                        'row': row_index,
                        'col': cell_index
                    })
                    
                    # Extract placeholders from cell text
                    placeholders = re.findall(placeholder_pattern, cell_text)
                    all_placeholders.update(placeholders)
                
                table_content['rows'].append(row_data)
            
            content.append(table_content)
        
        # Add template metadata
        template_info = {
            'filename': safe_name,
            'total_paragraphs': len([item for item in content if item['type'] == 'paragraph']),
            'total_tables': len([item for item in content if item['type'] == 'table']),
            'total_placeholders': len(all_placeholders),
            'page_count': len(doc.sections) if hasattr(doc, 'sections') else 1
        }
        
        return jsonify({
            'content': content,
            'placeholders': list(all_placeholders),
            'template_info': template_info
        })
    except Exception as e:
        return jsonify({'error': f'Failed to extract template content: {str(e)}'}), 500

@mvp.route('/templates/<template_name>/preview', methods=['POST'])
def generate_live_preview(template_name):
    """
    Generates a live preview of the filled template as PDF.
    Preserves original template formatting and styles.
    """
    import logging
    logger = logging.getLogger(__name__)
    
    # Sanitize filename to prevent directory traversal
    safe_name = os.path.basename(template_name)
    template_path = os.path.join(TEMPLATE_DIR, safe_name)
    
    logger.info(f"Attempting to generate preview for template: {safe_name}")
    logger.info(f"Template path: {template_path}")

    if not os.path.isfile(template_path):
        logger.error(f"Template not found at path: {template_path}")
        return jsonify({'error': 'Template not found'}), 404

    data = request.get_json()
    context = data.get('data', {}) if data else {}
    
    logger.info(f"Received context data: {context}")
    
    if not context:
        logger.warning("No data provided for preview")
        return jsonify({'error': 'No data provided for preview'}), 400

    file_ext = os.path.splitext(safe_name)[1].lower()
    if file_ext == '.docx':
        try:
            # Create a styled PDF that mimics the template using only reportlab
            buffer = io.BytesIO()
            pdf_doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Create custom styles to match template
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=20,
                alignment=1  # Center alignment
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12
            )
            
            # Read the original template to understand structure
            from docx import Document
            doc_original = Document(template_path)
            
            # Process paragraphs with preserved formatting
            for paragraph in doc_original.paragraphs:
                if paragraph.text.strip():
                    # Replace placeholders with actual values
                    text = paragraph.text
                    for key, value in context.items():
                        text = text.replace(f'{{{{{key}}}}}', str(value))
                    
                    if text.strip():
                        # Determine style based on paragraph properties
                        if paragraph.style and hasattr(paragraph.style, 'name') and paragraph.style.name and paragraph.style.name.startswith('Heading'):
                            p = Paragraph(text, title_style)
                        else:
                            p = Paragraph(text, normal_style)
                        story.append(p)
            
            # Process tables with preserved structure
            for table in doc_original.tables:
                from reportlab.platypus import Table, TableStyle
                from reportlab.lib import colors
                
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = ''
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                text = paragraph.text
                                for key, value in context.items():
                                    text = text.replace(f'{{{{{key}}}}}', str(value))
                                cell_text += text + ' '
                        row_data.append(cell_text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    pdf_table = Table(table_data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 20))
            
            pdf_doc.build(story)
            buffer.seek(0)
            
            # Convert to base64 for frontend display
            pdf_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
            
            logger.info("Successfully generated PDF preview using fallback method")
            
            return jsonify({
                'preview': f'data:application/pdf;base64,{pdf_base64}',
                'filename': f'preview_{uuid.uuid4().hex}.pdf'
            })
            
        except Exception as e:
            logger.error(f"Failed to generate preview: {str(e)}")
            return jsonify({'error': f'Failed to generate preview: {str(e)}'}), 500
    elif file_ext == '.tex':
        try:
            logger.info("Processing LaTeX template for preview")
            
            # Read LaTeX template
            with open(template_path, 'r', encoding='utf-8') as f:
                latex_content = f.read()
            
            logger.info(f"LaTeX template content length: {len(latex_content)}")
            logger.info(f"Context received: {list(context.keys()) if isinstance(context, dict) else type(context)}")
            
            # Import template converter
            from app.services.template_converter import TemplateConverter
            
            # Analyze and convert template syntax if needed
            syntax_type = TemplateConverter.analyze_template_syntax(latex_content)
            logger.info(f"Detected template syntax: {syntax_type}")
            
            if syntax_type in ['mustache', 'mixed']:
                logger.info("Converting Mustache syntax to Jinja2")
                latex_content = TemplateConverter.mustache_to_jinja2(latex_content)
                context = TemplateConverter.prepare_context_for_mustache_conversion(context)
            
            # Transform flat context data to nested structure
            nested_context = {}
            for key, value in context.items():
                if '.' in key:
                    parts = key.split('.')
                    current = nested_context
                    for part in parts[:-1]:
                        if part not in current:
                            current[part] = {}
                        current = current[part]
                    current[parts[-1]] = value
                else:
                    nested_context[key] = value
            
            logger.info(f"Nested context created: {list(nested_context.keys())}")
            
            # Use Jinja2 to render LaTeX with nested data and loops
            from jinja2 import Template, DebugUndefined
            
            # Use DebugUndefined to catch missing variables
            template = Template(latex_content, undefined=DebugUndefined)
            
            try:
                rendered_latex = template.render(nested_context)
                logger.info(f"LaTeX template rendered successfully. Length: {len(rendered_latex)}")
            except Exception as render_error:
                logger.error(f"Template rendering error: {str(render_error)}")
                # Try with a comprehensive safe context generated from template
                logger.info("Generating safe context from template placeholders")
                
                safe_context = TemplateConverter.generate_safe_context_from_template(latex_content)
                
                # Merge user context with safe defaults
                def deep_merge(base, user):
                    for key, value in user.items():
                        if key in base:
                            if isinstance(base[key], dict) and isinstance(value, dict):
                                deep_merge(base[key], value)
                            else:
                                base[key] = value
                        else:
                            base[key] = value
                    return base
                
                final_context = deep_merge(safe_context, nested_context)
                
                template_safe = Template(latex_content)
                rendered_latex = template_safe.render(final_context)
                logger.info("LaTeX template rendered with comprehensive safe defaults")
            
            # Create output directory if it doesn't exist
            output_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../static/generated'))
            os.makedirs(output_dir, exist_ok=True)
            
            # Save rendered LaTeX file for download
            latex_filename = f'preview_{uuid.uuid4().hex}.tex'
            latex_path = os.path.join(output_dir, latex_filename)
            with open(latex_path, 'w', encoding='utf-8') as f:
                f.write(rendered_latex)
            
            logger.info(f"LaTeX preview file saved: {latex_filename}")
            
            return jsonify({
                'success': True,
                'preview_type': 'latex',
                'download_url': f"/mvp/static/generated/{latex_filename}",
                'filename': latex_filename,
                'content_preview': rendered_latex[:1000] + "..." if len(rendered_latex) > 1000 else rendered_latex,
                'message': 'LaTeX preview generated successfully'
            })
                
        except Exception as latex_error:
            logger.error(f"Failed to process LaTeX template: {str(latex_error)}", exc_info=True)
            return jsonify({
                'error': f'Failed to generate LaTeX preview: {str(latex_error)}',
                'preview_type': 'error'
            }), 500
    else:
        return jsonify({'error': 'Unsupported template format for preview'}), 400

@mvp.route('/templates/list', methods=['GET'])
def list_templates():
    files = []
    for fname in os.listdir(TEMPLATE_DIR):
        if fname.lower().endswith(('.docx', '.tex')):
            files.append({
                'id': fname,
                'name': os.path.splitext(fname)[0],
                'description': '',
                'filename': fname
            })
    return jsonify(files)

@mvp.route('/generate-report', methods=['POST'])
def generate_report():
    try:
        data = request.get_json()
        current_app.logger.debug(f"Received data: {data}")
        
        template_filename = data.get('templateFilename')
        context = data.get('data', {})
        
        current_app.logger.debug(f"Template filename: {template_filename}")
        current_app.logger.debug(f"Context data: {context}")
        
        if not template_filename or not context:
            return jsonify({'error': 'Missing template filename or data'}), 400
        
        template_path = os.path.join(TEMPLATE_DIR, template_filename)
        current_app.logger.debug(f"Template path: {template_path}")
        
        if not os.path.isfile(template_path):
            return jsonify({'error': 'Template not found'}), 404
        
        # Generate unique output filename
        output_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../static/generated'))
        os.makedirs(output_dir, exist_ok=True)
        
        # Determine file type and handle accordingly
        file_ext = os.path.splitext(template_filename)[1].lower()
        if file_ext == '.docx':
            output_filename = f"report_{uuid.uuid4().hex}.docx"
            output_path = os.path.join(output_dir, output_filename)
            
            current_app.logger.debug(f"Output path: {output_path}")
            
            # Fill template using docxtpl
            from docxtpl import DocxTemplate
            current_app.logger.debug("Creating DocxTemplate instance")
            doc = DocxTemplate(template_path)
            
            # Transform flat context data to nested structure
            nested_context = {}
            for key, value in context.items():
                if '.' in key:
                    parts = key.split('.')
                    current = nested_context
                    for part in parts[:-1]:
                        if part not in current:
                            current[part] = {}
                        current = current[part]
                    current[parts[-1]] = value
                else:
                    nested_context[key] = value
            
            current_app.logger.debug(f"Nested context: {nested_context}")
            
            current_app.logger.debug("Rendering template with context")
            doc.render(nested_context)
            
            current_app.logger.debug("Saving rendered document")
            doc.save(output_path)
            
            # Create download URL
            download_url = f"/mvp/static/generated/{output_filename}"
            
            current_app.logger.debug(f"Report generated successfully: {download_url}")
            
            return jsonify({
                'downloadUrl': download_url,
                'message': f'Report generated successfully! File: {output_filename}',
                'filename': output_filename,
                'status': 'success'
            })
        elif file_ext == '.tex':
            output_filename = f"report_{uuid.uuid4().hex}.pdf"
            output_path = os.path.join(output_dir, output_filename)
            
            # Read LaTeX template
            with open(template_path, 'r', encoding='utf-8') as f:
                latex_content = f.read()
            
            # Transform flat context data to nested structure
            nested_context = {}
            for key, value in context.items():
                if '.' in key:
                    parts = key.split('.')
                    current = nested_context
                    for part in parts[:-1]:
                        if part not in current:
                            current[part] = {}
                        current = current[part]
                    current[parts[-1]] = value
                else:
                    nested_context[key] = value
            
            # Use Jinja2 to render LaTeX with nested data and loops
            from jinja2 import Template
            template = Template(latex_content)
            rendered_latex = template.render(nested_context)
            
            # Save rendered LaTeX file (PDF compilation would require LaTeX installation)
            latex_output_path = os.path.join(output_dir, f"report_{uuid.uuid4().hex}.tex")
            with open(latex_output_path, 'w', encoding='utf-8') as f:
                f.write(rendered_latex)
            
            # Create download URL for the LaTeX file
            download_url = f"/mvp/static/generated/{os.path.basename(latex_output_path)}"
            
            return jsonify({
                'downloadUrl': download_url,
                'message': f'Report generated successfully! File: {output_filename}',
                'filename': output_filename,
                'status': 'success'
            })
        else:
            return jsonify({'error': 'Unsupported template format'}), 400
        
    except Exception as e:
        current_app.logger.error(f"Error generating report: {str(e)}", exc_info=True)
        return jsonify({
            'error': f'Failed to generate report: {str(e)}',
            'status': 'error'
        }), 500

@mvp.route('/templates/<template_name>/download', methods=['GET'])
def download_template(template_name):
    """
    Download the original template file for verification.
    """
    # Sanitize filename to prevent directory traversal
    safe_name = os.path.basename(template_name)
    template_path = os.path.join(TEMPLATE_DIR, safe_name)

    if not os.path.isfile(template_path):
        return jsonify({'error': 'Template not found'}), 404

    try:
        return send_from_directory(
            TEMPLATE_DIR, 
            safe_name, 
            as_attachment=True,
            download_name=safe_name
        )
    except Exception as e:
        return jsonify({'error': f'Failed to download template: {str(e)}'}), 500 

@mvp.route('/static/generated/<filename>', methods=['GET'])
def serve_generated_file(filename):
    """
    Serve generated report files for download.
    """
    # Sanitize filename to prevent directory traversal
    safe_filename = os.path.basename(filename)
    generated_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../static/generated'))
    file_path = os.path.join(generated_dir, safe_filename)
    
    if not os.path.isfile(file_path):
        return jsonify({'error': 'File not found'}), 404
    
    try:
        return send_from_directory(
            generated_dir,
            safe_filename,
            as_attachment=True,
            download_name=safe_filename
        )
    except Exception as e:
        return jsonify({'error': f'Failed to serve file: {str(e)}'}), 500 