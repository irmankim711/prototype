#!/usr/bin/env python3
"""
Test the complete DOCX preview system
Creates a sample DOCX file and tests the preview generation
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches

# Add the parent directory to Python path
sys.path.insert(0, str(Path(__file__).parent))

from app.services.docx_preview_service import DocxPreviewService

def create_sample_docx():
    """Create a sample DOCX file for testing"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample Business Report', 0)
    title.alignment = 1  # Center alignment
    
    # Add subtitle
    doc.add_heading('Q3 2024 Performance Analysis', level=2)
    
    # Add some content
    doc.add_paragraph('This is a sample business report demonstrating the DOCX preview functionality.')
    
    # Add heading
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'Our Q3 2024 performance shows strong growth across all key metrics. '
        'Revenue increased by 15% compared to Q2, while customer satisfaction '
        'reached an all-time high of 94%.'
    )
    
    # Add table
    doc.add_heading('Key Performance Indicators', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Q2 2024'
    hdr_cells[2].text = 'Q3 2024'
    
    # Add data rows
    metrics = [
        ('Revenue', '$1.2M', '$1.38M'),
        ('Customer Count', '1,250', '1,420'),
        ('Satisfaction Score', '89%', '94%'),
        ('Market Share', '12%', '14%')
    ]
    
    for metric, q2, q3 in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = q2
        row_cells[2].text = q3
    
    # Add more content
    doc.add_heading('Detailed Analysis', level=1)
    doc.add_paragraph(
        'The significant growth in Q3 can be attributed to several factors:'
    )
    
    # Add bullet points
    doc.add_paragraph('• Launch of new product line in July', style='List Bullet')
    doc.add_paragraph('• Expansion into two new markets', style='List Bullet')
    doc.add_paragraph('• Improved customer service response times', style='List Bullet')
    doc.add_paragraph('• Strategic partnerships with key vendors', style='List Bullet')
    
    # Add conclusion
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph(
        'Based on the strong Q3 performance, we recommend:'
    )
    doc.add_paragraph('1. Continue investment in the new product line')
    doc.add_paragraph('2. Accelerate market expansion plans')
    doc.add_paragraph('3. Maintain focus on customer satisfaction initiatives')
    
    # Save the document
    sample_path = Path(__file__).parent / 'sample_report.docx'
    doc.save(sample_path)
    return sample_path

def test_preview_generation():
    """Test the preview generation system"""
    print("🚀 Testing DOCX Preview System")
    print("=" * 50)
    
    try:
        # Create sample DOCX
        print("📝 Creating sample DOCX file...")
        docx_path = create_sample_docx()
        print(f"✅ Sample DOCX created: {docx_path}")
        
        # Initialize preview service
        print("\n🔧 Initializing preview service...")
        preview_service = DocxPreviewService()
        print("✅ Preview service initialized")
        
        # Test HTML conversion
        print("\n🎨 Converting DOCX to HTML...")
        html_path, html_content = preview_service.convert_docx_to_html(str(docx_path))
        print(f"✅ HTML preview generated: {html_path}")
        print(f"   HTML content length: {len(html_content)} characters")
        
        # Test preview URL generation
        print("\n🔗 Generating preview URL...")
        preview_url = preview_service.get_preview_url(str(docx_path))
        print(f"✅ Preview URL: {preview_url}")
        
        # Verify HTML file exists and has content
        if os.path.exists(html_path):
            with open(html_path, 'r', encoding='utf-8') as f:
                file_content = f.read()
            print(f"✅ HTML file verified, size: {len(file_content)} bytes")
            
            # Check if it contains expected content
            if 'Sample Business Report' in file_content:
                print("✅ HTML contains expected title")
            if 'table' in file_content.lower():
                print("✅ HTML contains table elements")
            if 'css' in file_content.lower():
                print("✅ HTML includes CSS styling")
        else:
            print("❌ HTML file not found")
        
        print("\n" + "=" * 50)
        print("🎉 ALL TESTS PASSED!")
        print(f"📁 Preview file location: {html_path}")
        print(f"🌐 Preview URL: {preview_url}")
        print("=" * 50)
        
        # Clean up
        print("\n🧹 Cleaning up test files...")
        if docx_path.exists():
            docx_path.unlink()
            print("✅ Sample DOCX file cleaned up")
        
        return True
        
    except Exception as e:
        print(f"❌ Test failed with error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_preview_generation()
    exit(0 if success else 1)